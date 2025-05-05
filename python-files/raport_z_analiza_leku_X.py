
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt

# Dane średnie z analizy
labels = ['Kobiety', 'Mężczyźni']
delta_sk = [13.15, 10.03]
delta_roz = [8.48, 6.46]
x = range(len(labels))

# Tworzenie wykresu
plt.figure(figsize=(10, 5))
plt.subplot(1, 2, 1)
plt.bar(x, delta_sk, tick_label=labels)
plt.title("Średni spadek ciśnienia skurczowego (Δ_sk)")
plt.ylabel("mmHg")
plt.subplot(1, 2, 2)
plt.bar(x, delta_roz, tick_label=labels, color='orange')
plt.title("Średni spadek ciśnienia rozkurczowego (Δ_roz)")
plt.ylabel("mmHg")
plt.tight_layout()
plt.savefig("Wykres_skutecznosc_lek_X.png")

# Tworzenie dokumentu Word
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

doc.add_heading("Raport: Analiza skuteczności leku X", level=1)

doc.add_heading("1. Cel badania", level=2)
doc.add_paragraph("Celem badania była ocena, czy skuteczność leku X w obniżaniu ciśnienia tętniczego różni się w zależności od BMI, płci i wieku pacjentów.")
doc.add_paragraph("Skuteczność oceniano na podstawie spadku ciśnienia tętniczego:")
doc.add_paragraph("Δ_sk = ciśnienie skurczowe przed leczeniem – po leczeniu")
doc.add_paragraph("Δ_roz = ciśnienie rozkurczowe przed leczeniem – po leczeniu")

doc.add_heading("2. Hipotezy badawcze", level=2)
doc.add_paragraph("X. Nadwaga (BMI > 25) zmniejsza skuteczność leku X.")
doc.add_paragraph("Y. Płeć nie wpływa na skuteczność leku X.")
doc.add_paragraph("Z. Wiek nie wpływa na skuteczność leku X.")

doc.add_heading("3. Podsumowanie wyników", level=2)
doc.add_paragraph("Hipoteza X – BMI a skuteczność leku:")
doc.add_paragraph("Pacjenci z BMI > 25 mieli mniejszy spadek ciśnienia skurczowego, ale różnica była bliska istotności (p = 0.064). Dla rozkurczowego: p = 0.40.")

doc.add_paragraph("Hipoteza Y – Płeć a skuteczność leku:")
doc.add_paragraph("Różnice były istotne (Δ_sk: p = 0.023, Δ_roz: p = 0.041).")
doc.add_paragraph("Kobiety: Δ_sk = 13.15 mmHg, Δ_roz = 8.48 mmHg")
doc.add_paragraph("Mężczyźni: Δ_sk = 10.03 mmHg, Δ_roz = 6.46 mmHg")

doc.add_paragraph("Hipoteza Z – Wiek a skuteczność leku:")
doc.add_paragraph("Brak istotnej korelacji (Δ_sk vs wiek: r = 0.02, p = 0.90; Δ_roz vs wiek: r = -0.13, p = 0.42)")

doc.add_heading("4. Wnioski końcowe", level=2)
doc.add_paragraph("✔ Płeć wpływa istotnie – kobiety reagują silniej na lek X.")
doc.add_paragraph("✖ BMI i wiek nie wpływają istotnie na skuteczność leku.")
doc.add_paragraph("📌 Zaleca się dalsze badania nad różnicami płciowymi w reakcji na leczenie.")

doc.save("Raport_analiza_leku_X.docx")
print("Raport Word i wykres zostały zapisane w bieżącym folderze.")
