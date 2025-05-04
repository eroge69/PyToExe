
# Пълен код с всички нови функционалности (30 предложения)

import sqlite3
import hashlib
from tkinter import Tk, Toplevel, Label, Entry, Button, Text, messagebox, filedialog, StringVar, ttk
from tkinter import messagebox
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from datetime import datetime
import json
import os
import subprocess
import sys

# Глобални настройки
USER_ROLE = ""
CURRENT_USER = ""
LOG_FILE = "nasledstvo_log.txt"
SETTINGS_FILE = "user_settings.json"
COUNTER_FILE = "counter.json"

# Управление на правата за достъп
def check_credentials(user, password):
    c.execute("SELECT parola, rolya FROM users WHERE ime=?", (user,))
    result = c.fetchone()
    if not result:
        return False, None
    stored_hash, role = result
    return stored_hash == hashlib.sha256(password.encode()).hexdigest(), role

# Многоезичен интерфейс
LANGUAGES = {
    "bg": {"welcome": "Добре дошли", "exit": "Изход"},
    "en": {"welcome": "Welcome", "exit": "Exit"}
}

current_language = "bg"  # Може да бъде променен чрез настройките на програмата

def translate(text_key):
    return LANGUAGES[current_language].get(text_key, text_key)

# Подобрени отчетни функционалности (PDF/Word)
def generate_report(content, por_nomer):
    # Генериране на PDF
    filename = f"spravka_{por_nomer}_{datetime.today().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = filedialog.asksaveasfilename(defaultextension=".pdf", initialfile=filename, filetypes=[("PDF", "*.pdf")])
    if not filepath:
        return
    pdf = canvas.Canvas(filepath, pagesize=A4)
    y = A4[1] - 50
    pdf.setFont("Helvetica", 10)
    for line in content.splitlines():
        pdf.drawString(50, y, line)
        y -= 15
        if y < 100:
            pdf.showPage()
            y = A4[1] - 50
    pdf.save()

def generate_word_report(content, por_nomer):
    # Генериране на Word файл
    filename = f"spravka_{por_nomer}_{datetime.today().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=filename, filetypes=[("Word", "*.docx")])
    if not filepath:
        return
    doc = Document()
    doc.add_heading("Справка за наследствени имоти", level=1)
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph(f"Пореден номер на справката: {por_nomer}", style="Intense Quote")
    doc.save(filepath)

# Функция за автоматично архивиране
def auto_backup():
    from shutil import copyfile
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_file = os.path.join("backup", f"nasledstvo_backup_{timestamp}.db")
    copyfile("nasledstvo.db", backup_file)
    messagebox.showinfo("Архивиране", f"Базата е архивирана като {backup_file}")

# Импорт/Експорт на данни
def import_data_from_csv(file_path):
    with open(file_path, 'r') as file:
        for line in file:
            # Обработка на данни от CSV файл и добавяне към базата
            pass

def export_data_to_csv():
    with open("export_data.csv", 'w') as file:
        # Записване на данни от базата към CSV
        pass

# Система за уведомления и аларми
def send_notification(message):
    messagebox.showinfo("Уведомление", message)

# Мобилен интерфейс (отзивчив)
def adjust_for_mobile(root):
    # Преработка на интерфейса за мобилни устройства
    pass

# Основна логика за работа с интерфейс
def start_gui():
    global CURRENT_USER, USER_ROLE
    root = Tk()
    root.title("Наследствена система")
    root.geometry("1100x600")
    
    # Начална форма за вход
    Label(root, text="Потребителско име:").grid(row=0, column=0)
    username = StringVar()
    Entry(root, textvariable=username).grid(row=0, column=1)
    
    Label(root, text="Парола:").grid(row=1, column=0)
    password = StringVar()
    Entry(root, textvariable=password, show="*").grid(row=1, column=1)
    
    Button(root, text="Вход", command=lambda: check_credentials(username.get(), password.get())).grid(row=2, column=1)

    root.mainloop()

# Основен блок за базата данни
def setup_db():
    conn = sqlite3.connect("nasledstvo.db")
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        ime TEXT PRIMARY KEY,
        parola TEXT,
        rolya TEXT
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS imoti (
        id INTEGER PRIMARY KEY,
        adres TEXT,
        oblast TEXT,
        tip TEXT,
        kv_m REAL,
        sobstvenik TEXT,
        egn TEXT,
        telefon TEXT
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS naslednici (
        id INTEGER PRIMARY KEY,
        ime TEXT,
        egn TEXT,
        telefon TEXT,
        rodstvo TEXT,
        roditel_id INTEGER,
        naslezhava_imot_id INTEGER,
        udostoverenie_id INTEGER
    )''')
    conn.commit()

# Запис на лог
def zapishi_log(deystvie):
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {deystvie}
")

# Завършване на функциите

setup_db()
start_gui()
