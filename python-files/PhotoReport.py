import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from PIL import Image

class PhotoReportApp:
    def __init__(self, master):
        self.master = master
        self.master.title("Генератор фотоотчетов ГОСТ v22.0")
        self.master.geometry("1200x800")
        self.master.minsize(800, 600)
        
        # Настройка стиля
        self.style = ttk.Style()
        self.style.theme_use('vista')
        self.style.configure('TButton', font=('Times New Roman', 10))
        self.style.configure('TLabel', font=('Times New Roman', 10))
        
        # Переменные
        self.root_folder = tk.StringVar()
        self.logo_path = tk.StringVar()
        self.company_name = tk.StringVar(value="ООО «ЛУКОЙЛ-Инжиниринг»")
        self.project_name = tk.StringVar(value="Проект освоения Командиршорской группы месторождений")
        self.report_title = tk.StringVar(value="Еженедельный фотоотчет")
        self.document_number = tk.StringVar(value="№ ______")
        self.start_date = tk.StringVar(value=datetime.now().strftime("%d.%m.%Y"))
        self.end_date = tk.StringVar(value=datetime.now().strftime("%d.%m.%Y"))
        self.city = tk.StringVar(value="Усинск")
        self.year = tk.StringVar(value=datetime.now().strftime("%Y"))
        self.last_saved_file = None
        self.image_tree = {}
        
        # Создание интерфейса
        self.create_widgets()
        self.create_menu()

    def create_menu(self):
        menubar = tk.Menu(self.master)
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Создать отчет", command=self.create_report)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.master.quit)
        menubar.add_cascade(label="Файл", menu=file_menu)
        
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="О программе", command=self.show_about)
        menubar.add_cascade(label="Помощь", menu=help_menu)
        self.master.config(menu=menubar)

    def create_widgets(self):
        main_frame = ttk.Frame(self.master, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Блок настроек
        settings_frame = ttk.LabelFrame(main_frame, text="Настройки", padding=10)
        settings_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=10, pady=10)
        
        ttk.Label(settings_frame, text="Логотип:").grid(row=0, column=0, sticky=tk.W)
        self.logo_entry = ttk.Entry(settings_frame, textvariable=self.logo_path, width=40, state='readonly')
        self.logo_entry.grid(row=0, column=1, padx=5, sticky=tk.EW)
        ttk.Button(settings_frame, text="Выбрать", command=self.select_logo).grid(row=0, column=2, padx=5)
        
        ttk.Label(settings_frame, text="Компания:").grid(row=1, column=0, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.company_name, width=50).grid(row=1, column=1, columnspan=2, sticky=tk.EW)
        
        ttk.Label(settings_frame, text="Проект:").grid(row=2, column=0, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.project_name, width=50).grid(row=2, column=1, columnspan=2, sticky=tk.EW)
        
        ttk.Label(settings_frame, text="Название отчета:").grid(row=3, column=0, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.report_title, width=50).grid(row=3, column=1, columnspan=2, sticky=tk.EW)
        
        ttk.Label(settings_frame, text="Документ №:").grid(row=4, column=0, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.document_number, width=20).grid(row=4, column=1, sticky=tk.W)
        
        ttk.Label(settings_frame, text="Период:").grid(row=5, column=0, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.start_date, width=12).grid(row=5, column=1, sticky=tk.W)
        ttk.Label(settings_frame, text="-").grid(row=5, column=1, padx=5)
        ttk.Entry(settings_frame, textvariable=self.end_date, width=12).grid(row=5, column=2, sticky=tk.E)
        
        ttk.Label(settings_frame, text="Город:").grid(row=6, column=0, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.city, width=20).grid(row=6, column=1, sticky=tk.W)
        ttk.Label(settings_frame, text="Год:").grid(row=6, column=2, sticky=tk.W)
        ttk.Entry(settings_frame, textvariable=self.year, width=8).grid(row=6, column=2, sticky=tk.E)
        
        # Статус-бар
        self.progress = ttk.Progressbar(settings_frame, orient=tk.HORIZONTAL, mode='determinate', length=300)
        self.progress.grid(row=7, column=0, columnspan=3, pady=10)
        
        # Блок кнопок
        button_frame = ttk.Frame(settings_frame)
        button_frame.grid(row=8, column=0, columnspan=3, pady=5)
        
        ttk.Button(button_frame, text="Создать отчет", command=self.create_report).pack(side=tk.LEFT, padx=5)
        self.view_btn = ttk.Button(button_frame, text="Просмотреть отчет", command=self.view_report, state=tk.DISABLED)
        self.view_btn.pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Выход", command=self.master.quit).pack(side=tk.RIGHT, padx=5)
        
        # Блок выбора папки
        folder_frame = ttk.LabelFrame(main_frame, text="Источник данных", padding=5)
        folder_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), padx=10, pady=10)
        
        ttk.Button(folder_frame, text="Выбрать корневую папку", command=self.select_root_folder).pack(side=tk.LEFT, padx=5)
        self.folder_status = ttk.Label(folder_frame, text="Папка не выбрана", foreground='gray')
        self.folder_status.pack(side=tk.LEFT, padx=5)
        
        # Блок предпросмотра структуры
        tree_frame = ttk.LabelFrame(main_frame, text="Структура отчета", padding=5)
        tree_frame.grid(row=2, column=0, sticky=(tk.NSEW), padx=10, pady=10)
        
        self.treeview = ttk.Treeview(tree_frame, show='tree')
        self.treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.treeview.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.treeview.config(yscrollcommand=scrollbar.set)

    def select_logo(self):
        file_path = filedialog.askopenfilename(filetypes=[("Изображения", "*.jpg *.jpeg *.png")])
        if file_path:
            self.logo_path.set(file_path)
            self.logo_entry.config(textvariable=tk.StringVar(value=os.path.basename(file_path)))

    def select_root_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.root_folder.set(folder)
            self.folder_status.config(text=f"Выбрано: {os.path.basename(folder)}", foreground='black')
            self.scan_folder_structure(folder)

    def scan_folder_structure(self, path):
        self.treeview.delete(*self.treeview.get_children())
        
        def recursive_scan(current_path, parent_id=""):
            items = sorted(os.listdir(current_path), key=lambda x: self.file_priority(x))
            for item in items:
                full_path = os.path.join(current_path, item)
                if os.path.isdir(full_path):
                    node = self.treeview.insert(parent_id, 'end', text=item, open=True)
                    recursive_scan(full_path, node)
                else:
                    if item.lower().endswith(('.png', '.jpg', '.jpeg')):
                        self.treeview.insert(parent_id, 'end', text=item, tags=('image'))
        
        recursive_scan(path)
        self.treeview.tag_configure('image', foreground='blue')

    def file_priority(self, file_name):
        """Приоритет файлам с 'обзор'"""
        return 0 if 'обзор' in file_name.lower() else 1

    def create_report(self):
        if not self.root_folder.get():
            messagebox.showwarning("Ошибка", "Выберите корневую папку")
            return

        try:
            doc = Document()
            self.setup_document_styles(doc)
            self.add_title_page(doc)
            self.add_table_of_contents(doc)
            self.add_image_sections(doc, self.root_folder.get(), [])
            self.rebuild_toc(doc)
            self.configure_margins(doc)
            self.last_saved_file = self.save_report(doc)
            self.view_btn.config(state=tk.NORMAL if self.last_saved_file else tk.DISABLED)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")

    def setup_document_styles(self, doc):
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
        
        heading_styles = [
            ('Heading 1', Pt(14), True),
            ('Heading 2', Pt(12), True),
            ('Heading 3', Pt(11), False)
        ]
        for style_name, size, bold in heading_styles:
            if style_name not in doc.styles:
                doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style = doc.styles[style_name]
            style.font.name = 'Times New Roman'
            style.font.size = size
            style.font.bold = bold
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if 'Caption' not in doc.styles:
            doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style = doc.styles['Caption']
        caption_style.font.name = 'Times New Roman'
        caption_style.font.size = Pt(10)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(0, 0, 0)

    def add_title_page(self, doc):
        section = doc.sections[0]
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        
        if self.logo_path.get():
            header = section.header
            header.is_linked_to_previous = False
            logo_paragraph = header.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                logo_paragraph.add_run().add_picture(
                    self.logo_path.get(),
                    width=Cm(10)
                )
            except:
                pass

        title_text = (
            f"\n\n{self.company_name.get()}\n\n"
            f"{self.project_name.get().upper()}\n\n"
            f"{self.report_title.get().upper()}\n\n"
            f"За период с {self.start_date.get()} по {self.end_date.get()}\n\n"
            f"{self.city.get()} {self.year.get()}"
        )
        
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in title_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            
        doc.add_section(WD_SECTION.NEW_PAGE)
        new_section = doc.sections[1]
        new_section.header.is_linked_to_previous = False

    def add_table_of_contents(self, doc):
        doc.add_heading("Оглавление", 1)
        toc = doc.add_paragraph()
        fld_char = OxmlElement('w:fldSimple')
        fld_char.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
        toc._element.append(fld_char)
        doc.add_page_break()

    def rebuild_toc(self, doc):
        """Пересоздание оглавления"""
        for paragraph in doc.paragraphs:
            if paragraph.text == "Оглавление":
                toc_parent = paragraph._element.getparent()
                for child in toc_parent.getchildren():
                    if child.tag.endswith('fldSimple'):
                        toc_parent.remove(child)
                break
        self.add_table_of_contents(doc)

    def add_image_sections(self, doc, path, parent_numbers):
        items = sorted(os.listdir(path), key=lambda x: self.file_priority(x))
        total_images = self.count_images(path)
        processed = 0
        
        for idx, item in enumerate(items, 1):
            full_path = os.path.join(path, item)
            current_numbers = parent_numbers + [str(idx)]
            heading_level = len(current_numbers)
            heading_text = f"{'.'.join(current_numbers)} {item}"
            
            if os.path.isdir(full_path):
                doc.add_heading(heading_text, level=heading_level)
                self.add_image_sections(doc, full_path, current_numbers)
            else:
                if item.lower().endswith(('.png', '.jpg', '.jpeg')):
                    self.add_image_pair(doc, full_path)
                    processed += 1
                    self.progress['value'] = (processed / total_images) * 100
                    self.master.update_idletasks()
        if total_images == 0:
            self.progress['value'] = 100

    def count_images(self, path):
        count = 0
        for root, _, files in os.walk(path):
            count += len([f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
        return count

    def add_image_pair(self, doc, image_path):
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(15)
        
        image_cell = table.cell(0, 0)
        image_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            image_cell.paragraphs[0].add_run().add_picture(
                image_path,
                width=Cm(14)
            )
        except:
            image_cell.text = "Ошибка загрузки"

        caption = os.path.splitext(os.path.basename(image_path))[0]
        caption_cell = table.cell(1, 0)
        caption_cell.text = caption
        caption_cell.paragraphs[0].style = doc.styles['Caption']
        caption_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    def configure_margins(self, doc):
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
            
            footer = section.footer
            footer.is_linked_to_previous = False
            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fld_char = OxmlElement('w:fldSimple')
            fld_char.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
            paragraph._element.append(fld_char)
            paragraph.style.font.name = 'Times New Roman'
            paragraph.style.font.size = Pt(10)

    def save_report(self, doc):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документы Word", "*.docx")],
            initialfile=f"{self.report_title.get()}_{self.year.get()}"
        )
        if file_path:
            try:
                doc.save(file_path)
                messagebox.showinfo("Успех", f"Отчет сохранен: {file_path}")
                return file_path
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить: {str(e)}")
        return None

    def view_report(self):
        if self.last_saved_file and os.path.exists(self.last_saved_file):
            try:
                os.startfile(self.last_saved_file)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть отчет: {e}")
        else:
            messagebox.showwarning("Внимание", "Файл отчета не найден")

    def show_about(self):
        messagebox.showinfo(
            "О программе",
            "Генератор фотоотчетов ГОСТ v22.0\n\n"
            "Особенности:\n"
            "- Автоматическое оглавление\n"
            "- Файлы с 'обзор' в начале\n"
            "- Иерархическая нумерация\n"
            "- Номера страниц\n"
            "- Черный шрифт Times New Roman\n"
            "- Предпросмотр структуры\n"
            "- Статус-бар прогресса"
        )

    def file_priority(self, file_name):
        """Приоритет файлам с 'обзор'"""
        return 0 if 'обзор' in file_name.lower() else 1

    def add_image_sections(self, doc, path, parent_numbers):
        items = sorted(os.listdir(path), key=lambda x: self.file_priority(x))
        total_images = self.count_images(path)
        processed = 0
        
        for idx, item in enumerate(items, 1):
            full_path = os.path.join(path, item)
            current_numbers = parent_numbers + [str(idx)]
            heading_level = len(current_numbers)
            heading_text = f"{'.'.join(current_numbers)} {item}"
            
            if os.path.isdir(full_path):
                doc.add_heading(heading_text, level=heading_level)
                self.add_image_sections(doc, full_path, current_numbers)
            else:
                if item.lower().endswith(('.png', '.jpg', '.jpeg')):
                    self.add_image_pair(doc, full_path)
                    processed += 1
                    self.progress['value'] = (processed / total_images) * 100
                    self.master.update_idletasks()
        if total_images == 0:
            self.progress['value'] = 100

    def create_report(self):
        if not self.root_folder.get():
            messagebox.showwarning("Ошибка", "Выберите корневую папку")
            return

        try:
            doc = Document()
            self.setup_document_styles(doc)
            self.add_title_page(doc)
            self.add_table_of_contents(doc)
            self.add_image_sections(doc, self.root_folder.get(), [])
            self.rebuild_toc(doc)
            self.configure_margins(doc)
            self.last_saved_file = self.save_report(doc)
            self.view_btn.config(state=tk.NORMAL if self.last_saved_file else tk.DISABLED)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = PhotoReportApp(root)
    root.mainloop()