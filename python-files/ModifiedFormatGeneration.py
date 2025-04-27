import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.ttk import Style, Frame, Label, Entry, Button, Combobox, Checkbutton, Progressbar, Notebook
import pandas as pd
from docx import Document
import os

# Email via Outlook
try:
    import win32com.client as win32
    OUTLOOK_SUPPORT = True
except ImportError:
    OUTLOOK_SUPPORT = False

# Optional: convert to PDF
try:
    from docx2pdf import convert
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# Force use of openpyxl for Excel
pd_read_excel = lambda *args, **kwargs: pd.read_excel(*args, engine='openpyxl', **kwargs)

# Helper functions for placeholder replacement
def replace_text_in_paragraph(paragraph, data_dict):
    for key, val in data_dict.items():
        placeholder = f'{{{key}}}'
        if placeholder in paragraph.text:
            full_text = ''.join(run.text for run in paragraph.runs)
            if placeholder in full_text:
                new_text = full_text.replace(placeholder, val)
                for run in paragraph.runs:
                    run.text = ''
                paragraph.runs[0].text = new_text

def replace_placeholders(doc, data_dict):
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, data_dict)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, data_dict)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for p in hf.paragraphs:
                replace_text_in_paragraph(p, data_dict)

class LetterGeneratorApp:
    def __init__(self, root):
        self.root = root
        root.title('Format-Based Content Generator')
        root.geometry('700x720')
        root.resizable(False, False)
        root.configure(bg='#F3E5F5')

        style = Style()
        style.theme_use('vista')
        style.configure('TLabel', font=('Segoe UI', 11), background='#F3E5F5', foreground='#4A148C')
        style.configure('TButton', font=('Segoe UI', 11, 'bold'), padding=(20,10), background='#CE93D8', foreground='sky blue', borderwidth=0)
        style.map('TButton', background=[('active','#AB47BC')], foreground=[('disabled','#CCC')])
        style.configure('TEntry', font=('Segoe UI', 11), fieldbackground='white')
        style.configure('TCombobox', font=('Segoe UI', 11))
        style.configure('TCheckbutton', font=('Segoe UI', 10), background='#F3E5F5', foreground='#4A148C')

        self.excel_path = tk.StringVar()
        self.template_path = tk.StringVar()
        self.output_type = tk.StringVar(value='Both')
        self.same_folder = tk.BooleanVar()
        self.send_email = tk.BooleanVar()

        container = Frame(root, padding=20, style='TFrame')
        container.pack(fill='both', expand=True)

        Label(container, text='📄 Format-Based Generator', font=('Segoe UI Semibold', 16)).pack(pady=(0,20))

        self._add_file_picker(container, 'Excel Data File:', self.excel_path, self.load_excel)
        self._add_file_picker(container, 'Word Template File:', self.template_path, self.load_template)

        tabs = Notebook(container)
        general_tab = Frame(tabs, style='TFrame')
        email_tab = Frame(tabs, style='TFrame')
        tabs.add(general_tab, text='General')
        tabs.add(email_tab, text='Email')
        tabs.pack(fill='x', pady=15)

        Label(general_tab, text='Output Type:').grid(row=0, column=0, pady=5, sticky='w')
        Combobox(general_tab, textvariable=self.output_type, values=['Word','PDF','Both'], state='readonly', width=10).grid(row=0, column=1, pady=5, sticky='w')
        Checkbutton(general_tab, text='Save in same folder', variable=self.same_folder).grid(row=1, column=0, columnspan=2, sticky='w')

        Checkbutton(email_tab, text='Enable Email', variable=self.send_email, command=self.toggle_email_fields).grid(row=0, column=0, pady=5, sticky='w')
        Label(email_tab, text='Subject:').grid(row=1, column=0, sticky='w')
        self.subject_entry = Entry(email_tab, width=50, state='disabled')
        self.subject_entry.grid(row=1, column=1, padx=5, pady=5)
        Label(email_tab, text='Message:').grid(row=2, column=0, sticky='nw')
        self.email_text = tk.Text(email_tab, width=50, height=6, state='disabled', borderwidth=1, relief='solid')
        self.email_text.grid(row=2, column=1, padx=5, pady=5)

        Label(container, text='Progress:').pack(anchor='w', pady=(10,0))
        self.progress = Progressbar(container, length=650, mode='determinate')
        self.progress.pack(pady=5)
        self.progress_label = Label(container, text='0%')
        self.progress_label.pack()

        Button(container, text='Generate', command=self.generate).pack(pady=20, ipadx=10)

        Label(container, text='Developed by Bijay Pd. Shrestha', font=('Segoe UI', 9)).pack(anchor='e')

    def _add_file_picker(self, parent, text, var, command):
        frame = Frame(parent, style='TFrame')
        frame.pack(fill='x', pady=5)
        Label(frame, text=text).pack(anchor='w')
        inner = Frame(frame, style='TFrame')
        inner.pack(fill='x')
        Entry(inner, textvariable=var, width=50).pack(side='left', fill='x', expand=True)
        Button(inner, text='Browse', command=command).pack(side='right', padx=5)

    def load_excel(self):
        path = filedialog.askopenfilename(filetypes=[('Excel files', '*.xlsx *.xls')])
        if path:
            self.excel_path.set(path)

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[('Word files', '*.docx')])
        if path:
            self.template_path.set(path)

    def toggle_email_fields(self):
        state = 'normal' if self.send_email.get() else 'disabled'
        self.subject_entry.configure(state=state)
        self.email_text.configure(state=state)

    def generate(self):
        if not self.excel_path.get() or not self.template_path.get():
            messagebox.showwarning('Missing Input', 'Please select both Excel and Word template files.')
            return
        if self.send_email.get() and not OUTLOOK_SUPPORT:
            messagebox.showerror('Outlook Error', 'win32com.client not installed. Email disabled.')
            return

        try:
            df = pd_read_excel(self.excel_path.get(), dtype=str)
        except Exception as e:
            messagebox.showerror('Error', 'Failed to read Excel: ' + str(e))
            return

        if self.send_email.get() and 'Email' not in df.columns:
            messagebox.showerror('Missing Column', 'Excel must have an "Email" column to send emails.')
            return

        out_dir = os.path.dirname(self.excel_path.get()) if self.same_folder.get() else filedialog.askdirectory()
        if not out_dir:
            return

        total = len(df)
        self.progress['maximum'] = total
        self.progress['value'] = 0

        for i, row in df.iterrows():
            data_dict = {col: str(row[col]) for col in df.columns if pd.notna(row[col])}
            doc = Document(self.template_path.get())
            replace_placeholders(doc, data_dict)

            base = str(row[df.columns[0]])
            word_path = os.path.join(out_dir, base + '.docx')
            pdf_path = word_path.replace('.docx', '.pdf')
            out_type = self.output_type.get()

            if out_type in ('Word', 'Both'):
                doc.save(word_path)
            if out_type in ('PDF', 'Both'):
                temp_file = word_path if out_type == 'Both' else os.path.join(out_dir, '__temp__.docx')
                if out_type == 'PDF':
                    doc.save(temp_file)
                if PDF_SUPPORT:
                    convert(temp_file, pdf_path)
                    if out_type == 'PDF' and os.path.exists(temp_file):
                        os.remove(temp_file)
                else:
                    messagebox.showwarning('PDF Not Supported', 'Install docx2pdf to enable PDF conversion.')

            if self.send_email.get():
                try:
                    outlook = win32.Dispatch('outlook.application')
                    mail = outlook.CreateItem(0)
                    mail.To = row['Email']
                    mail.Subject = self.subject_entry.get().strip()
                    mail.Body = self.email_text.get('1.0', 'end').strip()
                    if out_type in ('Word', 'Both') and os.path.exists(word_path):
                        mail.Attachments.Add(word_path)
                    if out_type in ('PDF', 'Both') and os.path.exists(pdf_path):
                        mail.Attachments.Add(pdf_path)
                    mail.Send()
                except Exception as e:
                    messagebox.showerror('Email Error', f"Failed to send email for {base}: {e}")

            self.progress['value'] = i + 1
            percent = int((i + 1) / total * 100)
            self.progress_label.config(text=f"Progress: {percent}% ({i+1}/{total})")
            self.root.update_idletasks()

        messagebox.showinfo('Done', 'Process completed successfully.')

if __name__ == '__main__':
    root = tk.Tk()
    app = LetterGeneratorApp(root)
    root.mainloop()
