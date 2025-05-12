import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import docx
from docx.shared import Pt
import os
import json
import datetime
import calendar
import subprocess
import sys
from decimal import Decimal
import csv
import xml.etree.ElementTree as ET
import xml.dom.minidom as minidom
import time
import threading
import schedule
import webbrowser
import re
from pathlib import Path
import shutil

class RechnungsProgramm:
    def __init__(self, root):
        self.root = root
        self.root.title("Rechnungserstellung Pro")
        self.root.geometry("900x700")
        
        # Daten initialisieren
        self.firmendaten = {}
        self.adressbuch = {}
        self.artikelbuch = {}  # NEU: Gespeicherte Artikel
        self.artikelgruppen = {}  # NEU: Artikelgruppen für Vorlagen
        self.turnusrechnungen = {}  # NEU: Turnusrechnungen
        self.mahnungen = {}  # NEU: Gespeicherte Mahnvorlagen
        
        self.rechnungsdaten = {
            "empfaenger": "",
            "artikel": [],
            "datum": datetime.datetime.now().strftime("%d.%m.%Y")
        }
        
        # Rechnungsdatenbank initialisieren
        self.rechnungsdatenbank = {}  # Speichert alle erstellten Rechnungen
        
        # Daten laden, wenn vorhanden
        self.lade_daten()
        
        # Notebook für Tabs erstellen
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Tabs erstellen
        self.tab_firmendaten = ttk.Frame(self.notebook)
        self.tab_rechnung = ttk.Frame(self.notebook)
        self.tab_adressbuch = ttk.Frame(self.notebook)
        self.tab_angebote = ttk.Frame(self.notebook)  # NEU: Angebote-Tab
        self.tab_turnusrechnungen = ttk.Frame(self.notebook)  # NEU: Turnusrechnungen-Tab
        self.tab_mahnungen = ttk.Frame(self.notebook)  # NEU: Mahnungen-Tab
        self.tab_artikel = ttk.Frame(self.notebook)  # NEU: Artikelverwaltung 
        self.tab_rechnungsbearbeitung = ttk.Frame(self.notebook)  # NEU: Rechnungsbearbeitung
        
        self.notebook.add(self.tab_firmendaten, text="Firmendaten")
        self.notebook.add(self.tab_rechnung, text="Rechnung erstellen")
        self.notebook.add(self.tab_angebote, text="Angebote")  # NEU
        self.notebook.add(self.tab_adressbuch, text="Adressbuch")
        self.notebook.add(self.tab_turnusrechnungen, text="Turnusrechnungen")  # NEU
        self.notebook.add(self.tab_mahnungen, text="Mahnungen")  # NEU
        self.notebook.add(self.tab_artikel, text="Artikelverwaltung")  # NEU
        self.notebook.add(self.tab_rechnungsbearbeitung, text="Rechnungsbearbeitung")  # NEU
        
        # Maximale Artikelanzahl (kann nun dynamisch erweitert werden)
        self.max_artikel_default = 15
        
        # Tabs initialisieren
        self.init_firmendaten_tab()
        self.init_rechnung_tab()
        self.init_adressbuch_tab()
        self.init_angebote_tab()  # NEU
        self.init_turnusrechnungen_tab()  # NEU
        self.init_mahnungen_tab()  # NEU
        self.init_artikel_tab()  # NEU
        self.init_rechnungsbearbeitung_tab()  # NEU
        
        # Scheduler für Turnusrechnungen starten
        self.scheduler_thread = None
        self.scheduler_running = False
        self.starte_scheduler()
    
    def lade_daten(self):
        # Firmendaten laden
        if os.path.exists("firmendaten.json"):
            with open("firmendaten.json", "r", encoding="utf-8") as f:
                self.firmendaten = json.load(f)
        
        # Adressbuch laden
        if os.path.exists("adressbuch.json"):
            with open("adressbuch.json", "r", encoding="utf-8") as f:
                self.adressbuch = json.load(f)
        
        # Artikelbuch laden
        if os.path.exists("artikelbuch.json"):
            with open("artikelbuch.json", "r", encoding="utf-8") as f:
                self.artikelbuch = json.load(f)
        
        # Artikelgruppen laden
        if os.path.exists("artikelgruppen.json"):
            with open("artikelgruppen.json", "r", encoding="utf-8") as f:
                self.artikelgruppen = json.load(f)
        
        # Turnusrechnungen laden
        if os.path.exists("turnusrechnungen.json"):
            with open("turnusrechnungen.json", "r", encoding="utf-8") as f:
                self.turnusrechnungen = json.load(f)
        
        # Mahnungen laden
        if os.path.exists("mahnungen.json"):
            with open("mahnungen.json", "r", encoding="utf-8") as f:
                self.mahnungen = json.load(f)
        
        # Rechnungsdatenbank laden
        if os.path.exists("rechnungsdatenbank.json"):
            with open("rechnungsdatenbank.json", "r", encoding="utf-8") as f:
                self.rechnungsdatenbank = json.load(f)
    
    def speichere_daten(self):
        # Firmendaten speichern
        with open("firmendaten.json", "w", encoding="utf-8") as f:
            json.dump(self.firmendaten, f, ensure_ascii=False, indent=4)
        
        # Adressbuch speichern
        with open("adressbuch.json", "w", encoding="utf-8") as f:
            json.dump(self.adressbuch, f, ensure_ascii=False, indent=4)
        
        # Artikelbuch speichern
        with open("artikelbuch.json", "w", encoding="utf-8") as f:
            json.dump(self.artikelbuch, f, ensure_ascii=False, indent=4)
        
        # Artikelgruppen speichern
        with open("artikelgruppen.json", "w", encoding="utf-8") as f:
            json.dump(self.artikelgruppen, f, ensure_ascii=False, indent=4)
        
        # Turnusrechnungen speichern
        with open("turnusrechnungen.json", "w", encoding="utf-8") as f:
            json.dump(self.turnusrechnungen, f, ensure_ascii=False, indent=4)
        
        # Mahnungen speichern
        with open("mahnungen.json", "w", encoding="utf-8") as f:
            json.dump(self.mahnungen, f, ensure_ascii=False, indent=4)
        
        # Rechnungsdatenbank speichern
        with open("rechnungsdatenbank.json", "w", encoding="utf-8") as f:
            json.dump(self.rechnungsdatenbank, f, ensure_ascii=False, indent=4)
    
    def init_firmendaten_tab(self):
        frame = ttk.Frame(self.tab_firmendaten, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Felder für Firmendaten
        felder = [
            "Firmenname", "Straße", "PLZ", "Ort", "Telefon", "Email",
            "Steuernummer", "USt-ID", "IBAN", "BIC", "Bank", 
            "Website", "Geschäftsführer", "Handelsregister"  # Zusätzliche Felder
        ]
        
        self.firmendaten_eintraege = {}
        
        for i, feld in enumerate(felder):
            ttk.Label(frame, text=f"{feld}:").grid(row=i, column=0, sticky="w", pady=5)
            var = tk.StringVar(value=self.firmendaten.get(feld, ""))
            entry = ttk.Entry(frame, width=50, textvariable=var)
            entry.grid(row=i, column=1, sticky="ew", pady=5, padx=5)
            self.firmendaten_eintraege[feld] = var
        
        # Buttons
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=len(felder), column=0, columnspan=2, pady=20)
        
        ttk.Button(button_frame, text="Speichern", command=self.speichere_firmendaten).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Briefkopf erstellen", command=self.erstelle_briefkopf).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Briefkopf laden", command=self.lade_briefkopf).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Logo hinzufügen", command=self.logo_hinzufuegen).pack(side="left", padx=5)
    
    def logo_hinzufuegen(self):
        """Logo für Briefkopf auswählen"""
        filepath = filedialog.askopenfilename(
            filetypes=[("Bilddateien", "*.png;*.jpg;*.jpeg;*.bmp")],
            title="Logo auswählen"
        )
        
        if filepath:
            self.firmendaten["logo_pfad"] = filepath
            self.speichere_daten()
            messagebox.showinfo("Erfolg", "Logo wurde ausgewählt und wird in Dokumenten verwendet.")
    
    def init_rechnung_tab(self):
        # Hauptframe mit Scrollbar
        main_frame = ttk.Frame(self.tab_rechnung)
        main_frame.pack(fill='both', expand=True)
        
        # Canvas für Scrolling
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Content Frame im scrollbaren Bereich
        frame = ttk.Frame(scrollable_frame, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Empfänger-Auswahl
        ttk.Label(frame, text="Rechnungsempfänger:").grid(row=0, column=0, sticky="w", pady=5)
        self.empfaenger_var = tk.StringVar()
        self.empfaenger_combobox = ttk.Combobox(frame, textvariable=self.empfaenger_var, width=40)
        self.empfaenger_combobox.grid(row=0, column=1, sticky="ew", pady=5)
        self.aktualisiere_empfaenger_liste()
        
        # Typ der Rechnung (Lieferung oder Leistung)
        ttk.Label(frame, text="Typ:").grid(row=1, column=0, sticky="w", pady=5)
        self.rechnungstyp_var = tk.StringVar(value="Artikel")
        rechnungstyp_combo = ttk.Combobox(frame, textvariable=self.rechnungstyp_var, width=15, 
                                         values=["Artikel", "Leistung"])
        rechnungstyp_combo.grid(row=1, column=1, sticky="w", pady=5)
        rechnungstyp_combo.bind("<<ComboboxSelected>>", self.update_artikeltabelle_header)
        
        # Vorlage laden
        vorlagen_frame = ttk.Frame(frame)
        vorlagen_frame.grid(row=1, column=2, padx=10, sticky="w")
        
        ttk.Label(vorlagen_frame, text="Artikelgruppe laden:").pack(side="left")
        self.artikelgruppe_var = tk.StringVar()
        self.artikelgruppe_combo = ttk.Combobox(vorlagen_frame, textvariable=self.artikelgruppe_var, width=20)
        self.artikelgruppe_combo.pack(side="left", padx=5)
        self.aktualisiere_artikelgruppen_liste()
        
        ttk.Button(vorlagen_frame, text="Laden", command=self.lade_artikelgruppe).pack(side="left", padx=5)
        ttk.Button(vorlagen_frame, text="Speichern", command=self.speichere_als_artikelgruppe).pack(side="left", padx=5)
        
        # Artikeltabelle
        ttk.Label(frame, text="Artikel:").grid(row=2, column=0, sticky="nw", pady=5)
        
        # Artikelframe mit eigenem Scrollbereich
        self.artikel_canvas_frame = ttk.Frame(frame)
        self.artikel_canvas_frame.grid(row=2, column=1, columnspan=2, sticky="ew", pady=5)
        
        self.artikel_canvas = tk.Canvas(self.artikel_canvas_frame, height=300)
        artikel_scrollbar = ttk.Scrollbar(self.artikel_canvas_frame, orient="vertical", command=self.artikel_canvas.yview)
        
        self.artikel_scrollable_frame = ttk.Frame(self.artikel_canvas)
        
        self.artikel_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.artikel_canvas.configure(scrollregion=self.artikel_canvas.bbox("all"))
        )
        
        self.artikel_canvas.create_window((0, 0), window=self.artikel_scrollable_frame, anchor="nw")
        self.artikel_canvas.configure(yscrollcommand=artikel_scrollbar.set)
        
        self.artikel_canvas.pack(side="left", fill="both", expand=True)
        artikel_scrollbar.pack(side="right", fill="y")
        
        # Tabellenkopf im scrollbaren Bereich
        self.tabellenkopf_frame = ttk.Frame(self.artikel_scrollable_frame)
        self.tabellenkopf_frame.pack(fill="x")
        
        # Die eigentlichen Zeilen im scrollbaren Bereich
        self.artikel_container_frame = ttk.Frame(self.artikel_scrollable_frame)
        self.artikel_container_frame.pack(fill="both", expand=True)
        
        # Initialisiere die Tabelle
        self.artikel_eintraege = []
        self.init_artikeltabelle()
        
        # Buttons für Artikelverwaltung
        artikel_buttons_frame = ttk.Frame(frame)
        artikel_buttons_frame.grid(row=3, column=1, columnspan=2, sticky="w", pady=5)
        
        ttk.Button(artikel_buttons_frame, text="Zeile hinzufügen", 
                  command=self.artikel_zeile_hinzufuegen).pack(side="left", padx=5)
        ttk.Button(artikel_buttons_frame, text="Zeile entfernen", 
                  command=self.artikel_zeile_entfernen).pack(side="left", padx=5)
        ttk.Button(artikel_buttons_frame, text="Artikel aus Katalog", 
                  command=self.artikel_aus_katalog).pack(side="left", padx=5)
        
        # Rechnungsinformationen Frame
        info_frame = ttk.Frame(frame)
        info_frame.grid(row=4, column=0, columnspan=3, sticky="ew", pady=10)
        
        # Rechnungsnummer und Datum
        ttk.Label(info_frame, text="Rechnungsnummer:").grid(row=0, column=0, sticky="w", pady=5)
        self.rechnungsnr_var = tk.StringVar(value=self.generiere_rechnungsnummer())
        ttk.Entry(info_frame, textvariable=self.rechnungsnr_var, width=20).grid(row=0, column=1, sticky="w", pady=5)
        
        ttk.Label(info_frame, text="Rechnungsdatum:").grid(row=1, column=0, sticky="w", pady=5)
        self.rechnungsdatum_var = tk.StringVar(value=datetime.datetime.now().strftime("%d.%m.%Y"))
        ttk.Entry(info_frame, textvariable=self.rechnungsdatum_var, width=20).grid(row=1, column=1, sticky="w", pady=5)
        
        # Leistungszeitraum
        ttk.Label(info_frame, text="Leistungszeitraum:").grid(row=2, column=0, sticky="w", pady=5)
        
        leistung_frame = ttk.Frame(info_frame)
        leistung_frame.grid(row=2, column=1, sticky="w", pady=5)
        
        self.leistung_start_var = tk.StringVar()
        self.leistung_ende_var = tk.StringVar()
        
        ttk.Entry(leistung_frame, textvariable=self.leistung_start_var, width=10).pack(side="left")
        ttk.Label(leistung_frame, text=" - ").pack(side="left")
        ttk.Entry(leistung_frame, textvariable=self.leistung_ende_var, width=10).pack(side="left")
        
        # Monatsauswahl für Leistungszeitraum
        self.monat_var = tk.StringVar()
        today = datetime.datetime.now()
        monate = [f"{i:02d}.{today.year}" for i in range(1, 13)]
        monat_combo = ttk.Combobox(leistung_frame, textvariable=self.monat_var, width=15, values=monate)
        monat_combo.pack(side="left", padx=(10, 0))
        monat_combo.bind("<<ComboboxSelected>>", self.update_leistungszeitraum)
        
        # Zahlungsbedingungen
        ttk.Label(info_frame, text="Zahlungsbedingungen:").grid(row=3, column=0, sticky="w", pady=5)
        zahlungsbedingungen_frame = ttk.Frame(info_frame)
        zahlungsbedingungen_frame.grid(row=3, column=1, sticky="w", pady=5)
        
        ttk.Label(zahlungsbedingungen_frame, text="Zahlungsziel (Tage):").pack(side="left")
        self.zahlungsziel_var = tk.StringVar(value="14")
        ttk.Entry(zahlungsbedingungen_frame, textvariable=self.zahlungsziel_var, width=5).pack(side="left", padx=5)
        
        ttk.Label(zahlungsbedingungen_frame, text="Skonto (%):").pack(side="left", padx=(10,0))
        self.skonto_prozent_var = tk.StringVar(value="0")
        ttk.Entry(zahlungsbedingungen_frame, textvariable=self.skonto_prozent_var, width=5).pack(side="left", padx=5)
        
        ttk.Label(zahlungsbedingungen_frame, text="bei Zahlung innerhalb von (Tagen):").pack(side="left")
        self.skonto_tage_var = tk.StringVar(value="0")
        ttk.Entry(zahlungsbedingungen_frame, textvariable=self.skonto_tage_var, width=5).pack(side="left", padx=5)
        
        # Steuerschuldnerschaft-Option
        self.steuerschuldnerschaft_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(info_frame, text="Steuerschuldnerschaft des Leistungsempfängers", 
                         variable=self.steuerschuldnerschaft_var).grid(row=4, column=0, columnspan=2, sticky="w", pady=5)
        
        # Storno-Option
        storno_frame = ttk.Frame(info_frame)
        storno_frame.grid(row=5, column=0, columnspan=2, sticky="w", pady=5)
        
        self.storno_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(storno_frame, text="Storno/Korrektur erstellen", 
                         variable=self.storno_var, command=self.aktiviere_storno_auswahl).pack(side="left")
        
        ttk.Label(storno_frame, text="Original-Rechnungsnr.:").pack(side="left", padx=(10, 0))
        self.storno_rechnungsnr_var = tk.StringVar()
        self.storno_rechnung_combo = ttk.Combobox(storno_frame, textvariable=self.storno_rechnungsnr_var, width=15)
        self.storno_rechnung_combo.pack(side="left")
        
        # Buttons
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=6, column=0, columnspan=3, pady=10)
        
        ttk.Button(button_frame, text="Neu", command=self.neue_rechnung).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Vorschau", command=self.berechne_summen).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Rechnung erstellen", command=self.erstelle_rechnung).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Rechnung erstellen und per E-Mail versenden", 
                   command=self.erstelle_rechnung_und_versende).pack(side="left", padx=5)
        
        # Vorschau-Bereich
        self.vorschau_text = tk.Text(frame, height=8, width=90)
        self.vorschau_text.grid(row=7, column=0, columnspan=3, pady=10)
        self.vorschau_text.config(state=tk.DISABLED)
    
    def aktiviere_storno_auswahl(self):
        """Aktiviert oder deaktiviert die Storno-Rechnungsauswahl und lädt die verfügbaren Rechnungen"""
        if self.storno_var.get():
            # Nur bestehende Rechnungsnummern anzeigen
            rechnungsnummern = list(self.rechnungsdatenbank.keys())
            self.storno_rechnung_combo['values'] = rechnungsnummern
            self.storno_rechnung_combo['state'] = 'readonly'
        else:
            self.storno_rechnung_combo['state'] = 'disabled'
            self.storno_rechnungsnr_var.set("")
    
    def init_artikeltabelle(self):
        """Initialisiert die Artikeltabelle mit Standardwerten"""
        # Tabellenkopf erstellen
        self.tabellenkopf_frame.destroy()
        self.tabellenkopf_frame = ttk.Frame(self.artikel_scrollable_frame)
        self.tabellenkopf_frame.pack(fill="x", pady=(0, 5))
        
        ist_leistung = self.rechnungstyp_var.get() == "Leistung"
        
        col_widths = [30, 8, 10, 8, 10] if ist_leistung else [30, 10, 10, 8, 10]
        col_names = ["Bezeichnung", "Stunden", "Stundensatz €", "USt %", "Gesamt €"] if ist_leistung else ["Bezeichnung", "Anzahl", "Netto €", "USt %", "Gesamt €"]
        
        for i, (name, width) in enumerate(zip(col_names, col_widths)):
            ttk.Label(self.tabellenkopf_frame, text=name, width=width).grid(row=0, column=i, padx=2)
        
        # Lege die Anzahl der Start-Zeilen fest
        for _ in range(self.max_artikel_default):
            self.artikel_zeile_hinzufuegen()
    
    def update_artikeltabelle_header(self, event=None):
        """Aktualisiert den Tabellenkopf je nach Rechnungstyp"""
        # Alte Tabelle leeren
        for widget in self.artikel_container_frame.winfo_children():
            widget.destroy()
        
        self.artikel_eintraege = []
        
        # Tabellenkopf neu aufbauen
        self.init_artikeltabelle()
    
    def artikel_zeile_hinzufuegen(self):
        """Fügt eine neue Artikelzeile zur Tabelle hinzu"""
        ist_leistung = self.rechnungstyp_var.get() == "Leistung"
        
        row_frame = ttk.Frame(self.artikel_container_frame)
        row_frame.pack(fill="x", pady=2)
        
        bezeichnung_var = tk.StringVar()
        menge_var = tk.StringVar(value="1")
        preis_var = tk.StringVar(value="0.00")
        ust_var = tk.StringVar(value="19")
        
        bezeichnung_entry = ttk.Entry(row_frame, textvariable=bezeichnung_var, width=30)
        menge_entry = ttk.Entry(row_frame, textvariable=menge_var, width=8 if ist_leistung else 10)
        preis_entry = ttk.Entry(row_frame, textvariable=preis_var, width=10)
        ust_combo = ttk.Combobox(row_frame, textvariable=ust_var, width=8, values=["0", "7", "19"])
        
        # Gesamtpreis-Label (wird bei Berechnung aktualisiert)
        gesamt_label = ttk.Label(row_frame, text="0.00", width=10)
        
        bezeichnung_entry.grid(row=0, column=0, padx=2)
        menge_entry.grid(row=0, column=1, padx=2)
        preis_entry.grid(row=0, column=2, padx=2)
        ust_combo.grid(row=0, column=3, padx=2)
        gesamt_label.grid(row=0, column=4, padx=2)
        
        # Artikel-Auswahl-Button
        artikel_btn = ttk.Button(row_frame, text="...", width=3, 
                                command=lambda row=len(self.artikel_eintraege): 
                                            self.waehle_artikel_aus_katalog(row))
        artikel_btn.grid(row=0, column=5, padx=2)
        
        self.artikel_eintraege.append({
            "frame": row_frame,
            "bezeichnung": bezeichnung_var,
            "menge": menge_var,
            "preis": preis_var,
            "ust": ust_var,
            "gesamt_label": gesamt_label
        })
        
        # Aktualisiere die Scrollregion
        self.artikel_scrollable_frame.update_idletasks()
        self.artikel_canvas.configure(scrollregion=self.artikel_canvas.bbox("all"))
    
    def artikel_zeile_entfernen(self):
        """Entfernt die letzte Artikelzeile, wenn mehr als eine vorhanden ist"""
        if len(self.artikel_eintraege) > 1:
            # Letzten Eintrag entfernen
            letzter_eintrag = self.artikel_eintraege.pop()
            letzter_eintrag["frame"].destroy()
            
            # Aktualisiere die Scrollregion
            self.artikel_scrollable_frame.update_idletasks()
            self.artikel_canvas.configure(scrollregion=self.artikel_canvas.bbox("all"))
    
    def artikel_aus_katalog(self):
        """Öffnet ein Fenster zur Auswahl mehrerer Artikel aus dem Katalog"""
        if not self.artikelbuch:
            messagebox.showinfo("Information", "Der Artikelkatalog ist leer. Bitte fügen Sie zuerst Artikel hinzu.")
            return
        
        # Erstelle ein Toplevel-Fenster für die Artikelauswahl
        auswahl_fenster = tk.Toplevel(self.root)
        auswahl_fenster.title("Artikel aus Katalog auswählen")
        auswahl_fenster.geometry("800x500")
        
        # Linker Frame für die Artikel-Liste
        liste_frame = ttk.Frame(auswahl_fenster, padding=10)
        liste_frame.pack(side="left", fill="both", expand=True)
        
        ttk.Label(liste_frame, text="Verfügbare Artikel:").pack(anchor="w")
        
        # Suchfeld
        search_frame = ttk.Frame(liste_frame)
        search_frame.pack(fill="x", pady=5)
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        search_var = tk.StringVar()
        ttk.Entry(search_frame, textvariable=search_var, width=25).pack(side="left", padx=5)
        
        # Listbox mit Scrollbar
        list_container = ttk.Frame(liste_frame)
        list_container.pack(fill="both", expand=True)
        
        scrollbar = ttk.Scrollbar(list_container)
        scrollbar.pack(side="right", fill="y")
        
        artikel_listbox = tk.Listbox(list_container, yscrollcommand=scrollbar.set, height=20, width=50)
        artikel_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=artikel_listbox.yview)
        
        # Rechter Frame für die gewählten Artikel
        selected_frame = ttk.Frame(auswahl_fenster, padding=10)
        selected_frame.pack(side="right", fill="both", expand=True)
        
        ttk.Label(selected_frame, text="Ausgewählte Artikel:").pack(anchor="w")
        
        # Listbox für gewählte Artikel
        selected_list_container = ttk.Frame(selected_frame)
        selected_list_container.pack(fill="both", expand=True)
        
        selected_scrollbar = ttk.Scrollbar(selected_list_container)
        selected_scrollbar.pack(side="right", fill="y")
        
        selected_listbox = tk.Listbox(selected_list_container, yscrollcommand=selected_scrollbar.set, height=20, width=50)
        selected_listbox.pack(side="left", fill="both", expand=True)
        selected_scrollbar.config(command=selected_listbox.yview)
        
        # Buttons zwischen den Listen
        buttons_frame = ttk.Frame(auswahl_fenster)
        buttons_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        ttk.Button(buttons_frame, text=">>", 
                  command=lambda: self.add_to_selected(artikel_listbox, selected_listbox)).pack(pady=5)
        ttk.Button(buttons_frame, text="<<", 
                  command=lambda: self.remove_from_selected(selected_listbox)).pack(pady=5)
        
        # OK und Abbrechen Buttons
        bottom_frame = ttk.Frame(auswahl_fenster)
        bottom_frame.pack(side="bottom", fill="x", pady=10)
        
        ttk.Button(bottom_frame, text="Übernehmen", 
                  command=lambda: self.uebernehme_ausgewaehlte_artikel(selected_listbox, auswahl_fenster)).pack(side="right", padx=10)
        ttk.Button(bottom_frame, text="Abbrechen", 
                  command=auswahl_fenster.destroy).pack(side="right")
        
        # Artikel in Listbox füllen
        self.refresh_artikel_liste(artikel_listbox, search_var.get())
        
        # Suchfunktion
        search_var.trace("w", lambda name, index, mode, 
                         lb=artikel_listbox, sv=search_var: self.refresh_artikel_liste(lb, sv.get()))
        
        # Fenster modal machen
        auswahl_fenster.transient(self.root)
        auswahl_fenster.grab_set()
        self.root.wait_window(auswahl_fenster)
    
    def refresh_artikel_liste(self, listbox, suchbegriff=""):
        """Aktualisiert die Artikelliste basierend auf dem Suchbegriff"""
        listbox.delete(0, tk.END)
        
        for artikel_id, artikel_info in self.artikelbuch.items():
            artikel_name = artikel_info.get("bezeichnung", "")
            if suchbegriff.lower() in artikel_name.lower() or suchbegriff.lower() in artikel_id.lower():
                display_text = f"{artikel_name} ({artikel_info.get('preis', '0.00')}€)"
                listbox.insert(tk.END, display_text)
                # Hinterlege die Artikel-ID als Attribut
                listbox.itemconfig(listbox.size()-1, {'artikel_id': artikel_id})
    
    def add_to_selected(self, source_listbox, target_listbox):
        """Fügt ausgewählte Artikel zur Auswahlliste hinzu"""
        selected_indices = source_listbox.curselection()
        if not selected_indices:
            return
        
        for i in selected_indices:
            item_text = source_listbox.get(i)
            artikel_id = source_listbox.itemcget(i, 'artikel_id')
            if artikel_id:
                target_listbox.insert(tk.END, item_text)
                target_listbox.itemconfig(target_listbox.size()-1, {'artikel_id': artikel_id})
    
    def remove_from_selected(self, listbox):
        """Entfernt ausgewählte Artikel aus der Auswahlliste"""
        selected_indices = listbox.curselection()
        if not selected_indices:
            return
        
        # Von hinten nach vorne löschen, um Indexproblemen vorzubeugen
        for i in sorted(selected_indices, reverse=True):
            listbox.delete(i)
    
    def uebernehme_ausgewaehlte_artikel(self, listbox, fenster):
        """Übernimmt die ausgewählten Artikel in die Rechnungstabelle"""
        if listbox.size() == 0:
            fenster.destroy()
            return
        
        # Prüfe, ob genug Zeilen vorhanden sind, sonst füge hinzu
        while len(self.artikel_eintraege) < listbox.size():
            self.artikel_zeile_hinzufuegen()
        
        # Übernehme die Artikel in die Tabelle
        for i in range(listbox.size()):
            artikel_id = listbox.itemcget(i, 'artikel_id')
            if artikel_id in self.artikelbuch:
                artikel = self.artikelbuch[artikel_id]
                
                # Artikel in die i-te Zeile eintragen
                if i < len(self.artikel_eintraege):
                    self.artikel_eintraege[i]["bezeichnung"].set(artikel["bezeichnung"])
                    self.artikel_eintraege[i]["preis"].set(artikel["preis"])
                    self.artikel_eintraege[i]["ust"].set(artikel["ust"])
                    self.artikel_eintraege[i]["menge"].set("1")  # Standardmenge 1
        
        fenster.destroy()
    
    def waehle_artikel_aus_katalog(self, row_index):
        """Öffnet ein Fenster zur Auswahl eines einzelnen Artikels für eine bestimmte Zeile"""
        if not self.artikelbuch:
            messagebox.showinfo("Information", "Der Artikelkatalog ist leer. Bitte fügen Sie zuerst Artikel hinzu.")
            return
        
        # Erstelle ein Toplevel-Fenster für die Artikelauswahl
        auswahl_fenster = tk.Toplevel(self.root)
        auswahl_fenster.title("Artikel auswählen")
        auswahl_fenster.geometry("500x400")
        
        # Suchfeld
        search_frame = ttk.Frame(auswahl_fenster, padding=10)
        search_frame.pack(fill="x")
        
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        search_var = tk.StringVar()
        ttk.Entry(search_frame, textvariable=search_var, width=25).pack(side="left", padx=5)
        
        # Listbox mit Scrollbar
        list_container = ttk.Frame(auswahl_fenster, padding=10)
        list_container.pack(fill="both", expand=True)
        
        scrollbar = ttk.Scrollbar(list_container)
        scrollbar.pack(side="right", fill="y")
        
        artikel_listbox = tk.Listbox(list_container, yscrollcommand=scrollbar.set, height=15, width=50)
        artikel_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=artikel_listbox.yview)
        
        # OK und Abbrechen Buttons
        button_frame = ttk.Frame(auswahl_fenster, padding=10)
        button_frame.pack(fill="x")
        
        ttk.Button(button_frame, text="Auswählen", 
                  command=lambda: self.artikel_in_zeile_einfuegen(artikel_listbox, row_index, auswahl_fenster)).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Abbrechen", 
                  command=auswahl_fenster.destroy).pack(side="right")
        
        # Artikel in Listbox füllen
        self.refresh_artikel_liste(artikel_listbox)
        
        # Doppelklick-Bindung
        artikel_listbox.bind("<Double-1>", 
                            lambda event: self.artikel_in_zeile_einfuegen(artikel_listbox, row_index, auswahl_fenster))
        
        # Suchfunktion
        search_var.trace("w", lambda name, index, mode, 
                         lb=artikel_listbox, sv=search_var: self.refresh_artikel_liste(lb, sv.get()))
        
        # Fenster modal machen
        auswahl_fenster.transient(self.root)
        auswahl_fenster.grab_set()
        self.root.wait_window(auswahl_fenster)
    
    def artikel_in_zeile_einfuegen(self, listbox, row_index, fenster):
        """Fügt den ausgewählten Artikel in die angegebene Zeile ein"""
        selection = listbox.curselection()
        if not selection:
            return
        
        artikel_id = listbox.itemcget(selection[0], 'artikel_id')
        if artikel_id in self.artikelbuch:
            artikel = self.artikelbuch[artikel_id]
            
            # In die Zeile eintragen
            if 0 <= row_index < len(self.artikel_eintraege):
                self.artikel_eintraege[row_index]["bezeichnung"].set(artikel["bezeichnung"])
                self.artikel_eintraege[row_index]["preis"].set(artikel["preis"])
                self.artikel_eintraege[row_index]["ust"].set(artikel["ust"])
                self.artikel_eintraege[row_index]["menge"].set("1")  # Standardmenge 1
        
        fenster.destroy()
    
    def init_adressbuch_tab(self):
        frame = ttk.Frame(self.tab_adressbuch, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Linke Seite: Liste der Kunden
        liste_frame = ttk.Frame(frame)
        liste_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        ttk.Label(liste_frame, text="Gespeicherte Kunden:").pack(anchor="w")
        
        # Suchfeld
        search_frame = ttk.Frame(liste_frame)
        search_frame.pack(fill="x", pady=5)
        
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self.search_adressbuch)
        ttk.Entry(search_frame, textvariable=self.search_var, width=20).pack(side="left", padx=5)
        
        # Import/Export Buttons
        import_export_frame = ttk.Frame(liste_frame)
        import_export_frame.pack(fill="x", pady=5)
        
        ttk.Button(import_export_frame, text="Importieren", command=self.adressbuch_importieren).pack(side="left", padx=5)
        ttk.Button(import_export_frame, text="Exportieren", command=self.adressbuch_exportieren).pack(side="left", padx=5)
        
        # Kundenliste
        self.adressen_listbox = tk.Listbox(liste_frame, height=15, width=40)
        self.adressen_listbox.pack(side="top", fill="both", expand=True, pady=5)
        self.adressen_listbox.bind('<<ListboxSelect>>', self.adresse_auswaehlen)
        
        button_frame = ttk.Frame(liste_frame)
        button_frame.pack(side="top", fill="x", pady=5)
        
        ttk.Button(button_frame, text="Neu", command=self.neue_adresse).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Löschen", command=self.loesche_adresse).pack(side="left", padx=5)
        
        # Rechte Seite: Adressdetails
        details_frame = ttk.Frame(frame)
        details_frame.pack(side="right", fill="both", expand=True)
        
        felder = [
            "Vorname", "Nachname", "Firma", "Debitorennummer", "Straße", "PLZ", "Ort", "Land",
            "Telefon", "Handy", "Email", "USt-ID", "Kundengruppe", "Notizen"
        ]
        
        self.adress_eintraege = {}
        
        for i, feld in enumerate(felder):
            ttk.Label(details_frame, text=f"{feld}:").grid(row=i, column=0, sticky="w", pady=5)
            var = tk.StringVar()
            entry = ttk.Entry(details_frame, width=40, textvariable=var)
            entry.grid(row=i, column=1, sticky="ew", pady=5, padx=5)
            self.adress_eintraege[feld] = var
        
        ttk.Button(details_frame, text="Speichern", command=self.speichere_adresse).grid(
            row=len(felder), column=0, columnspan=2, pady=10
        )
        
        # Adressliste aktualisieren
        self.aktualisiere_adressliste()
    
    def adressbuch_importieren(self):
        """Importiert Adressen aus einer CSV-Datei"""
        filepath = filedialog.askopenfilename(
            title="Adressbuch importieren",
            filetypes=[("CSV-Dateien", "*.csv"), ("Alle Dateien", "*.*")]
        )
        
        if not filepath:
            return
        
        try:
            with open(filepath, 'r', newline='', encoding='utf-8') as file:
                reader = csv.DictReader(file, delimiter=';')
                
                import_count = 0
                for row in reader:
                    # Name für Dictionary-Key erstellen
                    if row.get('Nachname') and row.get('Vorname'):
                        name = f"{row.get('Nachname')}, {row.get('Vorname')}"
                    elif row.get('Nachname'):
                        name = row.get('Nachname')
                    elif row.get('Vorname'):
                        name = row.get('Vorname')
                    elif row.get('Firma'):
                        name = row.get('Firma')
                    else:
                        continue  # Überspringe Zeilen ohne Namen/Firma
                    
                    # Neue Adresse anlegen
                    self.adressbuch[name] = row
                    import_count += 1
                
                self.speichere_daten()
                self.aktualisiere_adressliste()
                self.aktualisiere_empfaenger_liste()
                
                messagebox.showinfo("Import erfolgreich", f"{import_count} Adressen wurden importiert.")
        except Exception as e:
            messagebox.showerror("Fehler beim Import", f"Ein Fehler ist aufgetreten: {str(e)}")
    
    def adressbuch_exportieren(self):
        """Exportiert Adressen in eine CSV-Datei"""
        if not self.adressbuch:
            messagebox.showinfo("Information", "Das Adressbuch ist leer.")
            return
        
        filepath = filedialog.asksaveasfilename(
            title="Adressbuch exportieren",
            defaultextension=".csv",
            filetypes=[("CSV-Dateien", "*.csv"), ("Alle Dateien", "*.*")]
        )
        
        if not filepath:
            return
        
        try:
            # Alle möglichen Felder ermitteln
            all_fields = set()
            for adresse in self.adressbuch.values():
                all_fields.update(adresse.keys())
            
            # Sortierte Liste aller Felder
            fieldnames = sorted(list(all_fields))
            
            with open(filepath, 'w', newline='', encoding='utf-8') as file:
                writer = csv.DictWriter(file, fieldnames=fieldnames, delimiter=';')
                writer.writeheader()
                
                # Alle Adressen schreiben
                for adresse in self.adressbuch.values():
                    writer.writerow(adresse)
            
            messagebox.showinfo("Export erfolgreich", f"Das Adressbuch wurde nach {filepath} exportiert.")
        except Exception as e:
            messagebox.showerror("Fehler beim Export", f"Ein Fehler ist aufgetreten: {str(e)}")
    
    def init_angebote_tab(self):
        """Initialisiert den Angebote-Tab mit ähnlicher Struktur wie den Rechnungen-Tab"""
        # Hauptframe mit Scrollbar
        main_frame = ttk.Frame(self.tab_angebote)
        main_frame.pack(fill='both', expand=True)
        
        # Canvas für Scrolling
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Content Frame im scrollbaren Bereich
        frame = ttk.Frame(scrollable_frame, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Empfänger-Auswahl
        ttk.Label(frame, text="Angebotsempfänger:").grid(row=0, column=0, sticky="w", pady=5)
        self.angebot_empfaenger_var = tk.StringVar()
        self.angebot_empfaenger_combobox = ttk.Combobox(frame, textvariable=self.angebot_empfaenger_var, width=40)
        self.angebot_empfaenger_combobox.grid(row=0, column=1, sticky="ew", columnspan=2, pady=5)
        
        # Angebotsverwaltung
        ttk.Label(frame, text="Gespeicherte Angebote:").grid(row=1, column=0, sticky="w", pady=5)
        
        angebote_frame = ttk.Frame(frame)
        angebote_frame.grid(row=1, column=1, columnspan=2, sticky="ew", pady=5)
        
        self.angebote_listbox = tk.Listbox(angebote_frame, height=5, width=50)
        self.angebote_listbox.pack(side="left", fill="both", expand=True)
        angebote_scrollbar = ttk.Scrollbar(angebote_frame, command=self.angebote_listbox.yview)
        angebote_scrollbar.pack(side="right", fill="y")
        self.angebote_listbox.config(yscrollcommand=angebote_scrollbar.set)
        self.angebote_listbox.bind('<<ListboxSelect>>', self.angebot_auswaehlen)
        
        angebot_buttons_frame = ttk.Frame(frame)
        angebot_buttons_frame.grid(row=2, column=1, columnspan=2, sticky="w", pady=5)
        
        ttk.Button(angebot_buttons_frame, text="Neues Angebot", 
                  command=self.neues_angebot).pack(side="left", padx=5)
        ttk.Button(angebot_buttons_frame, text="Angebot laden", 
                  command=self.angebot_laden).pack(side="left", padx=5)
        ttk.Button(angebot_buttons_frame, text="Angebot löschen", 
                  command=self.angebot_loeschen).pack(side="left", padx=5)
        ttk.Button(angebot_buttons_frame, text="Als Rechnung übernehmen", 
                  command=self.angebot_zu_rechnung).pack(side="left", padx=5)
        
        # Typ des Angebots (Lieferung oder Leistung)
        ttk.Label(frame, text="Typ:").grid(row=3, column=0, sticky="w", pady=5)
        self.angebot_typ_var = tk.StringVar(value="Artikel")
        angebot_typ_combo = ttk.Combobox(frame, textvariable=self.angebot_typ_var, width=15, 
                                        values=["Artikel", "Leistung"])
        angebot_typ_combo.grid(row=3, column=1, sticky="w", pady=5)
        angebot_typ_combo.bind("<<ComboboxSelected>>", self.update_angebot_artikeltabelle_header)
        
        # Artikeltabelle
        ttk.Label(frame, text="Artikel/Leistungen:").grid(row=4, column=0, sticky="nw", pady=5)
        
        # Artikelframe mit eigenem Scrollbereich
        self.angebot_artikel_canvas_frame = ttk.Frame(frame)
        self.angebot_artikel_canvas_frame.grid(row=4, column=1, columnspan=2, sticky="ew", pady=5)
        
        self.angebot_artikel_canvas = tk.Canvas(self.angebot_artikel_canvas_frame, height=300)
        angebot_artikel_scrollbar = ttk.Scrollbar(self.angebot_artikel_canvas_frame, 
                                                 orient="vertical", 
                                                 command=self.angebot_artikel_canvas.yview)
        
        self.angebot_artikel_scrollable_frame = ttk.Frame(self.angebot_artikel_canvas)
        
        self.angebot_artikel_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.angebot_artikel_canvas.configure(
                scrollregion=self.angebot_artikel_canvas.bbox("all"))
        )
        
        self.angebot_artikel_canvas.create_window((0, 0), 
                                                window=self.angebot_artikel_scrollable_frame, 
                                                anchor="nw")
        self.angebot_artikel_canvas.configure(yscrollcommand=angebot_artikel_scrollbar.set)
        
        self.angebot_artikel_canvas.pack(side="left", fill="both", expand=True)
        angebot_artikel_scrollbar.pack(side="right", fill="y")
        
        # Tabellenkopf im scrollbaren Bereich
        self.angebot_tabellenkopf_frame = ttk.Frame(self.angebot_artikel_scrollable_frame)
        self.angebot_tabellenkopf_frame.pack(fill="x")
        
        # Die eigentlichen Zeilen im scrollbaren Bereich
        self.angebot_artikel_container_frame = ttk.Frame(self.angebot_artikel_scrollable_frame)
        self.angebot_artikel_container_frame.pack(fill="both", expand=True)
        
        # Initialisiere die Tabelle
        self.angebot_artikel_eintraege = []
        self.init_angebot_artikeltabelle()
        
        # Buttons für Artikelverwaltung
        angebot_artikel_buttons_frame = ttk.Frame(frame)
        angebot_artikel_buttons_frame.grid(row=5, column=1, columnspan=2, sticky="w", pady=5)
        
        ttk.Button(angebot_artikel_buttons_frame, text="Zeile hinzufügen", 
                  command=self.angebot_artikel_zeile_hinzufuegen).pack(side="left", padx=5)
        ttk.Button(angebot_artikel_buttons_frame, text="Zeile entfernen", 
                  command=self.angebot_artikel_zeile_entfernen).pack(side="left", padx=5)
        ttk.Button(angebot_artikel_buttons_frame, text="Artikel aus Katalog", 
                  command=self.angebot_artikel_aus_katalog).pack(side="left", padx=5)
        
        # Angebotsinformationen Frame
        info_frame = ttk.Frame(frame)
        info_frame.grid(row=6, column=0, columnspan=3, sticky="ew", pady=10)
        
        # Angebotsnummer und Datum
        ttk.Label(info_frame, text="Angebotsnummer:").grid(row=0, column=0, sticky="w", pady=5)
        self.angebotsnr_var = tk.StringVar(value=self.generiere_angebotsnummer())
        ttk.Entry(info_frame, textvariable=self.angebotsnr_var, width=20).grid(row=0, column=1, sticky="w", pady=5)
        
        ttk.Label(info_frame, text="Angebotsdatum:").grid(row=1, column=0, sticky="w", pady=5)
        self.angebotsdatum_var = tk.StringVar(value=datetime.datetime.now().strftime("%d.%m.%Y"))
        ttk.Entry(info_frame, textvariable=self.angebotsdatum_var, width=20).grid(row=1, column=1, sticky="w", pady=5)
        
        # Gültigkeit
        ttk.Label(info_frame, text="Gültig bis:").grid(row=2, column=0, sticky="w", pady=5)
        
        # Standardmäßig 4 Wochen gültig
        gueltig_bis = datetime.datetime.now() + datetime.timedelta(days=28)
        self.angebot_gueltig_bis_var = tk.StringVar(value=gueltig_bis.strftime("%d.%m.%Y"))
        ttk.Entry(info_frame, textvariable=self.angebot_gueltig_bis_var, width=20).grid(row=2, column=1, sticky="w", pady=5)
        
        # Buttons
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=7, column=0, columnspan=3, pady=10)
        
        ttk.Button(button_frame, text="Vorschau", 
                  command=self.angebot_berechne_summen).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Angebot speichern", 
                  command=self.angebot_speichern).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Angebot erstellen", 
                  command=self.erstelle_angebot).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Angebot erstellen und per E-Mail versenden", 
                  command=self.erstelle_angebot_und_versende).pack(side="left", padx=5)
        
        # Vorschau-Bereich
        self.angebot_vorschau_text = tk.Text(frame, height=8, width=90)
        self.angebot_vorschau_text.grid(row=8, column=0, columnspan=3, pady=10)
        self.angebot_vorschau_text.config(state=tk.DISABLED)
        
        # Adressliste für Combobox aktualisieren
        self.aktualisiere_empfaenger_liste()
    
    def init_angebot_artikeltabelle(self):
        """Initialisiert die Artikeltabelle für Angebote mit Standardwerten"""
        # Tabellenkopf erstellen
        self.angebot_tabellenkopf_frame.destroy()
        self.angebot_tabellenkopf_frame = ttk.Frame(self.angebot_artikel_scrollable_frame)
        self.angebot_tabellenkopf_frame.pack(fill="x", pady=(0, 5))
        
        ist_leistung = self.angebot_typ_var.get() == "Leistung"
        
        col_widths = [30, 8, 10, 8, 10] if ist_leistung else [30, 10, 10, 8, 10]
        col_names = ["Bezeichnung", "Stunden", "Stundensatz €", "USt %", "Gesamt €"] if ist_leistung else ["Bezeichnung", "Anzahl", "Netto €", "USt %", "Gesamt €"]
        
        for i, (name, width) in enumerate(zip(col_names, col_widths)):
            ttk.Label(self.angebot_tabellenkopf_frame, text=name, width=width).grid(row=0, column=i, padx=2)
        
        # Lege die Anzahl der Start-Zeilen fest
        for _ in range(self.max_artikel_default):
            self.angebot_artikel_zeile_hinzufuegen()
    
    def update_angebot_artikeltabelle_header(self, event=None):
        """Aktualisiert den Tabellenkopf für Angebote je nach Typ"""
        # Alte Tabelle leeren
        for widget in self.angebot_artikel_container_frame.winfo_children():
            widget.destroy()
        
        self.angebot_artikel_eintraege = []
        
        # Tabellenkopf neu aufbauen
        self.init_angebot_artikeltabelle()
    
    def angebot_artikel_zeile_hinzufuegen(self):
        """Fügt eine neue Artikelzeile zur Angebots-Tabelle hinzu"""
        ist_leistung = self.angebot_typ_var.get() == "Leistung"
        
        row_frame = ttk.Frame(self.angebot_artikel_container_frame)
        row_frame.pack(fill="x", pady=2)
        
        bezeichnung_var = tk.StringVar()
        menge_var = tk.StringVar(value="1")
        preis_var = tk.StringVar(value="0.00")
        ust_var = tk.StringVar(value="19")
        
        bezeichnung_entry = ttk.Entry(row_frame, textvariable=bezeichnung_var, width=30)
        menge_entry = ttk.Entry(row_frame, textvariable=menge_var, width=8 if ist_leistung else 10)
        preis_entry = ttk.Entry(row_frame, textvariable=preis_var, width=10)
        ust_combo = ttk.Combobox(row_frame, textvariable=ust_var, width=8, values=["0", "7", "19"])
        
        # Gesamtpreis-Label (wird bei Berechnung aktualisiert)
        gesamt_label = ttk.Label(row_frame, text="0.00", width=10)
        
        bezeichnung_entry.grid(row=0, column=0, padx=2)
        menge_entry.grid(row=0, column=1, padx=2)
        preis_entry.grid(row=0, column=2, padx=2)
        ust_combo.grid(row=0, column=3, padx=2)
        gesamt_label.grid(row=0, column=4, padx=2)
        
        # Artikel-Auswahl-Button
        artikel_btn = ttk.Button(row_frame, text="...", width=3, 
                                command=lambda row=len(self.angebot_artikel_eintraege): 
                                            self.angebot_waehle_artikel_aus_katalog(row))
        artikel_btn.grid(row=0, column=5, padx=2)
        
        self.angebot_artikel_eintraege.append({
            "frame": row_frame,
            "bezeichnung": bezeichnung_var,
            "menge": menge_var,
            "preis": preis_var,
            "ust": ust_var,
            "gesamt_label": gesamt_label
        })
        
        # Aktualisiere die Scrollregion
        self.angebot_artikel_scrollable_frame.update_idletasks()
        self.angebot_artikel_canvas.configure(scrollregion=self.angebot_artikel_canvas.bbox("all"))
    
    def angebot_artikel_zeile_entfernen(self):
        """Entfernt die letzte Artikelzeile aus der Angebots-Tabelle, wenn mehr als eine vorhanden ist"""
        if len(self.angebot_artikel_eintraege) > 1:
            # Letzten Eintrag entfernen
            letzter_eintrag = self.angebot_artikel_eintraege.pop()
            letzter_eintrag["frame"].destroy()
            
            # Aktualisiere die Scrollregion
            self.angebot_artikel_scrollable_frame.update_idletasks()
            self.angebot_artikel_canvas.configure(scrollregion=self.angebot_artikel_canvas.bbox("all"))
    
    def angebot_artikel_aus_katalog(self):
        """Öffnet ein Fenster zur Auswahl mehrerer Artikel aus dem Katalog für Angebote"""
        if not self.artikelbuch:
            messagebox.showinfo("Information", "Der Artikelkatalog ist leer. Bitte fügen Sie zuerst Artikel hinzu.")
            return
        
        # Diese Funktion funktioniert analog zu artikel_aus_katalog, aber für Angebote
        # Erstelle ein Toplevel-Fenster für die Artikelauswahl
        auswahl_fenster = tk.Toplevel(self.root)
        auswahl_fenster.title("Artikel aus Katalog auswählen")
        auswahl_fenster.geometry("800x500")
        
        # Linker Frame für die Artikel-Liste
        liste_frame = ttk.Frame(auswahl_fenster, padding=10)
        liste_frame.pack(side="left", fill="both", expand=True)
        
        ttk.Label(liste_frame, text="Verfügbare Artikel:").pack(anchor="w")
        
        # Suchfeld
        search_frame = ttk.Frame(liste_frame)
        search_frame.pack(fill="x", pady=5)
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        search_var = tk.StringVar()
        ttk.Entry(search_frame, textvariable=search_var, width=25).pack(side="left", padx=5)
        
        # Listbox mit Scrollbar
        list_container = ttk.Frame(liste_frame)
        list_container.pack(fill="both", expand=True)
        
        scrollbar = ttk.Scrollbar(list_container)
        scrollbar.pack(side="right", fill="y")
        
        artikel_listbox = tk.Listbox(list_container, yscrollcommand=scrollbar.set, height=20, width=50)
        artikel_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=artikel_listbox.yview)
        
        # Rechter Frame für die gewählten Artikel
        selected_frame = ttk.Frame(auswahl_fenster, padding=10)
        selected_frame.pack(side="right", fill="both", expand=True)
        
        ttk.Label(selected_frame, text="Ausgewählte Artikel:").pack(anchor="w")
        
        # Listbox für gewählte Artikel
        selected_list_container = ttk.Frame(selected_frame)
        selected_list_container.pack(fill="both", expand=True)
        
        selected_scrollbar = ttk.Scrollbar(selected_list_container)
        selected_scrollbar.pack(side="right", fill="y")
        
        selected_listbox = tk.Listbox(selected_list_container, yscrollcommand=selected_scrollbar.set, height=20, width=50)
        selected_listbox.pack(side="left", fill="both", expand=True)
        selected_scrollbar.config(command=selected_listbox.yview)
        
        # Buttons zwischen den Listen
        buttons_frame = ttk.Frame(auswahl_fenster)
        buttons_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        ttk.Button(buttons_frame, text=">>", 
                  command=lambda: self.add_to_selected(artikel_listbox, selected_listbox)).pack(pady=5)
        ttk.Button(buttons_frame, text="<<", 
                  command=lambda: self.remove_from_selected(selected_listbox)).pack(pady=5)
        
        # OK und Abbrechen Buttons
        bottom_frame = ttk.Frame(auswahl_fenster)
        bottom_frame.pack(side="bottom", fill="x", pady=10)
        
        ttk.Button(bottom_frame, text="Übernehmen", 
                  command=lambda: self.angebot_uebernehme_ausgewaehlte_artikel(
                      selected_listbox, auswahl_fenster)).pack(side="right", padx=10)
        ttk.Button(bottom_frame, text="Abbrechen", 
                  command=auswahl_fenster.destroy).pack(side="right")
        
        # Artikel in Listbox füllen
        self.refresh_artikel_liste(artikel_listbox, search_var.get())
        
        # Suchfunktion
        search_var.trace("w", lambda name, index, mode, 
                         lb=artikel_listbox, sv=search_var: self.refresh_artikel_liste(lb, sv.get()))
        
        # Fenster modal machen
        auswahl_fenster.transient(self.root)
        auswahl_fenster.grab_set()
        self.root.wait_window(auswahl_fenster)
    
    def angebot_uebernehme_ausgewaehlte_artikel(self, listbox, fenster):
        """Übernimmt die ausgewählten Artikel in die Angebots-Tabelle"""
        if listbox.size() == 0:
            fenster.destroy()
            return
        
        # Prüfe, ob genug Zeilen vorhanden sind, sonst füge hinzu
        while len(self.angebot_artikel_eintraege) < listbox.size():
            self.angebot_artikel_zeile_hinzufuegen()
        
        # Übernehme die Artikel in die Tabelle
        for i in range(listbox.size()):
            artikel_id = listbox.itemcget(i, 'artikel_id')
            if artikel_id in self.artikelbuch:
                artikel = self.artikelbuch[artikel_id]
                
                # Artikel in die i-te Zeile eintragen
                if i < len(self.angebot_artikel_eintraege):
                    self.angebot_artikel_eintraege[i]["bezeichnung"].set(artikel["bezeichnung"])
                    self.angebot_artikel_eintraege[i]["preis"].set(artikel["preis"])
                    self.angebot_artikel_eintraege[i]["ust"].set(artikel["ust"])
                    self.angebot_artikel_eintraege[i]["menge"].set("1")  # Standardmenge 1
        
        fenster.destroy()
    
    def angebot_waehle_artikel_aus_katalog(self, row_index):
        """Öffnet ein Fenster zur Auswahl eines einzelnen Artikels für eine bestimmte Zeile im Angebot"""
        # Diese Funktion ist analog zu waehle_artikel_aus_katalog, aber für Angebote
        if not self.artikelbuch:
            messagebox.showinfo("Information", "Der Artikelkatalog ist leer. Bitte fügen Sie zuerst Artikel hinzu.")
            return
        
        # Erstelle ein Toplevel-Fenster für die Artikelauswahl
        auswahl_fenster = tk.Toplevel(self.root)
        auswahl_fenster.title("Artikel auswählen")
        auswahl_fenster.geometry("500x400")
        
        # Suchfeld
        search_frame = ttk.Frame(auswahl_fenster, padding=10)
        search_frame.pack(fill="x")
        
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        search_var = tk.StringVar()
        ttk.Entry(search_frame, textvariable=search_var, width=25).pack(side="left", padx=5)
        
        # Listbox mit Scrollbar
        list_container = ttk.Frame(auswahl_fenster, padding=10)
        list_container.pack(fill="both", expand=True)
        
        scrollbar = ttk.Scrollbar(list_container)
        scrollbar.pack(side="right", fill="y")
        
        artikel_listbox = tk.Listbox(list_container, yscrollcommand=scrollbar.set, height=15, width=50)
        artikel_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=artikel_listbox.yview)
        
        # OK und Abbrechen Buttons
        button_frame = ttk.Frame(auswahl_fenster, padding=10)
        button_frame.pack(fill="x")
        
        ttk.Button(button_frame, text="Auswählen", 
                  command=lambda: self.angebot_artikel_in_zeile_einfuegen(
                      artikel_listbox, row_index, auswahl_fenster)).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Abbrechen", 
                  command=auswahl_fenster.destroy).pack(side="right")
        
        # Artikel in Listbox füllen
        self.refresh_artikel_liste(artikel_listbox)
        
        # Doppelklick-Bindung
        artikel_listbox.bind("<Double-1>", 
                            lambda event: self.angebot_artikel_in_zeile_einfuegen(
                                artikel_listbox, row_index, auswahl_fenster))
        
        # Suchfunktion
        search_var.trace("w", lambda name, index, mode, 
                         lb=artikel_listbox, sv=search_var: self.refresh_artikel_liste(lb, sv.get()))
        
        # Fenster modal machen
        auswahl_fenster.transient(self.root)
        auswahl_fenster.grab_set()
        self.root.wait_window(auswahl_fenster)
    
    def angebot_artikel_in_zeile_einfuegen(self, listbox, row_index, fenster):
        """Fügt den ausgewählten Artikel in die angegebene Zeile des Angebots ein"""
        selection = listbox.curselection()
        if not selection:
            return
        
        artikel_id = listbox.itemcget(selection[0], 'artikel_id')
        if artikel_id in self.artikelbuch:
            artikel = self.artikelbuch[artikel_id]
            
            # In die Zeile eintragen
            if 0 <= row_index < len(self.angebot_artikel_eintraege):
                self.angebot_artikel_eintraege[row_index]["bezeichnung"].set(artikel["bezeichnung"])
                self.angebot_artikel_eintraege[row_index]["preis"].set(artikel["preis"])
                self.angebot_artikel_eintraege[row_index]["ust"].set(artikel["ust"])
                self.angebot_artikel_eintraege[row_index]["menge"].set("1")  # Standardmenge 1
        
        fenster.destroy()
    
    def generiere_angebotsnummer(self):
        """Generiert eine neue Angebotsnummer basierend auf Jahr, Monat und Zähler"""
        heute = datetime.datetime.now()
        jahr = heute.strftime("%Y")
        monat = heute.strftime("%m")
        
        # Zähler aus Datei lesen oder neu anlegen
        counter_file = f"angebot_counter_{jahr}.json"
        counter = {}
        
        if os.path.exists(counter_file):
            with open(counter_file, "r") as f:
                counter = json.load(f)
        
        if monat not in counter:
            counter[monat] = 0
        
        counter[monat] += 1
        
        # Zähler speichern
        with open(counter_file, "w") as f:
            json.dump(counter, f)
        
        # Angebotsnummer formatieren: A-Jahr-Monat-Zähler
        return f"A-{jahr}-{monat}-{counter[monat]:03d}"
    
    def angebot_berechne_summen(self):
        """Berechnet die Summen für die Angebots-Vorschau"""
        empfaenger = self.angebot_empfaenger_var.get()
        if not empfaenger or empfaenger not in self.adressbuch:
            messagebox.showerror("Fehler", "Bitte wählen Sie einen gültigen Empfänger!")
            return
        
        artikel = []
        nettosumme = Decimal('0.00')
        ust_betraege = {}  # Um Steuerbeträge nach Steuersätzen zu gruppieren
        
        for eintrag in self.angebot_artikel_eintraege:
            bezeichnung = eintrag["bezeichnung"].get()
            
            # Nur Artikel mit Bezeichnung berücksichtigen
            if bezeichnung:
                try:
                    menge = Decimal(eintrag["menge"].get().replace(',', '.'))
                    preis = Decimal(eintrag["preis"].get().replace(',', '.'))
                    ust_satz = Decimal(eintrag["ust"].get().replace(',', '.'))
                    
                    position_netto = menge * preis
                    ust_betrag = position_netto * (ust_satz / 100)
                    position_brutto = position_netto + ust_betrag
                    
                    # Gesamtpreis in Label aktualisieren
                    eintrag["gesamt_label"].config(text=f"{position_brutto:.2f}")
                    
                    artikel.append({
                        "bezeichnung": bezeichnung,
                        "menge": menge,
                        "preis": preis,
                        "ust_satz": ust_satz,
                        "ust_betrag": ust_betrag,
                        "gesamt": position_brutto
                    })
                    
                    nettosumme += position_netto
                    
                    # Steuerbeträge nach Sätzen gruppieren
                    ust_key = str(ust_satz)
                    if ust_key not in ust_betraege:
                        ust_betraege[ust_key] = Decimal('0.00')
                    ust_betraege[ust_key] += ust_betrag
                    
                except (ValueError, Exception) as e:
                    messagebox.showerror("Fehler", f"Ungültige Zahl in Artikel '{bezeichnung}'! {str(e)}")
                    return
        
        if not artikel:
            messagebox.showerror("Fehler", "Bitte geben Sie mindestens einen Artikel ein!")
            return
        
        # Gesamtbeträge berechnen
        gesamt_ust = sum(ust_betraege.values())
        bruttosumme = nettosumme + gesamt_ust
        
        # Vorschau anzeigen
        self.angebot_vorschau_text.config(state=tk.NORMAL)
        self.angebot_vorschau_text.delete(1.0, tk.END)
        
        self.angebot_vorschau_text.insert(tk.END, f"Empfänger: {empfaenger}\n\n")
        self.angebot_vorschau_text.insert(tk.END, f"Angebotsnummer: {self.angebotsnr_var.get()}\n")
        self.angebot_vorschau_text.insert(tk.END, f"Datum: {self.angebotsdatum_var.get()}\n")
        self.angebot_vorschau_text.insert(tk.END, f"Gültig bis: {self.angebot_gueltig_bis_var.get()}\n\n")
        self.angebot_vorschau_text.insert(tk.END, f"Positionen: {len(artikel)}\n")
        
        # Steueraufstellung
        self.angebot_vorschau_text.insert(tk.END, "\nSteueraufstellung:\n")
        for satz, betrag in ust_betraege.items():
            self.angebot_vorschau_text.insert(tk.END, f"  {satz}% USt: {betrag:.2f} €\n")
        
        # Summen
        self.angebot_vorschau_text.insert(tk.END, f"\nNettosumme: {nettosumme:.2f} €\n")
        self.angebot_vorschau_text.insert(tk.END, f"USt gesamt: {gesamt_ust:.2f} €\n")
        self.angebot_vorschau_text.insert(tk.END, f"Bruttosumme: {bruttosumme:.2f} €\n")
        
        self.angebot_vorschau_text.config(state=tk.DISABLED)
        
        return {
            "empfaenger": empfaenger,
            "artikel": artikel,
            "nettosumme": nettosumme,
            "ust_betraege": ust_betraege,
            "bruttosumme": bruttosumme
        }
    
    def angebot_speichern(self):
        """Speichert das aktuelle Angebot in der Datenbank"""
        # Daten sammeln
        empfaenger = self.angebot_empfaenger_var.get()
        if not empfaenger or empfaenger not in self.adressbuch:
            messagebox.showerror("Fehler", "Bitte wählen Sie einen gültigen Empfänger!")
            return
        
        angebotsnr = self.angebotsnr_var.get()
        if not angebotsnr:
            angebotsnr = self.generiere_angebotsnummer()
            self.angebotsnr_var.set(angebotsnr)
        
        # Artikel sammeln
        artikel = []
        for eintrag in self.angebot_artikel_eintraege:
            bezeichnung = eintrag["bezeichnung"].get()
            if bezeichnung:
                artikel.append({
                    "bezeichnung": bezeichnung,
                    "menge": eintrag["menge"].get(),
                    "preis": eintrag["preis"].get(),
                    "ust": eintrag["ust"].get()
                })
        
        if not artikel:
            messagebox.showerror("Fehler", "Bitte geben Sie mindestens einen Artikel ein!")
            return
        
        # Angebot in Datenbank speichern
        angebot_daten = {
            "nummer": angebotsnr,
            "empfaenger": empfaenger,
            "empfaenger_daten": self.adressbuch[empfaenger],
            "datum": self.angebotsdatum_var.get(),
            "gueltig_bis": self.angebot_gueltig_bis_var.get(),
            "artikel": artikel,
            "typ": self.angebot_typ_var.get()
        }
        
        # In Datei speichern
        if not os.path.exists("angebote"):
            os.makedirs("angebote")
        
        with open(f"angebote/{angebotsnr}.json", "w", encoding="utf-8") as f:
            json.dump(angebot_daten, f, ensure_ascii=False, indent=4)
        
        # Angebotsliste aktualisieren
        self.aktualisiere_angebote_liste()
        
        messagebox.showinfo("Erfolg", f"Angebot {angebotsnr} wurde gespeichert!")
    
    def aktualisiere_angebote_liste(self):
        """Aktualisiert die Liste der gespeicherten Angebote"""
        self.angebote_listbox.delete(0, tk.END)
        
        # Prüfe, ob der Angebote-Ordner existiert
        if not os.path.exists("angebote"):
            return
        
        # Alle Angebote aus dem Ordner laden
        angebote = []
        for filename in os.listdir("angebote"):
            if filename.endswith(".json"):
                try:
                    with open(os.path.join("angebote", filename), "r", encoding="utf-8") as f:
                        angebot = json.load(f)
                        angebote.append({
                            "nummer": angebot.get("nummer", ""),
                            "empfaenger": angebot.get("empfaenger", ""),
                            "datum": angebot.get("datum", ""),
                            "filename": filename
                        })
                except Exception:
                    pass
        
        # Nach Datum sortieren
        angebote.sort(key=lambda x: x["datum"], reverse=True)
        
        # In Listbox einfügen
        for angebot in angebote:
            display = f"{angebot['nummer']} - {angebot['datum']} - {angebot['empfaenger']}"
            self.angebote_listbox.insert(tk.END, display)
            self.angebote_listbox.itemconfig(self.angebote_listbox.size()-1, {'filename': angebot['filename']})
    
    def neues_angebot(self):
        """Erstellt ein neues leeres Angebot"""
        # Felder zurücksetzen
        self.angebot_empfaenger_var.set("")
        self.angebotsnr_var.set(self.generiere_angebotsnummer())
        self.angebotsdatum_var.set(datetime.datetime.now().strftime("%d.%m.%Y"))
        
        # Gültigkeit standardmäßig 4 Wochen
        gueltig_bis = datetime.datetime.now() + datetime.timedelta(days=28)
        self.angebot_gueltig_bis_var.set(gueltig_bis.strftime("%d.%m.%Y"))
        
        # Artikel zurücksetzen
        for eintrag in self.angebot_artikel_eintraege:
            eintrag["bezeichnung"].set("")
            eintrag["menge"].set("1")
            eintrag["preis"].set("0.00")
            eintrag["ust"].set("19")
            eintrag["gesamt_label"].config(text="0.00")
        
        # Vorschau leeren
        self.angebot_vorschau_text.config(state=tk.NORMAL)
        self.angebot_vorschau_text.delete(1.0, tk.END)
        self.angebot_vorschau_text.config(state=tk.DISABLED)
    
    def angebot_auswaehlen(self, event):
        """Handler für die Auswahl eines Angebots aus der Liste"""
        selection = self.angebote_listbox.curselection()
        if not selection:
            return
        
        # Dateinamen des ausgewählten Angebots ermitteln
        filename = self.angebote_listbox.itemcget(selection[0], 'filename')
        if not filename:
            return
        
        # Angebot laden
        self.lade_angebot_aus_datei(os.path.join("angebote", filename))
    
    def angebot_laden(self):
        """Lädt ein ausgewähltes Angebot aus der Liste"""
        selection = self.angebote_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie ein Angebot aus der Liste aus.")
            return
        
        # Dateinamen des ausgewählten Angebots ermitteln
        filename = self.angebote_listbox.itemcget(selection[0], 'filename')
        if not filename:
            return
        
        # Angebot laden
        self.lade_angebot_aus_datei(os.path.join("angebote", filename))
    
    def lade_angebot_aus_datei(self, filepath):
        """Lädt ein Angebot aus einer Datei"""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                angebot = json.load(f)
            
            # Felder setzen
            self.angebot_empfaenger_var.set(angebot.get("empfaenger", ""))
            self.angebotsnr_var.set(angebot.get("nummer", ""))
            self.angebotsdatum_var.set(angebot.get("datum", ""))
            self.angebot_gueltig_bis_var.set(angebot.get("gueltig_bis", ""))
            
            # Typ setzen
            typ = angebot.get("typ", "Artikel")
            self.angebot_typ_var.set(typ)
            
            # Tabelle zurücksetzen
            for widget in self.angebot_artikel_container_frame.winfo_children():
                widget.destroy()
            
            self.angebot_artikel_eintraege = []
            self.init_angebot_artikeltabelle()
            
            # Artikel befüllen
            artikel_liste = angebot.get("artikel", [])
            
            # Bei Bedarf Zeilen hinzufügen
            while len(self.angebot_artikel_eintraege) < len(artikel_liste):
                self.angebot_artikel_zeile_hinzufuegen()
            
            # Artikel übertragen
            for i, artikel_item in enumerate(artikel_liste):
                if i < len(self.angebot_artikel_eintraege):
                    self.angebot_artikel_eintraege[i]["bezeichnung"].set(artikel_item.get("bezeichnung", ""))
                    self.angebot_artikel_eintraege[i]["menge"].set(artikel_item.get("menge", "1"))
                    self.angebot_artikel_eintraege[i]["preis"].set(artikel_item.get("preis", "0.00"))
                    self.angebot_artikel_eintraege[i]["ust"].set(artikel_item.get("ust", "19"))
            
            # Summen neuberechnen
            self.angebot_berechne_summen()
            
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Laden des Angebots: {str(e)}")
    
    def angebot_loeschen(self):
        """Löscht ein ausgewähltes Angebot"""
        selection = self.angebote_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie ein Angebot aus der Liste aus.")
            return
        
        # Dateinamen des ausgewählten Angebots ermitteln
        filename = self.angebote_listbox.itemcget(selection[0], 'filename')
        if not filename:
            return
        
        # Angebotsnummer aus Listbox extrahieren
        display_text = self.angebote_listbox.get(selection[0])
        angebots_nr = display_text.split(' - ')[0]
        
        # Bestätigung einholen
        if not messagebox.askyesno("Löschen bestätigen", 
                                  f"Möchten Sie das Angebot {angebots_nr} wirklich löschen?"):
            return
        
        # Angebot löschen
        try:
            os.remove(os.path.join("angebote", filename))
            messagebox.showinfo("Erfolg", f"Angebot {angebots_nr} wurde gelöscht!")
            self.aktualisiere_angebote_liste()
            self.neues_angebot()
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Löschen des Angebots: {str(e)}")
    
    def angebot_zu_rechnung(self):
        """Übernimmt ein Angebot als neue Rechnung"""
        selection = self.angebote_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie ein Angebot aus der Liste aus.")
            return
        
        # Dateinamen des ausgewählten Angebots ermitteln
        filename = self.angebote_listbox.itemcget(selection[0], 'filename')
        if not filename:
            return
        
        try:
            # Angebot laden
            with open(os.path.join("angebote", filename), "r", encoding="utf-8") as f:
                angebot = json.load(f)
            
            # Zum Rechnungs-Tab wechseln
            self.notebook.select(self.tab_rechnung)
            
            # Felder übernehmen
            self.empfaenger_var.set(angebot.get("empfaenger", ""))
            self.rechnungstyp_var.set(angebot.get("typ", "Artikel"))
            
            # Tabelle neu aufbauen wegen möglicher Typen-Änderung
            self.update_artikeltabelle_header()
            
            # Artikel übernehmen
            artikel_liste = angebot.get("artikel", [])
            
            # Bei Bedarf Zeilen hinzufügen
            while len(self.artikel_eintraege) < len(artikel_liste):
                self.artikel_zeile_hinzufuegen()
            
            # Artikel übertragen
            for i, artikel_item in enumerate(artikel_liste):
                if i < len(self.artikel_eintraege):
                    self.artikel_eintraege[i]["bezeichnung"].set(artikel_item.get("bezeichnung", ""))
                    self.artikel_eintraege[i]["menge"].set(artikel_item.get("menge", "1"))
                    self.artikel_eintraege[i]["preis"].set(artikel_item.get("preis", "0.00"))
                    self.artikel_eintraege[i]["ust"].set(artikel_item.get("ust", "19"))
            
            # Rechnungsnummer und Datum neu setzen
            self.rechnungsnr_var.set(self.generiere_rechnungsnummer())
            self.rechnungsdatum_var.set(datetime.datetime.now().strftime("%d.%m.%Y"))
            
            # Summen neuberechnen
            self.berechne_summen()
            
            messagebox.showinfo("Erfolg", f"Angebot wurde als Rechnung übernommen.")
            
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Übernehmen des Angebots: {str(e)}")
    
    def erstelle_angebot(self, zum_versenden=False):
        """Erstellt ein Angebotsdokument basierend auf den eingegebenen Daten"""
        # Prüfen, ob ein Briefkopf existiert
        if not self.firmendaten.get("briefkopf_pfad") or not os.path.exists(self.firmendaten["briefkopf_pfad"]):
            messagebox.showerror("Fehler", "Bitte erstellen oder laden Sie zuerst einen Briefkopf!")
            return None
        
        # Empfänger prüfen
        empfaenger_name = self.angebot_empfaenger_var.get()
        if not empfaenger_name or empfaenger_name not in self.adressbuch:
            messagebox.showerror("Fehler", "Bitte wählen Sie einen gültigen Empfänger!")
            return None
        
        empfaenger = self.adressbuch[empfaenger_name]
        
        # Angebot berechnen
        ergebnis = self.angebot_berechne_summen()
        if not ergebnis:
            return None
        
        nettosumme = ergebnis["nettosumme"]
        ust_betraege = ergebnis["ust_betraege"]
        bruttosumme = ergebnis["bruttosumme"]
        artikel = ergebnis["artikel"]
        
        # Angebotsdokument erstellen
        try:
            # Vorlage kopieren
            doc = docx.Document(self.firmendaten["briefkopf_pfad"])
            
            # Empfängerdaten eintragen
            for i, para in enumerate(doc.paragraphs):
                if "Empfänger:" in para.text:
                    # Die nächsten Paragraphen mit Empfängerdaten überschreiben
                    empfaenger_zeilen = []
                    
                    # Name oder Firma
                    if empfaenger.get("Firma"):
                        empfaenger_zeilen.append(empfaenger["Firma"])
                        if empfaenger.get("Vorname") and empfaenger.get("Nachname"):
                            empfaenger_zeilen.append(f"{empfaenger['Vorname']} {empfaenger['Nachname']}")
                        elif empfaenger.get("Vorname"):
                            empfaenger_zeilen.append(empfaenger["Vorname"])
                        elif empfaenger.get("Nachname"):
                            empfaenger_zeilen.append(empfaenger["Nachname"])
                    else:
                        name_str = ""
                        if empfaenger.get("Vorname"):
                            name_str += empfaenger["Vorname"] + " "
                        if empfaenger.get("Nachname"):
                            name_str += empfaenger["Nachname"]
                        empfaenger_zeilen.append(name_str.strip())
                    
                    # Adresse
                    empfaenger_zeilen.append(empfaenger.get("Straße", ""))
                    empfaenger_zeilen.append(f"{empfaenger.get('PLZ', '')} {empfaenger.get('Ort', '')}")
                    
                    # In Dokument schreiben
                    for j, zeile in enumerate(empfaenger_zeilen):
                        if i+j+1 < len(doc.paragraphs):
                            doc.paragraphs[i+j+1].text = zeile
                    break
            
            # Text "RECHNUNG" durch "ANGEBOT" ersetzen und Gültigkeit hinzufügen
            for para in doc.paragraphs:
                if "RECHNUNG" in para.text:
                    para.text = "ANGEBOT"
                    para.runs[0].bold = True
                    
                    # Gültigkeitsvermerk als neuen Absatz hinzufügen
                    gueltig_bis = self.angebot_gueltig_bis_var.get()
                    if gueltig_bis:
                        p = doc.add_paragraph()
                        p.text = f"Dieses Angebot ist gültig bis zum {gueltig_bis}."
                        p.style = 'Normal'
                    break
            
            # Angebotsinformationen
            angebotsnr = self.angebotsnr_var.get()
            datum = self.angebotsdatum_var.get()
            
            for table in doc.tables:
                if len(table.rows) >= 3 and "Rechnungsnummer:" in table.rows[0].cells[0].text:
                    # Text "Rechnungsnummer" zu "Angebotsnummer" ändern
                    table.rows[0].cells[0].text = "Angebotsnummer:"
                    table.rows[0].cells[1].text = angebotsnr
                    
                    # Text "Datum" bleibt
                    table.rows[1].cells[1].text = datum
                    
                    # Debitorennummer als Kundennummer verwenden, falls vorhanden
                    if empfaenger.get("Debitorennummer"):
                        kundennr = empfaenger["Debitorennummer"]
                    else:
                        # Sonst eine einfache Kundennummer generieren
                        kundennr = f"{empfaenger_name[0].upper()}{hash(empfaenger_name) % 10000:04d}"
                    
                    table.rows[2].cells[1].text = kundennr
                    
                    # Leistungszeitraum-Zeile entfernen, wenn vorhanden
                    if len(table.rows) >= 4 and "Leistungszeitraum:" in table.rows[3].cells[0].text:
                        # Zeile entfernen oder Text ändern
                        table.rows[3].cells[0].text = "Gültig bis:"
                        table.rows[3].cells[1].text = self.angebot_gueltig_bis_var.get()
                    
                    break
            
            # Artikel in Tabelle einfügen
            artikel_tabelle = None
            for table in doc.tables:
                if len(table.rows) >= 1 and len(table.rows[0].cells) >= 5:
                    if "Bezeichnung" in table.rows[0].cells[0].text and "Anzahl" in table.rows[0].cells[1].text:
                        artikel_tabelle = table
                        break
            
            if artikel_tabelle:
                # Bei Leistungen Header anpassen
                ist_leistung = self.angebot_typ_var.get() == "Leistung"
                if ist_leistung and "Anzahl" in artikel_tabelle.rows[0].cells[1].text:
                    artikel_tabelle.rows[0].cells[1].text = "Stunden"
                    artikel_tabelle.rows[0].cells[2].text = "Stundensatz €"
                
                # Beispielzeile entfernen, falls vorhanden
                if len(artikel_tabelle.rows) > 1:
                    if "<ARTIKEL>" in artikel_tabelle.rows[1].cells[0].text:
                        artikel_tabelle._tbl.remove(artikel_tabelle.rows[1]._tr)
                
                # Artikel hinzufügen
                for artikel_item in artikel:
                    row = artikel_tabelle.add_row().cells
                    row[0].text = artikel_item["bezeichnung"]
                    row[1].text = f"{artikel_item['menge']}"
                    row[2].text = f"{artikel_item['preis']:.2f}"
                    row[3].text = f"{artikel_item['ust_satz']} %"
                    row[4].text = f"{artikel_item['gesamt']:.2f}"
            
            # Summentabelle finden und aktualisieren
            sum_table = None
            for table in doc.tables:
                if len(table.rows) >= 3 and "Nettobetrag:" in table.rows[0].cells[0].text:
                    sum_table = table
                    break
            
            if sum_table:
                sum_table.rows[0].cells[1].text = f"{nettosumme:.2f} €"
                
                # Bei mehreren Steuersätzen detailliert aufschlüsseln
                if len(ust_betraege) > 1:
                    # Alle vorhandenen Zeilen ab der zweiten löschen (Umsatzsteuer, Gesamtbetrag)
                    while len(sum_table.rows) > 1:
                        sum_table._tbl.remove(sum_table.rows[-1]._tr)
                    
                    # Zeilen für jeden Steuersatz einfügen
                    for satz, betrag in sorted(ust_betraege.items()):
                        row_cells = sum_table.add_row().cells
                        row_cells[0].text = f"USt {satz} %:"
                        row_cells[1].text = f"{betrag:.2f} €"
                    
                    # Zeile für Gesamtbetrag hinzufügen
                    row_cells = sum_table.add_row().cells
                    row_cells[0].text = "Gesamtbetrag:"
                    row_cells[1].text = f"{bruttosumme:.2f} €"
                else:
                    # Standard mit einem Steuersatz
                    # Sicherstellen, dass genau 3 Zeilen existieren
                    while len(sum_table.rows) > 3:
                        sum_table._tbl.remove(sum_table.rows[-1]._tr)
                    
                    while len(sum_table.rows) < 3:
                        sum_table.add_row()
                    
                    # Werte setzen
                    if len(ust_betraege) > 0:
                        gesamt_ust = sum(ust_betraege.values())
                        sum_table.rows[1].cells[0].text = "Umsatzsteuer:"
                        sum_table.rows[1].cells[1].text = f"{gesamt_ust:.2f} €"
                    else:
                        sum_table.rows[1].cells[0].text = "Umsatzsteuer:"
                        sum_table.rows[1].cells[1].text = "0.00 €"
                    
                    sum_table.rows[2].cells[0].text = "Gesamtbetrag:"
                    sum_table.rows[2].cells[1].text = f"{bruttosumme:.2f} €"
            
            # Angebot speichern
            if zum_versenden:
                # Temporäres Speichern für E-Mail-Versand
                temp_docx = os.path.join(os.environ.get("TEMP", "."), f"Angebot_{angebotsnr}.docx")
                doc.save(temp_docx)
                return {"docx": temp_docx, "angebotsnr": angebotsnr, "empfaenger": empfaenger}
            else:
                filepath = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Dokument", "*.docx")],
                    initialfile=f"Angebot_{angebotsnr}.docx",
                    title="Angebot speichern"
                )
                
                if filepath:
                    doc.save(filepath)
                    messagebox.showinfo("Erfolg", f"Angebot wurde unter {filepath} gespeichert!")
                    
                    # Angebot in Datenbank speichern
                    self.angebot_speichern()
                    
                return filepath
        
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Erstellen des Angebots: {str(e)}")
            return None
    
    def erstelle_angebot_und_versende(self):
        """Erstellt ein Angebot und öffnet den E-Mail-Client zum Versenden"""
        # Angebot erstellen
        ergebnis = self.erstelle_angebot(zum_versenden=True)
        
        if ergebnis:
            docx_pfad = ergebnis["docx"]
            angebotsnr = ergebnis["angebotsnr"]
            empfaenger = ergebnis["empfaenger"]
            
            email = empfaenger.get("Email", "")
            
            # E-Mail-Betreff
            betreff = f"Angebot {angebotsnr}"
            
            # E-Mail-Text
            email_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "vielen Dank für Ihre Anfrage. Anbei erhalten Sie unser Angebot. "
                "Bei Rückfragen stehen wir Ihnen gerne zur Verfügung.\n\n"
                "Mit freundlichen Grüßen"
            )
            
            # Mail-Client öffnen
            try:
                # Outlook öffnen
                if sys.platform == "win32":
                    # Windows mit Outlook
                    os.system(f'start outlook.exe /c ipm.note /m "{email}" /a "{docx_pfad}" /s /t "{betreff} {email_text}"')
                else:
                    # Andere Betriebssysteme - Standard-Mail-Client
                    import webbrowser
                    mailto_link = f"mailto:{email}?subject={betreff}&body={email_text}"
                    webbrowser.open(mailto_link)
                    messagebox.showinfo("Hinweis", 
                                       f"E-Mail-Client wurde geöffnet. Bitte hängen Sie das Angebot manuell an:\n{docx_pfad}")
            
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Öffnen des E-Mail-Clients: {str(e)}")
            
            # Angebot in Datenbank speichern
            self.angebot_speichern()
    
    def init_artikel_tab(self):
        """Initialisiert den Tab für die Artikelverwaltung"""
        frame = ttk.Frame(self.tab_artikel, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Linke Seite: Liste der Artikel
        liste_frame = ttk.Frame(frame)
        liste_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        ttk.Label(liste_frame, text="Gespeicherte Artikel:").pack(anchor="w")
        
        # Suchfeld
        search_frame = ttk.Frame(liste_frame)
        search_frame.pack(fill="x", pady=5)
        
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        self.artikel_search_var = tk.StringVar()
        self.artikel_search_var.trace("w", self.search_artikelbuch)
        ttk.Entry(search_frame, textvariable=self.artikel_search_var, width=20).pack(side="left", padx=5)
        
        # Import/Export Buttons
        import_export_frame = ttk.Frame(liste_frame)
        import_export_frame.pack(fill="x", pady=5)
        
        ttk.Button(import_export_frame, text="Importieren", 
                  command=self.artikelbuch_importieren).pack(side="left", padx=5)
        ttk.Button(import_export_frame, text="Exportieren", 
                  command=self.artikelbuch_exportieren).pack(side="left", padx=5)
        
        # Artikelliste
        self.artikel_listbox = tk.Listbox(liste_frame, height=15, width=40)
        self.artikel_listbox.pack(side="top", fill="both", expand=True, pady=5)
        self.artikel_listbox.bind('<<ListboxSelect>>', self.artikel_auswaehlen)
        
        button_frame = ttk.Frame(liste_frame)
        button_frame.pack(side="top", fill="x", pady=5)
        
        ttk.Button(button_frame, text="Neu", command=self.neuer_artikel).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Löschen", command=self.loesche_artikel).pack(side="left", padx=5)
        
        # Rechte Seite: Artikeldetails
        details_frame = ttk.Frame(frame)
        details_frame.pack(side="right", fill="both", expand=True)
        
        felder = [
            "Artikelnummer", "Bezeichnung", "Beschreibung", "Einheit", 
            "Netto-Preis", "USt-Satz", "Kategorie"
        ]
        
        self.artikel_eintraege_form = {}
        
        for i, feld in enumerate(felder):
            ttk.Label(details_frame, text=f"{feld}:").grid(row=i, column=0, sticky="w", pady=5)
            var = tk.StringVar()
            
            if feld == "USt-Satz":
                widget = ttk.Combobox(details_frame, textvariable=var, width=40, values=["0", "7", "19"])
                var.set("19")  # Standardwert
            else:
                widget = ttk.Entry(details_frame, width=40, textvariable=var)
            
            widget.grid(row=i, column=1, sticky="ew", pady=5, padx=5)
            self.artikel_eintraege_form[feld] = var
        
        ttk.Button(details_frame, text="Speichern", command=self.speichere_artikel).grid(
            row=len(felder), column=0, columnspan=2, pady=10
        )
        
        # Artikelgruppen-Verwaltung
        gruppe_frame = ttk.LabelFrame(details_frame, text="Artikelgruppen-Verwaltung")
        gruppe_frame.grid(row=len(felder)+1, column=0, columnspan=2, pady=10, sticky="ew")
        
        ttk.Label(gruppe_frame, text="Gruppe:").grid(row=0, column=0, sticky="w", pady=5)
        self.gruppe_name_var = tk.StringVar()
        self.gruppe_combo = ttk.Combobox(gruppe_frame, textvariable=self.gruppe_name_var, width=30)
        self.gruppe_combo.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        gruppe_btn_frame = ttk.Frame(gruppe_frame)
        gruppe_btn_frame.grid(row=1, column=0, columnspan=2, pady=5)
        
        ttk.Button(gruppe_btn_frame, text="Neue Gruppe", 
                 command=self.neue_artikelgruppe).pack(side="left", padx=5)
        ttk.Button(gruppe_btn_frame, text="Löschen", 
                 command=self.loesche_artikelgruppe).pack(side="left", padx=5)
        ttk.Button(gruppe_btn_frame, text="Artikel hinzufügen", 
                 command=self.artikel_zu_gruppe_hinzufuegen).pack(side="left", padx=5)
        ttk.Button(gruppe_btn_frame, text="Aus Gruppe entfernen", 
                 command=self.artikel_aus_gruppe_entfernen).pack(side="left", padx=5)
        
        # Artikelliste für diese Gruppe
        ttk.Label(gruppe_frame, text="Artikel in dieser Gruppe:").grid(row=2, column=0, columnspan=2, sticky="w", pady=5)
        
        self.gruppe_artikel_listbox = tk.Listbox(gruppe_frame, height=5)
        self.gruppe_artikel_listbox.grid(row=3, column=0, columnspan=2, sticky="ew", pady=5)
        
        self.aktualisiere_artikelliste()
        self.aktualisiere_artikelgruppen_liste()
        self.gruppe_combo.bind("<<ComboboxSelected>>", self.aktualisiere_gruppe_artikelliste)
    
    def aktualisiere_artikelliste(self, filter_text=None):
        """Aktualisiert die Artikelliste basierend auf dem Suchbegriff"""
        self.artikel_listbox.delete(0, tk.END)
        
        # Wir speichern die Artikel-IDs in einer separaten Liste
        self.artikel_listbox_id_mapping = []
        
        # Sortierte und formatierte Liste erstellen
        artikel_liste = []
        for artikel_id, artikel in self.artikelbuch.items():
            artikelnummer = artikel.get("Artikelnummer", "")
            bezeichnung = artikel.get("Bezeichnung", "")
            preis = artikel.get("Netto-Preis", "0.00")
            
            display_text = f"{artikelnummer} - {bezeichnung} ({preis}€)"
            
            # Bei Suchfilter überprüfen
            if filter_text is None or filter_text.lower() in display_text.lower():
                artikel_liste.append((artikelnummer, bezeichnung, display_text, artikel_id))
        
        # Nach Artikelnummer, Bezeichnung sortieren
        artikel_liste.sort()
        
        # Anzeigen
        for _, _, display_text, artikel_id in artikel_liste:
            self.artikel_listbox.insert(tk.END, display_text)
            # Artikel-ID in separater Liste speichern
            self.artikel_listbox_id_mapping.append(artikel_id)
    
    def search_artikelbuch(self, *args):
        """Sucht im Artikelbuch nach dem eingegebenen Suchbegriff"""
        search_text = self.artikel_search_var.get()
        self.aktualisiere_artikelliste(search_text)
    
    def aktualisiere_artikelgruppen_liste(self):
        """Aktualisiert die Combobox mit den verfügbaren Artikelgruppen"""
        gruppen = sorted(list(self.artikelgruppen.keys()))
        
        # Combobox im Artikel-Tab (nur aktualisieren, wenn sie bereits existiert)
        if hasattr(self, 'gruppe_combo'):
            self.gruppe_combo['values'] = gruppen
    
        # Combobox im Rechnungs-Tab (nur aktualisieren, wenn sie bereits existiert)
        if hasattr(self, 'artikelgruppe_combo'):
            self.artikelgruppe_combo['values'] = gruppen
    
    def aktualisiere_gruppe_artikelliste(self, event=None):
        """Aktualisiert die Liste der Artikel in der ausgewählten Gruppe"""
        gruppe = self.gruppe_name_var.get()
        self.gruppe_artikel_listbox.delete(0, tk.END)
        
        if gruppe and gruppe in self.artikelgruppen:
            artikel_ids = self.artikelgruppen[gruppe]
            
            for artikel_id in artikel_ids:
                if artikel_id in self.artikelbuch:
                    artikel = self.artikelbuch[artikel_id]
                    bezeichnung = artikel.get("Bezeichnung", "")
                    preis = artikel.get("Netto-Preis", "0.00")
                    
                    display_text = f"{bezeichnung} ({preis}€)"
                    self.gruppe_artikel_listbox.insert(tk.END, display_text)
                    # Artikel-ID als Attribut speichern
                    self.gruppe_artikel_listbox.itemconfig(
                        self.gruppe_artikel_listbox.size()-1, {'artikel_id': artikel_id})
    
    def artikel_auswaehlen(self, event):
        """Handler für die Auswahl eines Artikels aus der Liste"""
        selection = self.artikel_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        if 0 <= index < len(self.artikel_listbox_id_mapping):
            artikel_id = self.artikel_listbox_id_mapping[index]
            
            if artikel_id in self.artikelbuch:
                artikel = self.artikelbuch[artikel_id]
                
                # Formular befüllen
                self.artikel_eintraege_form["Artikelnummer"].set(artikel.get("Artikelnummer", ""))
                self.artikel_eintraege_form["Bezeichnung"].set(artikel.get("Bezeichnung", ""))
                self.artikel_eintraege_form["Beschreibung"].set(artikel.get("Beschreibung", ""))
                self.artikel_eintraege_form["Einheit"].set(artikel.get("Einheit", ""))
                self.artikel_eintraege_form["Netto-Preis"].set(artikel.get("Netto-Preis", "0.00"))
                self.artikel_eintraege_form["USt-Satz"].set(artikel.get("USt-Satz", "19"))
                self.artikel_eintraege_form["Kategorie"].set(artikel.get("Kategorie", ""))
    
    def neuer_artikel(self):
        """Leert das Artikelformular für einen neuen Eintrag"""
        for var in self.artikel_eintraege_form.values():
            var.set("")
        
        # Standardwerte setzen
        self.artikel_eintraege_form["USt-Satz"].set("19")
    
    def loesche_artikel(self):
        """Löscht den ausgewählten Artikel"""
        selection = self.artikel_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie einen Artikel aus.")
            return
        
        index = selection[0]
        if 0 <= index < len(self.artikel_listbox_id_mapping):
            artikel_id = self.artikel_listbox_id_mapping[index]
            display_text = self.artikel_listbox.get(index)
            
            if messagebox.askyesno("Löschen bestätigen", 
                                  f"Möchten Sie den Artikel '{display_text}' wirklich löschen?"):
                
                # Aus Artikelbuch entfernen
                if artikel_id in self.artikelbuch:
                    del self.artikelbuch[artikel_id]
                
                # Aus allen Gruppen entfernen
                for gruppe, artikel_ids in self.artikelgruppen.items():
                    if artikel_id in artikel_ids:
                        self.artikelgruppen[gruppe] = [aid for aid in artikel_ids if aid != artikel_id]
                
                # Daten speichern
                self.speichere_daten()
                
                # Listen aktualisieren
                self.aktualisiere_artikelliste()
                self.aktualisiere_gruppe_artikelliste()
                
                # Formular leeren
                self.neuer_artikel()
    
    def speichere_artikel(self):
        """Speichert den aktuellen Artikel im Artikelbuch"""
        artikelnummer = self.artikel_eintraege_form["Artikelnummer"].get()
        bezeichnung = self.artikel_eintraege_form["Bezeichnung"].get()
        
        if not bezeichnung:
            messagebox.showerror("Fehler", "Bitte geben Sie mindestens eine Bezeichnung ein!")
            return
        
        # Artikel-ID erstellen oder bestehende verwenden
        if not artikelnummer:
            # Generiere eine ID basierend auf der Bezeichnung
            base_id = re.sub(r'[^\w]', '_', bezeichnung.lower())
            artikel_id = base_id
            
            # Falls ID bereits existiert, füge Nummer hinzu
            counter = 1
            while artikel_id in self.artikelbuch and counter < 100:
                artikel_id = f"{base_id}_{counter}"
                counter += 1
        else:
            artikel_id = re.sub(r'[^\w]', '_', artikelnummer.lower())
        
        # Artikel-Daten sammeln
        artikel = {
            "Artikelnummer": artikelnummer,
            "Bezeichnung": bezeichnung,
            "Beschreibung": self.artikel_eintraege_form["Beschreibung"].get(),
            "Einheit": self.artikel_eintraege_form["Einheit"].get(),
            "Netto-Preis": self.artikel_eintraege_form["Netto-Preis"].get().replace(',', '.'),
            "USt-Satz": self.artikel_eintraege_form["USt-Satz"].get(),
            "Kategorie": self.artikel_eintraege_form["Kategorie"].get(),
            
            # Felder für Verwendung in Rechnungen
            "bezeichnung": bezeichnung,
            "preis": self.artikel_eintraege_form["Netto-Preis"].get().replace(',', '.'),
            "ust": self.artikel_eintraege_form["USt-Satz"].get()
        }
        
        # Im Artikelbuch speichern
        self.artikelbuch[artikel_id] = artikel
        
        # Daten speichern
        self.speichere_daten()
        
        # Artikelliste aktualisieren
        self.aktualisiere_artikelliste()
        
        messagebox.showinfo("Erfolg", f"Artikel '{bezeichnung}' wurde gespeichert!")
    
    def neue_artikelgruppe(self):
        """Erstellt eine neue Artikelgruppe"""
        gruppe_name = simpledialog.askstring("Neue Artikelgruppe", 
                                            "Bitte geben Sie den Namen der neuen Artikelgruppe ein:")
        
        if not gruppe_name:
            return
        
        if gruppe_name in self.artikelgruppen:
            messagebox.showerror("Fehler", f"Die Artikelgruppe '{gruppe_name}' existiert bereits!")
            return
        
        # Neue leere Gruppe erstellen
        self.artikelgruppen[gruppe_name] = []
        
        # Daten speichern
        self.speichere_daten()
        
        # Listen aktualisieren
        self.aktualisiere_artikelgruppen_liste()
        self.gruppe_name_var.set(gruppe_name)
        self.aktualisiere_gruppe_artikelliste()
    
    def loesche_artikelgruppe(self):
        """Löscht die ausgewählte Artikelgruppe"""
        gruppe_name = self.gruppe_name_var.get()
        
        if not gruppe_name or gruppe_name not in self.artikelgruppen:
            messagebox.showinfo("Information", "Bitte wählen Sie eine gültige Artikelgruppe aus.")
            return
        
        if messagebox.askyesno("Löschen bestätigen", 
                              f"Möchten Sie die Artikelgruppe '{gruppe_name}' wirklich löschen?"):
            
            # Aus Artikelgruppen entfernen
            del self.artikelgruppen[gruppe_name]
            
            # Daten speichern
            self.speichere_daten()
            
            # Listen aktualisieren
            self.aktualisiere_artikelgruppen_liste()
            self.gruppe_name_var.set("")
            self.gruppe_artikel_listbox.delete(0, tk.END)
    
    def artikel_zu_gruppe_hinzufuegen(self):
        """Fügt den ausgewählten Artikel zur aktuellen Gruppe hinzu"""
        gruppe_name = self.gruppe_name_var.get()
        
        if not gruppe_name or gruppe_name not in self.artikelgruppen:
            messagebox.showinfo("Information", "Bitte wählen Sie eine gültige Artikelgruppe aus.")
            return
        
        selection = self.artikel_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie einen Artikel aus.")
            return
        
        index = selection[0]
        if 0 <= index < len(self.artikel_listbox_id_mapping):
            artikel_id = self.artikel_listbox_id_mapping[index]
            
            # Prüfen, ob Artikel bereits in der Gruppe ist
            if artikel_id in self.artikelgruppen[gruppe_name]:
                messagebox.showinfo("Information", "Dieser Artikel ist bereits in der Gruppe enthalten.")
                return
            
            # Artikel zur Gruppe hinzufügen
            self.artikelgruppen[gruppe_name].append(artikel_id)
            
            # Daten speichern
            self.speichere_daten()
            
            # Gruppenliste aktualisieren
            self.aktualisiere_gruppe_artikelliste()
    
    def artikel_aus_gruppe_entfernen(self):
        """Entfernt den ausgewählten Artikel aus der aktuellen Gruppe"""
        gruppe_name = self.gruppe_name_var.get()
        
        if not gruppe_name or gruppe_name not in self.artikelgruppen:
            messagebox.showinfo("Information", "Bitte wählen Sie eine gültige Artikelgruppe aus.")
            return
        
        selection = self.gruppe_artikel_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie einen Artikel aus der Gruppenliste aus.")
            return
        
        index = selection[0]
        
        if 0 <= index < len(self.artikelgruppen[gruppe_name]):
            # Artikel-ID aus Attribut holen
            artikel_id = self.gruppe_artikel_listbox.itemcget(index, 'artikel_id')
            
            if artikel_id:
                # Artikel aus Gruppe entfernen
                self.artikelgruppen[gruppe_name] = [
                    aid for aid in self.artikelgruppen[gruppe_name] if aid != artikel_id
                ]
                
                # Daten speichern
                self.speichere_daten()
                
                # Gruppenliste aktualisieren
                self.aktualisiere_gruppe_artikelliste()
    
    def lade_artikelgruppe(self):
        """Lädt eine Artikelgruppe in die aktuelle Rechnung"""
        gruppe_name = self.artikelgruppe_var.get()
        
        if not gruppe_name or gruppe_name not in self.artikelgruppen:
            messagebox.showinfo("Information", "Bitte wählen Sie eine gültige Artikelgruppe aus.")
            return
        
        # Artikel aus der Gruppe laden
        artikel_ids = self.artikelgruppen[gruppe_name]
        
        if not artikel_ids:
            messagebox.showinfo("Information", f"Die Artikelgruppe '{gruppe_name}' enthält keine Artikel.")
            return
        
        # Bestehende Artikel in der Rechnung durch die Gruppe ersetzen?
        if any(eintrag["bezeichnung"].get() for eintrag in self.artikel_eintraege):
            if not messagebox.askyesno("Bestätigung", 
                                      "Möchten Sie die bestehenden Artikel in der Rechnung ersetzen?"):
                return
        
        # Prüfen, ob genug Zeilen vorhanden sind, sonst hinzufügen
        while len(self.artikel_eintraege) < len(artikel_ids):
            self.artikel_zeile_hinzufuegen()
        
        # Alle Zeilen zunächst leeren
        for eintrag in self.artikel_eintraege:
            eintrag["bezeichnung"].set("")
            eintrag["menge"].set("1")
            eintrag["preis"].set("0.00")
            eintrag["ust"].set("19")
        
        # Artikel aus der Gruppe in die Tabelle übernehmen
        for i, artikel_id in enumerate(artikel_ids):
            if artikel_id in self.artikelbuch:
                artikel = self.artikelbuch[artikel_id]
                
                # In die i-te Zeile eintragen
                if i < len(self.artikel_eintraege):
                    self.artikel_eintraege[i]["bezeichnung"].set(artikel.get("bezeichnung", ""))
                    self.artikel_eintraege[i]["preis"].set(artikel.get("preis", "0.00"))
                    self.artikel_eintraege[i]["ust"].set(artikel.get("ust", "19"))
                    self.artikel_eintraege[i]["menge"].set("1")  # Standardmenge 1
        
        # Summen neuberechnen
        self.berechne_summen()
    
    def speichere_als_artikelgruppe(self):
        """Speichert die aktuellen Artikel als Artikelgruppe"""
        # Sammle alle nicht-leeren Artikel aus der aktuellen Rechnung
        artikel_eintraege = []
        
        for eintrag in self.artikel_eintraege:
            bezeichnung = eintrag["bezeichnung"].get()
            if bezeichnung:
                artikel_eintraege.append({
                    "bezeichnung": bezeichnung,
                    "preis": eintrag["preis"].get(),
                    "ust": eintrag["ust"].get()
                })
        
        if not artikel_eintraege:
            messagebox.showinfo("Information", "Es sind keine Artikel in der Rechnung vorhanden.")
            return
        
        # Gruppennamen abfragen
        gruppe_name = simpledialog.askstring("Artikelgruppe speichern", 
                                           "Bitte geben Sie den Namen der Artikelgruppe ein:")
        
        if not gruppe_name:
            return
        
        # Bestätigung bei existierender Gruppe
        if gruppe_name in self.artikelgruppen and self.artikelgruppen[gruppe_name]:
            if not messagebox.askyesno("Bestätigung", 
                                     f"Die Artikelgruppe '{gruppe_name}' existiert bereits. Überschreiben?"):
                return
        
        # Für jeden Artikel in der Rechnung:
        # 1. Prüfen, ob er bereits im Artikelbuch ist (nach Bezeichnung, Preis und USt)
        # 2. Falls nicht, neuen Artikel anlegen
        # 3. Artikel-ID zur Gruppe hinzufügen
        
        artikel_ids = []
        
        for eintrag in artikel_eintraege:
            bezeichnung = eintrag["bezeichnung"]
            preis = eintrag["preis"]
            ust = eintrag["ust"]
            
            # Nach passendem Artikel im Artikelbuch suchen
            gefunden = False
            for artikel_id, artikel in self.artikelbuch.items():
                if (artikel.get("bezeichnung") == bezeichnung and
                    artikel.get("preis") == preis and
                    artikel.get("ust") == ust):
                    
                    artikel_ids.append(artikel_id)
                    gefunden = True
                    break
            
            if not gefunden:
                # Neuen Artikel anlegen
                base_id = re.sub(r'[^\w]', '_', bezeichnung.lower())
                artikel_id = base_id
                
                # Falls ID bereits existiert, füge Nummer hinzu
                counter = 1
                while artikel_id in self.artikelbuch and counter < 100:
                    artikel_id = f"{base_id}_{counter}"
                    counter += 1
                
                # Artikel-Daten
                artikel = {
                    "Artikelnummer": "",
                    "Bezeichnung": bezeichnung,
                    "Beschreibung": "",
                    "Einheit": "",
                    "Netto-Preis": preis,
                    "USt-Satz": ust,
                    "Kategorie": "",
                    
                    # Felder für Verwendung in Rechnungen
                    "bezeichnung": bezeichnung,
                    "preis": preis,
                    "ust": ust
                }
                
                # Im Artikelbuch speichern
                self.artikelbuch[artikel_id] = artikel
                
                # Zur Gruppe hinzufügen
                artikel_ids.append(artikel_id)
        
        # Gruppe speichern
        self.artikelgruppen[gruppe_name] = artikel_ids
        
        # Daten speichern
        self.speichere_daten()
        
        # Listen aktualisieren
        self.aktualisiere_artikelliste()
        self.aktualisiere_artikelgruppen_liste()
        
        messagebox.showinfo("Erfolg", f"Artikelgruppe '{gruppe_name}' mit {len(artikel_ids)} Artikeln gespeichert!")
    
    def artikelbuch_importieren(self):
        """Importiert Artikel aus einer CSV-Datei"""
        filepath = filedialog.askopenfilename(
            title="Artikelbuch importieren",
            filetypes=[("CSV-Dateien", "*.csv"), ("Alle Dateien", "*.*")]
        )
        
        if not filepath:
            return
        
        try:
            with open(filepath, 'r', newline='', encoding='utf-8') as file:
                reader = csv.DictReader(file, delimiter=';')
                
                import_count = 0
                for row in reader:
                    # Artikelnummer oder Bezeichnung muss vorhanden sein
                    artikelnummer = row.get('Artikelnummer', '').strip()
                    bezeichnung = row.get('Bezeichnung', '').strip()
                    
                    if not bezeichnung:
                        continue  # Überspringe Zeilen ohne Bezeichnung
                    
                    # Artikel-ID erstellen
                    if artikelnummer:
                        artikel_id = re.sub(r'[^\w]', '_', artikelnummer.lower())
                    else:
                        base_id = re.sub(r'[^\w]', '_', bezeichnung.lower())
                        artikel_id = base_id
                        
                        # Falls ID bereits existiert, füge Nummer hinzu
                        counter = 1
                        while artikel_id in self.artikelbuch and counter < 100:
                            artikel_id = f"{base_id}_{counter}"
                            counter += 1
                    
                    # Felder für Verwendung in Rechnungen hinzufügen
                    row['bezeichnung'] = bezeichnung
                    row['preis'] = row.get('Netto-Preis', '0.00')
                    row['ust'] = row.get('USt-Satz', '19')
                    
                    # Artikel speichern
                    self.artikelbuch[artikel_id] = row
                    import_count += 1
                
                self.speichere_daten()
                self.aktualisiere_artikelliste()
                
                messagebox.showinfo("Import erfolgreich", f"{import_count} Artikel wurden importiert.")
        except Exception as e:
            messagebox.showerror("Fehler beim Import", f"Ein Fehler ist aufgetreten: {str(e)}")
    
    def artikelbuch_exportieren(self):
        """Exportiert Artikel in eine CSV-Datei"""
        if not self.artikelbuch:
            messagebox.showinfo("Information", "Das Artikelbuch ist leer.")
            return
        
        filepath = filedialog.asksaveasfilename(
            title="Artikelbuch exportieren",
            defaultextension=".csv",
            filetypes=[("CSV-Dateien", "*.csv"), ("Alle Dateien", "*.*")]
        )
        
        if not filepath:
            return
        
        try:
            # Alle möglichen Felder ermitteln
            all_fields = set()
            for artikel in self.artikelbuch.values():
                all_fields.update(artikel.keys())
            
            # Einige Felder ausschließen, die nur für die interne Verwendung sind
            exclude_fields = {'bezeichnung', 'preis', 'ust'}
            fieldnames = sorted([f for f in all_fields if f not in exclude_fields])
            
            with open(filepath, 'w', newline='', encoding='utf-8') as file:
                writer = csv.DictWriter(file, fieldnames=fieldnames, delimiter=';')
                writer.writeheader()
                
                # Alle Artikel schreiben
                for artikel_id, artikel in self.artikelbuch.items():
                    # Nur die gewünschten Felder exportieren
                    export_row = {f: artikel.get(f, "") for f in fieldnames}
                    writer.writerow(export_row)
            
            messagebox.showinfo("Export erfolgreich", f"Das Artikelbuch wurde nach {filepath} exportiert.")
        except Exception as e:
            messagebox.showerror("Fehler beim Export", f"Ein Fehler ist aufgetreten: {str(e)}")
    
    def init_turnusrechnungen_tab(self):
        """Initialisiert den Tab für die Turnusrechnungen"""
        # Hauptframe
        frame = ttk.Frame(self.tab_turnusrechnungen, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Linke Seite: Liste der Turnusrechnungen
        liste_frame = ttk.Frame(frame)
        liste_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        ttk.Label(liste_frame, text="Gespeicherte Turnusrechnungen:").pack(anchor="w")
        
        # Suchfeld
        search_frame = ttk.Frame(liste_frame)
        search_frame.pack(fill="x", pady=5)
        
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        self.turnus_search_var = tk.StringVar()
        self.turnus_search_var.trace("w", self.search_turnusrechnungen)
        ttk.Entry(search_frame, textvariable=self.turnus_search_var, width=20).pack(side="left", padx=5)
        
        # Status-Anzeige
        status_frame = ttk.Frame(liste_frame)
        status_frame.pack(fill="x", pady=5)
        
        self.scheduler_status_var = tk.StringVar(value="Scheduler: Inaktiv")
        ttk.Label(status_frame, textvariable=self.scheduler_status_var).pack(side="left")
        
        ttk.Button(status_frame, text="Scheduler starten", 
                  command=self.starte_scheduler).pack(side="left", padx=5)
        ttk.Button(status_frame, text="Scheduler stoppen", 
                  command=self.stoppe_scheduler).pack(side="left", padx=5)
        
        # Turnusrechnungsliste
        turnus_list_frame = ttk.Frame(liste_frame)
        turnus_list_frame.pack(fill="both", expand=True, pady=5)
        
        self.turnus_listbox = tk.Listbox(turnus_list_frame, height=15, width=40)
        turnus_scrollbar = ttk.Scrollbar(turnus_list_frame, orient="vertical", 
                                       command=self.turnus_listbox.yview)
        
        self.turnus_listbox.pack(side="left", fill="both", expand=True)
        turnus_scrollbar.pack(side="right", fill="y")
        self.turnus_listbox.config(yscrollcommand=turnus_scrollbar.set)
        
        self.turnus_listbox.bind('<<ListboxSelect>>', self.turnus_auswaehlen)
        
        button_frame = ttk.Frame(liste_frame)
        button_frame.pack(fill="x", pady=5)
        
        ttk.Button(button_frame, text="Neu", 
                  command=self.neue_turnusrechnung).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Löschen", 
                  command=self.loesche_turnusrechnung).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Jetzt ausführen", 
                  command=self.turnusrechnung_jetzt_ausfuehren).pack(side="left", padx=5)
        
        # Rechte Seite: Turnusrechnungsdetails
        details_frame = ttk.Frame(frame)
        details_frame.pack(side="right", fill="both", expand=True)
        
        # Grunddaten
        ttk.Label(details_frame, text="Bezeichnung:").grid(row=0, column=0, sticky="w", pady=5)
        self.turnus_bezeichnung_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.turnus_bezeichnung_var, width=40).grid(
            row=0, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Empfänger:").grid(row=1, column=0, sticky="w", pady=5)
        self.turnus_empfaenger_var = tk.StringVar()
        self.turnus_empfaenger_combo = ttk.Combobox(details_frame, 
                                                  textvariable=self.turnus_empfaenger_var, width=40)
        self.turnus_empfaenger_combo.grid(row=1, column=1, sticky="ew", pady=5, padx=5)
        self.aktualisiere_empfaenger_liste()  # Füllt die Combobox
        
        # Turnus-Einstellungen
        turnus_frame = ttk.LabelFrame(details_frame, text="Turnus-Einstellungen")
        turnus_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=10, padx=5)
        
        ttk.Label(turnus_frame, text="Turnus:").grid(row=0, column=0, sticky="w", pady=5)
        self.turnus_typ_var = tk.StringVar(value="monatlich")
        turnus_combo = ttk.Combobox(turnus_frame, textvariable=self.turnus_typ_var, width=15,
                                   values=["täglich", "wöchentlich", "monatlich", "vierteljährlich", "jährlich"])
        turnus_combo.grid(row=0, column=1, sticky="w", pady=5, padx=5)
        turnus_combo.bind("<<ComboboxSelected>>", self.update_turnus_optionen)
        
        # Tag/Datum-Auswahl
        self.turnus_tag_frame = ttk.Frame(turnus_frame)
        self.turnus_tag_frame.grid(row=1, column=0, columnspan=2, sticky="w", pady=5, padx=5)
        
        # Wochentag für wöchentlichen Turnus
        self.turnus_wochentag_var = tk.StringVar(value="Montag")
        self.turnus_wochentag_combo = ttk.Combobox(self.turnus_tag_frame, 
                                                 textvariable=self.turnus_wochentag_var, width=15,
                                                 values=["Montag", "Dienstag", "Mittwoch", "Donnerstag", 
                                                        "Freitag", "Samstag", "Sonntag"])
        
        # Tag im Monat für monatlichen Turnus
        self.turnus_monatstag_var = tk.StringVar(value="1")
        self.turnus_monatstag_combo = ttk.Combobox(self.turnus_tag_frame, 
                                                 textvariable=self.turnus_monatstag_var, width=15,
                                                 values=[str(i) for i in range(1, 32)])
        
        # Startzeit
        ttk.Label(turnus_frame, text="Uhrzeit:").grid(row=2, column=0, sticky="w", pady=5)
        
        zeit_frame = ttk.Frame(turnus_frame)
        zeit_frame.grid(row=2, column=1, sticky="w", pady=5, padx=5)
        
        self.turnus_stunde_var = tk.StringVar(value="9")
        self.turnus_minute_var = tk.StringVar(value="00")
        
        ttk.Combobox(zeit_frame, textvariable=self.turnus_stunde_var, width=5,
                    values=[f"{i:02d}" for i in range(24)]).pack(side="left")
        ttk.Label(zeit_frame, text=":").pack(side="left")
        ttk.Combobox(zeit_frame, textvariable=self.turnus_minute_var, width=5,
                    values=["00", "15", "30", "45"]).pack(side="left")
        
        # Artikel-Vorlage
        ttk.Label(details_frame, text="Artikelgruppe:").grid(row=3, column=0, sticky="w", pady=5)
        self.turnus_artikelgruppe_var = tk.StringVar()
        self.turnus_artikelgruppe_combo = ttk.Combobox(details_frame, 
                                                     textvariable=self.turnus_artikelgruppe_var, width=40)
        self.turnus_artikelgruppe_combo.grid(row=3, column=1, sticky="ew", pady=5, padx=5)
        self.aktualisiere_artikelgruppen_liste()  # Füllt die Combobox
        
        # E-Mail-Optionen
        email_frame = ttk.LabelFrame(details_frame, text="E-Mail-Versand")
        email_frame.grid(row=4, column=0, columnspan=2, sticky="ew", pady=10, padx=5)
        
        self.turnus_email_senden_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(email_frame, text="Rechnung per E-Mail versenden", 
                        variable=self.turnus_email_senden_var).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=5)
        
        ttk.Label(email_frame, text="Betreff:").grid(row=1, column=0, sticky="w", pady=5)
        self.turnus_email_betreff_var = tk.StringVar(
            value="Ihre Rechnung {rechnungsnr} vom {datum}")
        ttk.Entry(email_frame, textvariable=self.turnus_email_betreff_var, width=40).grid(
            row=1, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(email_frame, text="Text:").grid(row=2, column=0, sticky="nw", pady=5)
        self.turnus_email_text_var = tk.StringVar(
            value="Sehr geehrte Damen und Herren,\n\nhiermit übersenden wir Ihnen unsere Rechnung.\n\nMit freundlichen Grüßen")
        
        email_text_frame = ttk.Frame(email_frame)
        email_text_frame.grid(row=2, column=1, sticky="ew", pady=5, padx=5)
        
        email_text = tk.Text(email_text_frame, height=6, width=40)
        email_text.pack(side="left", fill="both", expand=True)
        email_scrollbar = ttk.Scrollbar(email_text_frame, orient="vertical", command=email_text.yview)
        email_scrollbar.pack(side="right", fill="y")
        email_text.config(yscrollcommand=email_scrollbar.set)
        
        # Text einfügen und an Variable binden
        email_text.insert("1.0", self.turnus_email_text_var.get())
        
        # Callback, um Text in Variable zu speichern
        def update_email_text(*args):
            self.turnus_email_text_var.set(email_text.get("1.0", "end-1c"))
        
        # Text-Widget mit Variable verknüpfen
        self.turnus_email_text = email_text
        email_text.bind("<KeyRelease>", update_email_text)
        
        # Speichern-Button
        ttk.Button(details_frame, text="Speichern", 
                  command=self.speichere_turnusrechnung).grid(
            row=5, column=0, columnspan=2, pady=10)
        
        # Initialisiere die Details mit den Default-Werten
        self.update_turnus_optionen()
        
        # Turnusrechnungen aktualisieren
        self.aktualisiere_turnusrechnungen_liste()
    
    def update_turnus_optionen(self, event=None):
        """Aktualisiert die angezeigten Optionen je nach gewähltem Turnus"""
        # Erst alle Widgets aus dem Frame entfernen
        for widget in self.turnus_tag_frame.winfo_children():
            widget.destroy()
        
        turnus_typ = self.turnus_typ_var.get()
        
        if turnus_typ == "täglich":
            # Keine speziellen Optionen für täglich
            ttk.Label(self.turnus_tag_frame, text="Jeden Tag").pack(side="left")
            
        elif turnus_typ == "wöchentlich":
            ttk.Label(self.turnus_tag_frame, text="Wochentag:").pack(side="left")
            self.turnus_wochentag_combo = ttk.Combobox(self.turnus_tag_frame, 
                                                    textvariable=self.turnus_wochentag_var, width=15,
                                                    values=["Montag", "Dienstag", "Mittwoch", "Donnerstag", 
                                                           "Freitag", "Samstag", "Sonntag"])
            self.turnus_wochentag_combo.pack(side="left", padx=5)
            
        elif turnus_typ == "monatlich":
            ttk.Label(self.turnus_tag_frame, text="Tag im Monat:").pack(side="left")
            self.turnus_monatstag_combo = ttk.Combobox(self.turnus_tag_frame, 
                                                    textvariable=self.turnus_monatstag_var, width=5,
                                                    values=[str(i) for i in range(1, 32)])
            self.turnus_monatstag_combo.pack(side="left", padx=5)
            
        elif turnus_typ == "vierteljährlich":
            ttk.Label(self.turnus_tag_frame, text="Quartale:").pack(side="left")
            
            self.turnus_quartal_vars = []
            quartale = ["Q1 (Jan-Mär)", "Q2 (Apr-Jun)", "Q3 (Jul-Sep)", "Q4 (Okt-Dez)"]
            
            for i, quartal in enumerate(quartale):
                var = tk.BooleanVar(value=True)
                self.turnus_quartal_vars.append(var)
                ttk.Checkbutton(self.turnus_tag_frame, text=quartal, variable=var).pack(side="left", padx=5)
            
            ttk.Label(self.turnus_tag_frame, text="Tag:").pack(side="left", padx=(10, 0))
            self.turnus_monatstag_combo = ttk.Combobox(self.turnus_tag_frame, 
                                                    textvariable=self.turnus_monatstag_var, width=5,
                                                    values=[str(i) for i in range(1, 32)])
            self.turnus_monatstag_combo.pack(side="left", padx=5)
            
        elif turnus_typ == "jährlich":
            ttk.Label(self.turnus_tag_frame, text="Monat:").pack(side="left")
            
            self.turnus_monat_var = tk.StringVar(value="1")
            monate = ["1 (Januar)", "2 (Februar)", "3 (März)", "4 (April)", "5 (Mai)", "6 (Juni)", 
                     "7 (Juli)", "8 (August)", "9 (September)", "10 (Oktober)", "11 (November)", "12 (Dezember)"]
            
            ttk.Combobox(self.turnus_tag_frame, textvariable=self.turnus_monat_var, width=15,
                        values=monate).pack(side="left", padx=5)
            
            ttk.Label(self.turnus_tag_frame, text="Tag:").pack(side="left", padx=(10, 0))
            self.turnus_monatstag_combo = ttk.Combobox(self.turnus_tag_frame, 
                                                    textvariable=self.turnus_monatstag_var, width=5,
                                                    values=[str(i) for i in range(1, 32)])
            self.turnus_monatstag_combo.pack(side="left", padx=5)
    
    def aktualisiere_turnusrechnungen_liste(self, filter_text=None):
        """Aktualisiert die Liste der Turnusrechnungen"""
        self.turnus_listbox.delete(0, tk.END)
        
        # Sortierte und formatierte Liste erstellen
        turnus_liste = []
        
        for name, turnus in self.turnusrechnungen.items():
            display_text = f"{name} - {turnus.get('turnus_typ', 'unbekannt')} - {turnus.get('empfaenger', '')}"
            
            # Bei Suchfilter überprüfen
            if filter_text is None or filter_text.lower() in display_text.lower():
                turnus_liste.append((name, display_text))
        
        # Nach Namen sortieren
        turnus_liste.sort()
        
        # Anzeigen
        for name, display_text in turnus_liste:
            self.turnus_listbox.insert(tk.END, display_text)
            # Name als Attribut speichern
            self.turnus_listbox.itemconfig(self.turnus_listbox.size()-1, {'name': name})
    
    def search_turnusrechnungen(self, *args):
        """Sucht in den Turnusrechnungen nach dem eingegebenen Suchbegriff"""
        search_text = self.turnus_search_var.get()
        self.aktualisiere_turnusrechnungen_liste(search_text)
    
    def turnus_auswaehlen(self, event):
        """Handler für die Auswahl einer Turnusrechnung aus der Liste"""
        selection = self.turnus_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        # Namen aus Attribut holen
        name = self.turnus_listbox.itemcget(index, 'name')
        
        if name and name in self.turnusrechnungen:
            turnus = self.turnusrechnungen[name]
            
            # Formular befüllen
            self.turnus_bezeichnung_var.set(name)
            self.turnus_empfaenger_var.set(turnus.get('empfaenger', ''))
            
            # Turnus-Einstellungen
            turnus_typ = turnus.get('turnus_typ', 'monatlich')
            self.turnus_typ_var.set(turnus_typ)
            
            # UI für Turnus-Typ aktualisieren
            self.update_turnus_optionen()
            
            # Werte für spezifischen Turnus-Typ setzen
            if turnus_typ == "wöchentlich":
                self.turnus_wochentag_var.set(turnus.get('wochentag', 'Montag'))
            elif turnus_typ in ["monatlich", "vierteljährlich", "jährlich"]:
                self.turnus_monatstag_var.set(turnus.get('monatstag', '1'))
                
                if turnus_typ == "vierteljährlich":
                    # Quartale aktivieren
                    quartale = turnus.get('quartale', [True, True, True, True])
                    for i, var in enumerate(self.turnus_quartal_vars):
                        if i < len(quartale):
                            var.set(quartale[i])
                
                if turnus_typ == "jährlich":
                    self.turnus_monat_var.set(turnus.get('monat', '1'))
            
            # Uhrzeit
            self.turnus_stunde_var.set(turnus.get('stunde', '9'))
            self.turnus_minute_var.set(turnus.get('minute', '00'))
            
            # Artikelgruppe
            self.turnus_artikelgruppe_var.set(turnus.get('artikelgruppe', ''))
            
            # E-Mail-Optionen
            self.turnus_email_senden_var.set(turnus.get('email_senden', True))
            self.turnus_email_betreff_var.set(turnus.get('email_betreff', 
                                               'Ihre Rechnung {rechnungsnr} vom {datum}'))
            
            email_text = turnus.get('email_text', 
                                   'Sehr geehrte Damen und Herren,\n\nhiermit übersenden wir Ihnen unsere Rechnung.\n\nMit freundlichen Grüßen')
            
            self.turnus_email_text.delete('1.0', tk.END)
            self.turnus_email_text.insert('1.0', email_text)
            self.turnus_email_text_var.set(email_text)
    
    def neue_turnusrechnung(self):
        """Leert das Formular für eine neue Turnusrechnung"""
        # Bezeichnung zurücksetzen
        self.turnus_bezeichnung_var.set("")
        
        # Empfänger leeren
        self.turnus_empfaenger_var.set("")
        
        # Turnus-Einstellungen auf Standard setzen
        self.turnus_typ_var.set("monatlich")
        self.update_turnus_optionen()
        
        # Monatstag auf 1 setzen
        self.turnus_monatstag_var.set("1")
        
        # Uhrzeit auf 9:00 Uhr setzen
        self.turnus_stunde_var.set("9")
        self.turnus_minute_var.set("00")
        
        # Artikelgruppe leeren
        self.turnus_artikelgruppe_var.set("")
        
        # E-Mail-Optionen auf Standard setzen
        self.turnus_email_senden_var.set(True)
        self.turnus_email_betreff_var.set("Ihre Rechnung {rechnungsnr} vom {datum}")
        
        email_text = "Sehr geehrte Damen und Herren,\n\nhiermit übersenden wir Ihnen unsere Rechnung.\n\nMit freundlichen Grüßen"
        self.turnus_email_text.delete('1.0', tk.END)
        self.turnus_email_text.insert('1.0', email_text)
        self.turnus_email_text_var.set(email_text)
    
    def loesche_turnusrechnung(self):
        """Löscht die ausgewählte Turnusrechnung"""
        selection = self.turnus_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie eine Turnusrechnung aus.")
            return
        
        index = selection[0]
        # Namen aus Attribut holen
        name = self.turnus_listbox.itemcget(index, 'name')
        
        if name and name in self.turnusrechnungen:
            if messagebox.askyesno("Löschen bestätigen", 
                                  f"Möchten Sie die Turnusrechnung '{name}' wirklich löschen?"):
                
                # Aus Turnusrechnungen entfernen
                del self.turnusrechnungen[name]
                
                # Daten speichern
                self.speichere_daten()
                
                # Liste aktualisieren
                self.aktualisiere_turnusrechnungen_liste()
                
                # Formular leeren
                self.neue_turnusrechnung()
    
    def speichere_turnusrechnung(self):
        """Speichert die aktuelle Turnusrechnung"""
        name = self.turnus_bezeichnung_var.get().strip()
        empfaenger = self.turnus_empfaenger_var.get()
        
        if not name:
            messagebox.showerror("Fehler", "Bitte geben Sie eine Bezeichnung ein!")
            return
        
        if not empfaenger or empfaenger not in self.adressbuch:
            messagebox.showerror("Fehler", "Bitte wählen Sie einen gültigen Empfänger!")
            return
        
        # Sammelt Daten für die Turnusrechnung
        turnus_daten = {
            'empfaenger': empfaenger,
            'turnus_typ': self.turnus_typ_var.get(),
            'stunde': self.turnus_stunde_var.get(),
            'minute': self.turnus_minute_var.get(),
            'artikelgruppe': self.turnus_artikelgruppe_var.get(),
            'email_senden': self.turnus_email_senden_var.get(),
            'email_betreff': self.turnus_email_betreff_var.get(),
            'email_text': self.turnus_email_text.get('1.0', 'end-1c')
        }
        
        # Spezifische Daten je nach Turnus-Typ
        turnus_typ = self.turnus_typ_var.get()
        
        if turnus_typ == "wöchentlich":
            turnus_daten['wochentag'] = self.turnus_wochentag_var.get()
            
        elif turnus_typ in ["monatlich", "vierteljährlich", "jährlich"]:
            turnus_daten['monatstag'] = self.turnus_monatstag_var.get()
            
            if turnus_typ == "vierteljährlich":
                # Quartale speichern
                turnus_daten['quartale'] = [var.get() for var in self.turnus_quartal_vars]
            
            if turnus_typ == "jährlich":
                turnus_daten['monat'] = self.turnus_monat_var.get()
        
        # Turnusrechnung speichern
        self.turnusrechnungen[name] = turnus_daten
        
        # Daten speichern
        self.speichere_daten()
        
        # Liste aktualisieren
        self.aktualisiere_turnusrechnungen_liste()
        
        messagebox.showinfo("Erfolg", f"Turnusrechnung '{name}' wurde gespeichert!")
    
    def turnusrechnung_jetzt_ausfuehren(self):
        """Führt die ausgewählte Turnusrechnung sofort aus"""
        selection = self.turnus_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie eine Turnusrechnung aus.")
            return
        
        index = selection[0]
        # Namen aus Attribut holen
        name = self.turnus_listbox.itemcget(index, 'name')
        
        if name and name in self.turnusrechnungen:
            if messagebox.askyesno("Bestätigung", 
                                  f"Möchten Sie die Turnusrechnung '{name}' jetzt ausführen?"):
                
                # Turnusrechnung ausführen
                ergebnis = self.fuehre_turnusrechnung_aus(name)
                
                if ergebnis:
                    messagebox.showinfo("Erfolg", 
                                      f"Turnusrechnung '{name}' wurde erfolgreich ausgeführt.\n"
                                      f"Rechnung {ergebnis} wurde erstellt.")
                else:
                    messagebox.showerror("Fehler", 
                                       f"Fehler beim Ausführen der Turnusrechnung '{name}'.")
    
    def fuehre_turnusrechnung_aus(self, name):
        """Führt eine Turnusrechnung aus und gibt die erstellte Rechnungsnummer zurück"""
        if name not in self.turnusrechnungen:
            return None
        
        turnus = self.turnusrechnungen[name]
        empfaenger_name = turnus.get('empfaenger')
        
        if not empfaenger_name or empfaenger_name not in self.adressbuch:
            return None
        
        # Wechsle zum Rechnungs-Tab, um dortige Funktionen zu nutzen
        self.notebook.select(self.tab_rechnung)
        
        # Empfänger setzen
        self.empfaenger_var.set(empfaenger_name)
        
        # Artikelgruppe laden, falls angegeben
        artikelgruppe = turnus.get('artikelgruppe')
        if artikelgruppe and artikelgruppe in self.artikelgruppen:
            self.artikelgruppe_var.set(artikelgruppe)
            self.lade_artikelgruppe()
        
        # Rechnungsinformationen setzen
        rechnungsnr = self.generiere_rechnungsnummer()
        self.rechnungsnr_var.set(rechnungsnr)
        
        rechnungsdatum = datetime.datetime.now().strftime("%d.%m.%Y")
        self.rechnungsdatum_var.set(rechnungsdatum)
        
        # Leistungszeitraum berechnen
        heute = datetime.datetime.now()
        
        # Für monatliche Rechnungen: Vormonat
        if turnus.get('turnus_typ') == "monatlich":
            # Erster und letzter Tag des Vormonats
            if heute.month == 1:  # Januar
                vormonat = datetime.date(heute.year - 1, 12, 1)
            else:
                vormonat = datetime.date(heute.year, heute.month - 1, 1)
            
            letzter_tag = calendar.monthrange(vormonat.year, vormonat.month)[1]
            ende_vormonat = datetime.date(vormonat.year, vormonat.month, letzter_tag)
            
            self.leistung_start_var.set(vormonat.strftime("%d.%m.%Y"))
            self.leistung_ende_var.set(ende_vormonat.strftime("%d.%m.%Y"))
        
        # Rechnung erstellen
        if turnus.get('email_senden', True):
            # Daten für E-Mail-Versand vorbereiten
            empfaenger = self.adressbuch[empfaenger_name]
            email = empfaenger.get("Email", "")
            
            # E-Mail-Betreff und -Text mit Platzhaltern ersetzen
            betreff_template = turnus.get('email_betreff', "Ihre Rechnung {rechnungsnr} vom {datum}")
            betreff = betreff_template.format(
                rechnungsnr=rechnungsnr,
                datum=rechnungsdatum
            )
            
            text_template = turnus.get('email_text', "Sehr geehrte Damen und Herren,\n\nhiermit übersenden wir Ihnen unsere Rechnung.\n\nMit freundlichen Grüßen")
            email_text = text_template.format(
                rechnungsnr=rechnungsnr,
                datum=rechnungsdatum
            )
            
            # Rechnung erstellen und per E-Mail versenden
            try:
                # Rechnung erstellen
                ergebnis = self.erstelle_rechnung(zum_versenden=True)
                
                if ergebnis:
                    docx_pfad = ergebnis["docx"]
                    
                    # Zusätzlich X-Rechnung erstellen, falls gewünscht
                    xrechnung_pfad = self.erstelle_xrechnung(ergebnis)
                    
                    # Mail-Client öffnen
                    try:
                        attachments = f'"{docx_pfad}"'
                        if xrechnung_pfad:
                            attachments += f' "{xrechnung_pfad}"'
                        
                        if sys.platform == "win32":
                            # Windows mit Outlook
                            os.system(f'start outlook.exe /c ipm.note /m "{email}" /a {attachments} /s /t "{betreff}" {email_text}')
                        else:
                            # Andere Betriebssysteme - Standard-Mail-Client
                            import webbrowser
                            mailto_link = f"mailto:{email}?subject={betreff}&body={email_text}"
                            webbrowser.open(mailto_link)
                            print(f"E-Mail-Client geöffnet. Bitte hängen Sie die Rechnung manuell an: {docx_pfad}")
                            if xrechnung_pfad:
                                print(f"Und die X-Rechnung: {xrechnung_pfad}")
                        
                        # Rechnung in Datenbank speichern
                        self.speichere_rechnung_in_datenbank(rechnungsnr, empfaenger_name)
                        
                        return rechnungsnr
                    except Exception as e:
                        print(f"Fehler beim Öffnen des E-Mail-Clients: {str(e)}")
                        return None
            except Exception as e:
                print(f"Fehler beim Erstellen der Turnusrechnung: {str(e)}")
                return None
        else:
            # Rechnung ohne E-Mail-Versand erstellen
            try:
                filepath = self.erstelle_rechnung()
                if filepath:
                    # Rechnung in Datenbank speichern
                    self.speichere_rechnung_in_datenbank(rechnungsnr, empfaenger_name)
                    return rechnungsnr
                return None
            except Exception as e:
                print(f"Fehler beim Erstellen der Turnusrechnung: {str(e)}")
                return None
    
    def speichere_rechnung_in_datenbank(self, rechnungsnr, empfaenger_name):
        """Speichert eine erstellte Rechnung in der Datenbank"""
        if not rechnungsnr or not empfaenger_name:
            return
        
        # Rechnungsdaten sammeln
        artikel = []
        for eintrag in self.artikel_eintraege:
            bezeichnung = eintrag["bezeichnung"].get()
            if bezeichnung:
                artikel.append({
                    "bezeichnung": bezeichnung,
                    "menge": eintrag["menge"].get(),
                    "preis": eintrag["preis"].get(),
                    "ust": eintrag["ust"].get()
                })
        
        rechnung = {
            "nummer": rechnungsnr,
            "datum": self.rechnungsdatum_var.get(),
            "empfaenger": empfaenger_name,
            "empfaenger_daten": self.adressbuch.get(empfaenger_name, {}),
            "leistungszeitraum_start": self.leistung_start_var.get(),
            "leistungszeitraum_ende": self.leistung_ende_var.get(),
            "artikel": artikel,
            "steuerschuldnerschaft": self.steuerschuldnerschaft_var.get(),
            "storno": self.storno_var.get(),
            "storno_rechnungsnr": self.storno_rechnungsnr_var.get(),
            "zahlungsziel": self.zahlungsziel_var.get(),
            "skonto_prozent": self.skonto_prozent_var.get(),
            "skonto_tage": self.skonto_tage_var.get(),
            "typ": self.rechnungstyp_var.get(),
            "bezahlt": False,
            "bezahlt_am": "",
            "mahnstufe": 0,
            "letzte_mahnung": ""
        }
        
        # In Datenbank speichern
        self.rechnungsdatenbank[rechnungsnr] = rechnung
        
        # Daten speichern
        self.speichere_daten()
    
    def starte_scheduler(self):
        """Startet den Scheduler für Turnusrechnungen"""
        if self.scheduler_running:
            messagebox.showinfo("Information", "Der Scheduler läuft bereits.")
            return
        
        # Scheduler starten
        self.scheduler_running = True
        self.scheduler_status_var.set("Scheduler: Aktiv")
        
        # Thread starten
        self.scheduler_thread = threading.Thread(target=self.scheduler_loop, daemon=True)
        self.scheduler_thread.start()
        
        messagebox.showinfo("Erfolg", "Scheduler wurde gestartet!")
    
    def stoppe_scheduler(self):
        """Stoppt den Scheduler für Turnusrechnungen"""
        if not self.scheduler_running:
            messagebox.showinfo("Information", "Der Scheduler ist bereits gestoppt.")
            return
        
        # Scheduler stoppen
        self.scheduler_running = False
        self.scheduler_status_var.set("Scheduler: Inaktiv")
        
        messagebox.showinfo("Erfolg", "Scheduler wurde gestoppt!")
    
    def scheduler_loop(self):
        """Haupt-Loop für den Scheduler zum Ausführen von Turnusrechnungen"""
        while self.scheduler_running:
            # Prüfe, ob Turnusrechnungen anstehen
            self.prüfe_turnusrechnungen()
            
            # Alle 5 Minuten prüfen
            time.sleep(300)
    
    def prüfe_turnusrechnungen(self):
        """Prüft, ob Turnusrechnungen ausgeführt werden müssen"""
        now = datetime.datetime.now()
        
        for name, turnus in self.turnusrechnungen.items():
            # Uhrzeit prüfen
            stunde = int(turnus.get('stunde', '9'))
            minute = int(turnus.get('minute', '0'))
            
            # Wenn die aktuelle Zeit innerhalb der letzten 5 Minuten liegt
            zeit_passt = (now.hour == stunde and 
                          now.minute >= minute and 
                          now.minute < minute + 5)
            
            if not zeit_passt:
                continue
            
            # Turnus-Typ prüfen
            turnus_typ = turnus.get('turnus_typ', 'monatlich')
            ausführen = False
            
            if turnus_typ == "täglich":
                # Täglich zur angegebenen Uhrzeit
                ausführen = True
                
            elif turnus_typ == "wöchentlich":
                # Wochentag prüfen
                wochentag = turnus.get('wochentag', 'Montag')
                wochentag_nr = {
                    "Montag": 0, "Dienstag": 1, "Mittwoch": 2, "Donnerstag": 3,
                    "Freitag": 4, "Samstag": 5, "Sonntag": 6
                }.get(wochentag, 0)
                
                ausführen = now.weekday() == wochentag_nr
                
            elif turnus_typ == "monatlich":
                # Tag im Monat prüfen
                monatstag = int(turnus.get('monatstag', '1'))
                ausführen = now.day == monatstag
                
            elif turnus_typ == "vierteljährlich":
                # Tag im Monat prüfen
                monatstag = int(turnus.get('monatstag', '1'))
                
                # Monat prüfen (Q1=1,2,3; Q2=4,5,6; Q3=7,8,9; Q4=10,11,12)
                quartale = turnus.get('quartale', [True, True, True, True])
                
                q1 = quartale[0] if len(quartale) > 0 else True
                q2 = quartale[1] if len(quartale) > 1 else True
                q3 = quartale[2] if len(quartale) > 2 else True
                q4 = quartale[3] if len(quartale) > 3 else True
                
                if ((q1 and now.month in [1, 4, 7, 10]) or
                    (q2 and now.month in [2, 5, 8, 11]) or
                    (q3 and now.month in [3, 6, 9, 12])):
                    ausführen = now.day == monatstag
                
            elif turnus_typ == "jährlich":
                # Tag und Monat prüfen
                monatstag = int(turnus.get('monatstag', '1'))
                monat = int(turnus.get('monat', '1'))
                
                ausführen = now.day == monatstag and now.month == monat
            
            # Turnusrechnung ausführen, wenn Bedingungen erfüllt sind
            if ausführen:
                try:
                    self.fuehre_turnusrechnung_aus(name)
                    print(f"Turnusrechnung '{name}' wurde ausgeführt.")
                except Exception as e:
                    print(f"Fehler beim Ausführen der Turnusrechnung '{name}': {str(e)}")
    
    def init_mahnungen_tab(self):
        """Initialisiert den Tab für Mahnungen"""
        frame = ttk.Frame(self.tab_mahnungen, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Linke Seite: Liste der offenen Rechnungen
        liste_frame = ttk.LabelFrame(frame, text="Offene Rechnungen")
        liste_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        # Suchfeld
        search_frame = ttk.Frame(liste_frame)
        search_frame.pack(fill="x", pady=5)
        
        ttk.Label(search_frame, text="Suche:").pack(side="left")
        self.mahnung_search_var = tk.StringVar()
        self.mahnung_search_var.trace("w", self.search_offene_rechnungen)
        ttk.Entry(search_frame, textvariable=self.mahnung_search_var, width=20).pack(side="left", padx=5)
        
        # Status-Filter
        ttk.Label(search_frame, text="Status:").pack(side="left", padx=(10, 0))
        self.mahnung_status_var = tk.StringVar(value="Alle")
        ttk.Combobox(search_frame, textvariable=self.mahnung_status_var, width=15,
                    values=["Alle", "Fällig", "Überfällig", "Gemahnt"]).pack(side="left", padx=5)
        self.mahnung_status_var.trace("w", self.search_offene_rechnungen)
        
        # Rechnungsliste
        rechnungs_list_frame = ttk.Frame(liste_frame)
        rechnungs_list_frame.pack(fill="both", expand=True, pady=5)
        
        self.offene_rechnungen_listbox = tk.Listbox(rechnungs_list_frame, height=15, width=50)
        rechnung_scrollbar = ttk.Scrollbar(rechnungs_list_frame, orient="vertical", 
                                          command=self.offene_rechnungen_listbox.yview)
        
        self.offene_rechnungen_listbox.pack(side="left", fill="both", expand=True)
        rechnung_scrollbar.pack(side="right", fill="y")
        self.offene_rechnungen_listbox.config(yscrollcommand=rechnung_scrollbar.set)
        
        self.offene_rechnungen_listbox.bind('<<ListboxSelect>>', self.rechnung_fuer_mahnung_auswaehlen)
        
        # Buttons für Mahnlauf
        button_frame = ttk.Frame(liste_frame)
        button_frame.pack(fill="x", pady=5)
        
        ttk.Button(button_frame, text="Rechnung aktualisieren", 
                  command=self.aktualisiere_offene_rechnungen).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Als bezahlt markieren", 
                  command=self.rechnung_als_bezahlt_markieren).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Automatischer Mahnlauf", 
                  command=self.automatischer_mahnlauf).pack(side="left", padx=5)
        
        # Rechte Seite: Mahnungsdetails
        details_frame = ttk.LabelFrame(frame, text="Mahnungsdetails")
        details_frame.pack(side="right", fill="both", expand=True)
        
        # Ausgewählte Rechnung
        ttk.Label(details_frame, text="Rechnungsnummer:").grid(row=0, column=0, sticky="w", pady=5)
        self.mahnung_rechnungsnr_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_rechnungsnr_var, width=20, state="readonly").grid(
            row=0, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Rechnungsdatum:").grid(row=1, column=0, sticky="w", pady=5)
        self.mahnung_rechnungsdatum_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_rechnungsdatum_var, width=20, state="readonly").grid(
            row=1, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Zahlungsziel:").grid(row=2, column=0, sticky="w", pady=5)
        self.mahnung_zahlungsziel_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_zahlungsziel_var, width=20, state="readonly").grid(
            row=2, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Betrag:").grid(row=3, column=0, sticky="w", pady=5)
        self.mahnung_betrag_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_betrag_var, width=20, state="readonly").grid(
            row=3, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Empfänger:").grid(row=4, column=0, sticky="w", pady=5)
        self.mahnung_empfaenger_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_empfaenger_var, width=40, state="readonly").grid(
            row=4, column=1, sticky="ew", pady=5, padx=5)
        
        # Mahnungsstatus
        ttk.Label(details_frame, text="Mahnstufe:").grid(row=5, column=0, sticky="w", pady=5)
        self.mahnung_stufe_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_stufe_var, width=20, state="readonly").grid(
            row=5, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Letzte Mahnung:").grid(row=6, column=0, sticky="w", pady=5)
        self.mahnung_letzte_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.mahnung_letzte_var, width=20, state="readonly").grid(
            row=6, column=1, sticky="w", pady=5, padx=5)
        
        # Mahnung erstellen
        mahnung_frame = ttk.LabelFrame(details_frame, text="Mahnung erstellen")
        mahnung_frame.grid(row=7, column=0, columnspan=2, sticky="ew", pady=10, padx=5)
        
        ttk.Label(mahnung_frame, text="Mahnstufe:").grid(row=0, column=0, sticky="w", pady=5)
        self.mahnung_neue_stufe_var = tk.StringVar(value="1")
        ttk.Combobox(mahnung_frame, textvariable=self.mahnung_neue_stufe_var, width=15,
                    values=["1", "2", "3", "Letzte Mahnung"]).grid(
            row=0, column=1, sticky="w", pady=5, padx=5)
        
        self.mahnung_gebuehr_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(mahnung_frame, text="Mahngebühr erheben", 
                       variable=self.mahnung_gebuehr_var).grid(
            row=1, column=0, columnspan=2, sticky="w", pady=5)
        
        ttk.Label(mahnung_frame, text="Mahngebühr:").grid(row=2, column=0, sticky="w", pady=5)
        self.mahnung_gebuehr_betrag_var = tk.StringVar(value="5.00")
        ttk.Entry(mahnung_frame, textvariable=self.mahnung_gebuehr_betrag_var, width=10).grid(
            row=2, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(mahnung_frame, text="Zahlungsfrist (Tage):").grid(row=3, column=0, sticky="w", pady=5)
        self.mahnung_frist_var = tk.StringVar(value="7")
        ttk.Entry(mahnung_frame, textvariable=self.mahnung_frist_var, width=10).grid(
            row=3, column=1, sticky="w", pady=5, padx=5)
        
        # E-Mail-Optionen
        email_frame = ttk.LabelFrame(mahnung_frame, text="E-Mail-Versand")
        email_frame.grid(row=4, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        
        self.mahnung_email_senden_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(email_frame, text="Mahnung per E-Mail versenden", 
                       variable=self.mahnung_email_senden_var).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=5)
        
        ttk.Label(email_frame, text="Betreff:").grid(row=1, column=0, sticky="w", pady=5)
        self.mahnung_email_betreff_var = tk.StringVar(
            value="Mahnung zur Rechnung {rechnungsnr}")
        ttk.Entry(email_frame, textvariable=self.mahnung_email_betreff_var, width=40).grid(
            row=1, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(email_frame, text="Text:").grid(row=2, column=0, sticky="nw", pady=5)
        
        email_text_frame = ttk.Frame(email_frame)
        email_text_frame.grid(row=2, column=1, sticky="ew", pady=5, padx=5)
        
        self.mahnung_email_text = tk.Text(email_text_frame, height=6, width=40)
        self.mahnung_email_text.pack(side="left", fill="both", expand=True)
        email_scrollbar = ttk.Scrollbar(email_text_frame, orient="vertical", 
                                      command=self.mahnung_email_text.yview)
        email_scrollbar.pack(side="right", fill="y")
        self.mahnung_email_text.config(yscrollcommand=email_scrollbar.set)
        
        # Standard-Text für die 1. Mahnung
        standard_text = (
            "Sehr geehrte Damen und Herren,\n\n"
            "leider konnten wir bislang keinen Zahlungseingang zu unserer Rechnung {rechnungsnr} vom {datum} "
            "feststellen. Wir bitten Sie, den Rechnungsbetrag in Höhe von {betrag} € "
            "innerhalb der nächsten {frist} Tage zu begleichen.\n\n"
            "Mit freundlichen Grüßen"
        )
        
        self.mahnung_email_text.insert("1.0", standard_text)
        
        # Buttons
        button_frame = ttk.Frame(details_frame)
        button_frame.grid(row=8, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="Mahnung erstellen", 
                  command=self.erstelle_mahnung).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Mahnvorlagen verwalten", 
                  command=self.mahnvorlagen_verwalten).pack(side="left", padx=5)
        
        # Liste der offenen Rechnungen aktualisieren
        self.aktualisiere_offene_rechnungen()
    
    def aktualisiere_offene_rechnungen(self, filter_text=None):
        """Aktualisiert die Liste der offenen Rechnungen"""
        self.offene_rechnungen_listbox.delete(0, tk.END)
        
        # Aktuelles Datum für Fälligkeitsberechnung
        heute = datetime.datetime.now().date()
        
        # Status-Filter
        status_filter = self.mahnung_status_var.get()
        
        # Sortierte und formatierte Liste erstellen
        rechnungs_liste = []
        
        for rechnungsnr, rechnung in self.rechnungsdatenbank.items():
            # Nur offene Rechnungen anzeigen
            if rechnung.get("bezahlt", False):
                continue
            
            # Rechnungsdaten extrahieren
            datum = rechnung.get("datum", "")
            empfaenger = rechnung.get("empfaenger", "")
            zahlungsziel = int(rechnung.get("zahlungsziel", "14"))
            
            # Fälligkeitsdatum berechnen
            try:
                rechnungsdatum = datetime.datetime.strptime(datum, "%d.%m.%Y").date()
                faelligkeit = rechnungsdatum + datetime.timedelta(days=zahlungsziel)
                tage_ueberfaellig = (heute - faelligkeit).days if heute > faelligkeit else 0
            except:
                # Wenn Datum nicht geparst werden kann
                faelligkeit = None
                tage_ueberfaellig = 0
            
            # Mahnstatus
            mahnstufe = rechnung.get("mahnstufe", 0)
            status = ""
            
            if mahnstufe > 0:
                status = f"Gemahnt ({mahnstufe})"
            elif tage_ueberfaellig > 0:
                status = f"Überfällig ({tage_ueberfaellig} Tage)"
            elif faelligkeit:
                tage_bis_faellig = (faelligkeit - heute).days
                if tage_bis_faellig <= 7:
                    status = f"Fällig in {tage_bis_faellig} Tagen"
                else:
                    status = "Offen"
            else:
                status = "Offen"
            
            # Status-Filter anwenden
            if status_filter != "Alle":
                if status_filter == "Fällig" and "Fällig" not in status:
                    continue
                elif status_filter == "Überfällig" and "Überfällig" not in status:
                    continue
                elif status_filter == "Gemahnt" and "Gemahnt" not in status:
                    continue
            
            # Display-Text
            display_text = f"{rechnungsnr} - {datum} - {empfaenger} - {status}"
            
            # Bei Suchfilter überprüfen
            if filter_text is None or filter_text.lower() in display_text.lower():
                rechnungs_liste.append((rechnungsnr, datum, empfaenger, status, display_text))
        
        # Nach Datum sortieren
        rechnungs_liste.sort(key=lambda x: x[1], reverse=True)
        
        # Anzeigen
        for rechnungsnr, _, _, _, display_text in rechnungs_liste:
            self.offene_rechnungen_listbox.insert(tk.END, display_text)
            # Rechnungsnummer als Attribut speichern
            self.offene_rechnungen_listbox.itemconfig(
                self.offene_rechnungen_listbox.size()-1, {'rechnungsnr': rechnungsnr})
    
    def search_offene_rechnungen(self, *args):
        """Sucht in den offenen Rechnungen nach dem eingegebenen Suchbegriff"""
        search_text = self.mahnung_search_var.get()
        self.aktualisiere_offene_rechnungen(search_text)
    
    def rechnung_fuer_mahnung_auswaehlen(self, event):
        """Handler für die Auswahl einer Rechnung aus der Liste"""
        selection = self.offene_rechnungen_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        # Rechnungsnummer aus Attribut holen
        rechnungsnr = self.offene_rechnungen_listbox.itemcget(index, 'rechnungsnr')
        
        if rechnungsnr and rechnungsnr in self.rechnungsdatenbank:
            rechnung = self.rechnungsdatenbank[rechnungsnr]
            
            # Mahnung-Details befüllen
            self.mahnung_rechnungsnr_var.set(rechnungsnr)
            self.mahnung_rechnungsdatum_var.set(rechnung.get("datum", ""))
            
            # Zahlungsziel berechnen
            zahlungsziel = rechnung.get("zahlungsziel", "14")
            datum = rechnung.get("datum", "")
            
            try:
                rechnungsdatum = datetime.datetime.strptime(datum, "%d.%m.%Y").date()
                faelligkeit = rechnungsdatum + datetime.timedelta(days=int(zahlungsziel))
                self.mahnung_zahlungsziel_var.set(faelligkeit.strftime("%d.%m.%Y"))
            except:
                self.mahnung_zahlungsziel_var.set(f"{zahlungsziel} Tage nach Rechnungsdatum")
            
            # Betrag berechnen
            betrag = self.berechne_rechnungsbetrag(rechnung)
            self.mahnung_betrag_var.set(f"{betrag:.2f} €")
            
            # Empfänger
            self.mahnung_empfaenger_var.set(rechnung.get("empfaenger", ""))
            
            # Mahnstatus
            mahnstufe = rechnung.get("mahnstufe", 0)
            self.mahnung_stufe_var.set(str(mahnstufe))
            
            letzte_mahnung = rechnung.get("letzte_mahnung", "")
            self.mahnung_letzte_var.set(letzte_mahnung)
            
            # Vorschlag für nächste Mahnstufe
            naechste_stufe = mahnstufe + 1
            if naechste_stufe > 3:
                self.mahnung_neue_stufe_var.set("Letzte Mahnung")
            else:
                self.mahnung_neue_stufe_var.set(str(naechste_stufe))
            
            # Text für die Mahnung aktualisieren
            self.aktualisiere_mahnung_text(rechnung)
    
    def berechne_rechnungsbetrag(self, rechnung):
        """Berechnet den Gesamtbetrag einer Rechnung"""
        artikel = rechnung.get("artikel", [])
        
        nettosumme = Decimal('0.00')
        ust_betraege = {}
        
        for artikel_item in artikel:
            try:
                menge = Decimal(artikel_item.get("menge", "1").replace(',', '.'))
                preis = Decimal(artikel_item.get("preis", "0.00").replace(',', '.'))
                ust_satz = Decimal(artikel_item.get("ust", "19").replace(',', '.'))
                
                position_netto = menge * preis
                ust_betrag = position_netto * (ust_satz / 100)
                
                nettosumme += position_netto
                
                ust_key = str(ust_satz)
                if ust_key not in ust_betraege:
                    ust_betraege[ust_key] = Decimal('0.00')
                ust_betraege[ust_key] += ust_betrag
                
            except:
                continue
        
        # Gesamtbeträge berechnen
        gesamt_ust = sum(ust_betraege.values())
        bruttosumme = nettosumme + gesamt_ust
        
        return bruttosumme
    
    def aktualisiere_mahnung_text(self, rechnung):
        """Aktualisiert den Text für die Mahnung basierend auf der Mahnstufe"""
        rechnungsnr = rechnung.get("nummer", "")
        datum = rechnung.get("datum", "")
        betrag = self.berechne_rechnungsbetrag(rechnung)
        frist = self.mahnung_frist_var.get()
        
        # Mahnstufe aus UI
        mahnstufe = self.mahnung_neue_stufe_var.get()
        
        # Standard-Texte je nach Mahnstufe
        if mahnstufe == "1":
            standard_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "leider konnten wir bislang keinen Zahlungseingang zu unserer Rechnung {rechnungsnr} vom {datum} "
                "feststellen. Wir bitten Sie, den Rechnungsbetrag in Höhe von {betrag} € "
                "innerhalb der nächsten {frist} Tage zu begleichen.\n\n"
                "Sollten Sie die Zahlung bereits veranlasst haben, betrachten Sie dieses Schreiben bitte als gegenstandslos.\n\n"
                "Mit freundlichen Grüßen"
            )
        elif mahnstufe == "2":
            standard_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "wir müssen leider feststellen, dass Sie unsere Rechnung {rechnungsnr} vom {datum} "
                "trotz unserer ersten Mahnung noch nicht beglichen haben. "
                "Wir fordern Sie auf, den Rechnungsbetrag in Höhe von {betrag} € unverzüglich, "
                "spätestens jedoch innerhalb der nächsten {frist} Tage zu begleichen.\n\n"
                "Für die entstandenen Umstände berechnen wir eine Mahngebühr.\n\n"
                "Mit freundlichen Grüßen"
            )
        elif mahnstufe == "3":
            standard_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "trotz wiederholter Mahnungen haben Sie unsere Rechnung {rechnungsnr} vom {datum} "
                "noch immer nicht beglichen. Dies ist nun unsere dritte Mahnung. "
                "Wir fordern Sie letztmalig auf, den Rechnungsbetrag in Höhe von {betrag} € zuzüglich "
                "der Mahngebühren innerhalb der nächsten {frist} Tage zu begleichen.\n\n"
                "Bei weiterer Zahlungsverzögerung werden wir rechtliche Schritte einleiten müssen.\n\n"
                "Mit freundlichen Grüßen"
            )
        elif mahnstufe == "Letzte Mahnung":
            standard_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "hiermit mahnen wir Sie letztmalig an, den ausstehenden Betrag für unsere Rechnung "
                "{rechnungsnr} vom {datum} in Höhe von {betrag} € zuzüglich der angefallenen Mahngebühren "
                "innerhalb der nächsten {frist} Tage zu begleichen.\n\n"
                "Sollte der Betrag nicht fristgerecht bei uns eingehen, werden wir ohne weitere "
                "Ankündigung rechtliche Schritte einleiten und den Vorgang an unser Inkassounternehmen übergeben.\n\n"
                "Mit freundlichen Grüßen"
            )
        else:
            standard_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "bitte begleichen Sie die offene Rechnung {rechnungsnr} vom {datum} "
                "in Höhe von {betrag} € innerhalb der nächsten {frist} Tage.\n\n"
                "Mit freundlichen Grüßen"
            )
        
        # Platzhalter ersetzen
        text = standard_text.format(
            rechnungsnr=rechnungsnr,
            datum=datum,
            betrag=f"{betrag:.2f}",
            frist=frist
        )
        
        # Text in das Textfeld einfügen
        self.mahnung_email_text.delete("1.0", tk.END)
        self.mahnung_email_text.insert("1.0", text)
    
    def rechnung_als_bezahlt_markieren(self):
        """Markiert die ausgewählte Rechnung als bezahlt"""
        selection = self.offene_rechnungen_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie eine Rechnung aus.")
            return
        
        index = selection[0]
        # Rechnungsnummer aus Attribut holen
        rechnungsnr = self.offene_rechnungen_listbox.itemcget(index, 'rechnungsnr')
        
        if rechnungsnr and rechnungsnr in self.rechnungsdatenbank:
            if messagebox.askyesno("Bestätigung", 
                                  f"Möchten Sie die Rechnung {rechnungsnr} als bezahlt markieren?"):
                
                # Rechnung als bezahlt markieren
                self.rechnungsdatenbank[rechnungsnr]["bezahlt"] = True
                self.rechnungsdatenbank[rechnungsnr]["bezahlt_am"] = datetime.datetime.now().strftime("%d.%m.%Y")
                
                # Daten speichern
                self.speichere_daten()
                
                # Liste aktualisieren
                self.aktualisiere_offene_rechnungen()
                
                # Formularfelder leeren
                self.mahnung_rechnungsnr_var.set("")
                self.mahnung_rechnungsdatum_var.set("")
                self.mahnung_zahlungsziel_var.set("")
                self.mahnung_betrag_var.set("")
                self.mahnung_empfaenger_var.set("")
                self.mahnung_stufe_var.set("")
                self.mahnung_letzte_var.set("")
    
    def erstelle_mahnung(self):
        """Erstellt eine Mahnung für die ausgewählte Rechnung"""
        rechnungsnr = self.mahnung_rechnungsnr_var.get()
        
        if not rechnungsnr or rechnungsnr not in self.rechnungsdatenbank:
            messagebox.showinfo("Information", "Bitte wählen Sie eine gültige Rechnung aus.")
            return
        
        rechnung = self.rechnungsdatenbank[rechnungsnr]
        
        # Mahnstufe
        mahnstufe_text = self.mahnung_neue_stufe_var.get()
        if mahnstufe_text == "Letzte Mahnung":
            mahnstufe = 4
        else:
            try:
                mahnstufe = int(mahnstufe_text)
            except:
                mahnstufe = 1
        
        # Prüfen, ob die Mahnstufe erhöht wird
        alte_stufe = rechnung.get("mahnstufe", 0)
        if mahnstufe <= alte_stufe:
            if not messagebox.askyesno("Bestätigung", 
                                     f"Die Mahnstufe wird nicht erhöht. Trotzdem eine neue Mahnung erstellen?"):
                return
        
        # Mahngebühr
        mahngebuehr = Decimal('0.00')
        if self.mahnung_gebuehr_var.get():
            try:
                mahngebuehr = Decimal(self.mahnung_gebuehr_betrag_var.get().replace(',', '.'))
            except:
                mahngebuehr = Decimal('5.00')
        
        # Mahnung erstellen
        try:
            # Prüfen, ob ein Briefkopf existiert
            if not self.firmendaten.get("briefkopf_pfad") or not os.path.exists(self.firmendaten["briefkopf_pfad"]):
                messagebox.showerror("Fehler", "Bitte erstellen oder laden Sie zuerst einen Briefkopf!")
                return None
            
            # Empfänger-Daten
            empfaenger_name = rechnung.get("empfaenger", "")
            if not empfaenger_name or empfaenger_name not in self.adressbuch:
                messagebox.showerror("Fehler", "Der Empfänger wurde nicht im Adressbuch gefunden!")
                return None
            
            empfaenger = self.adressbuch[empfaenger_name]
            
            # Mahndokument erstellen
            # Vorlage kopieren
            doc = docx.Document(self.firmendaten["briefkopf_pfad"])
            
            # Empfängerdaten eintragen
            for i, para in enumerate(doc.paragraphs):
                if "Empfänger:" in para.text:
                    # Die nächsten Paragraphen mit Empfängerdaten überschreiben
                    empfaenger_zeilen = []
                    
                    # Name oder Firma
                    if empfaenger.get("Firma"):
                        empfaenger_zeilen.append(empfaenger["Firma"])
                        if empfaenger.get("Vorname") and empfaenger.get("Nachname"):
                            empfaenger_zeilen.append(f"{empfaenger['Vorname']} {empfaenger['Nachname']}")
                        elif empfaenger.get("Vorname"):
                            empfaenger_zeilen.append(empfaenger["Vorname"])
                        elif empfaenger.get("Nachname"):
                            empfaenger_zeilen.append(empfaenger["Nachname"])
                    else:
                        name_str = ""
                        if empfaenger.get("Vorname"):
                            name_str += empfaenger["Vorname"] + " "
                        if empfaenger.get("Nachname"):
                            name_str += empfaenger["Nachname"]
                        empfaenger_zeilen.append(name_str.strip())
                    
                    # Adresse
                    empfaenger_zeilen.append(empfaenger.get("Straße", ""))
                    empfaenger_zeilen.append(f"{empfaenger.get('PLZ', '')} {empfaenger.get('Ort', '')}")
                    
                    # In Dokument schreiben
                    for j, zeile in enumerate(empfaenger_zeilen):
                        if i+j+1 < len(doc.paragraphs):
                            doc.paragraphs[i+j+1].text = zeile
                    break
            
            # Text "RECHNUNG" durch "MAHNUNG" ersetzen
            for para in doc.paragraphs:
                if "RECHNUNG" in para.text:
                    if mahnstufe == 1:
                        para.text = "MAHNUNG"
                    elif mahnstufe == 2:
                        para.text = "ZWEITE MAHNUNG"
                    elif mahnstufe == 3:
                        para.text = "DRITTE MAHNUNG"
                    elif mahnstufe >= 4:
                        para.text = "LETZTE MAHNUNG"
                    else:
                        para.text = "MAHNUNG"
                    
                    para.runs[0].bold = True
                    break
            
            # Infos zur Mahnung hinzufügen
            p = doc.add_paragraph()
            p.text = f"Bezug: Rechnung {rechnungsnr} vom {rechnung.get('datum', '')}"
            
            # Mahnungstext aus Textfeld
            mahnungstext = self.mahnung_email_text.get("1.0", "end-1c")
            p = doc.add_paragraph()
            p.text = mahnungstext
            
            # Zahlungsfrist
            frist = int(self.mahnung_frist_var.get())
            frist_datum = datetime.datetime.now() + datetime.timedelta(days=frist)
            
            p = doc.add_paragraph()
            p.text = f"Bitte begleichen Sie den offenen Betrag bis spätestens {frist_datum.strftime('%d.%m.%Y')}."
            
            # Mahngebühr
            if mahngebuehr > 0:
                p = doc.add_paragraph()
                p.text = f"Für diese Mahnung berechnen wir eine Mahngebühr in Höhe von {mahngebuehr:.2f} €."
            
            # Betrag
            betrag = self.berechne_rechnungsbetrag(rechnung)
            gesamt = betrag + mahngebuehr
            
            # Summentabelle
            sum_table = doc.add_table(rows=3, cols=2)
            sum_table.autofit = True
            
            row = sum_table.rows[0].cells
            row[0].text = "Offener Rechnungsbetrag:"
            row[1].text = f"{betrag:.2f} €"
            
            row = sum_table.rows[1].cells
            row[0].text = "Mahngebühr:"
            row[1].text = f"{mahngebuehr:.2f} €"
            
            row = sum_table.rows[2].cells
            row[0].text = "Gesamtbetrag:"
            row[1].text = f"{gesamt:.2f} €"
            
            # Mahnungsdatum
            mahnung_datum = datetime.datetime.now().strftime("%d.%m.%Y")
            
            # Speichern
            if self.mahnung_email_senden_var.get():
                # Temporäres Speichern für E-Mail-Versand
                temp_docx = os.path.join(os.environ.get("TEMP", "."), f"Mahnung_{rechnungsnr}_{mahnstufe}.docx")
                doc.save(temp_docx)
                
                # E-Mail-Versand
                email = empfaenger.get("Email", "")
                
                # E-Mail-Betreff
                betreff_template = self.mahnung_email_betreff_var.get()
                betreff = betreff_template.format(
                    rechnungsnr=rechnungsnr,
                    datum=rechnung.get("datum", ""),
                    betrag=f"{betrag:.2f}",
                    frist=frist
                )
                
                # Mail-Client öffnen
                try:
                    if sys.platform == "win32":
                        # Windows mit Outlook
                        os.system(f'start outlook.exe /c ipm.note /m "{email}" /a "{temp_docx}" /s /t "{betreff}"')
                    else:
                        # Andere Betriebssysteme - Standard-Mail-Client
                        import webbrowser
                        mailto_link = f"mailto:{email}?subject={betreff}"
                        webbrowser.open(mailto_link)
                        messagebox.showinfo("Hinweis", 
                                           f"E-Mail-Client wurde geöffnet. Bitte hängen Sie die Mahnung manuell an:\n{temp_docx}")
                
                except Exception as e:
                    messagebox.showerror("Fehler", f"Fehler beim Öffnen des E-Mail-Clients: {str(e)}")
            else:
                # Speichern als Datei
                filepath = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Dokument", "*.docx")],
                    initialfile=f"Mahnung_{rechnungsnr}_{mahnstufe}.docx",
                    title="Mahnung speichern"
                )
                
                if filepath:
                    doc.save(filepath)
            
            # Rechnung aktualisieren
            self.rechnungsdatenbank[rechnungsnr]["mahnstufe"] = mahnstufe
            self.rechnungsdatenbank[rechnungsnr]["letzte_mahnung"] = mahnung_datum
            
            # Mahngebühr zur Rechnung hinzufügen (optional)
            if mahngebuehr > 0:
                artikel = rechnung.get("artikel", [])
                artikel.append({
                    "bezeichnung": f"Mahngebühr ({mahnstufe}. Mahnung)",
                    "menge": "1",
                    "preis": str(mahngebuehr),
                    "ust": "0"  # Mahngebühren sind umsatzsteuerfrei
                })
                self.rechnungsdatenbank[rechnungsnr]["artikel"] = artikel
            
            # Daten speichern
            self.speichere_daten()
            
            # Liste aktualisieren
            self.aktualisiere_offene_rechnungen()
            
            messagebox.showinfo("Erfolg", f"Mahnung für Rechnung {rechnungsnr} (Stufe {mahnstufe}) wurde erstellt!")
            
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Erstellen der Mahnung: {str(e)}")
    
    def mahnvorlagen_verwalten(self):
        """Öffnet ein Fenster zur Verwaltung von Mahnvorlagen"""
        vorlagen_fenster = tk.Toplevel(self.root)
        vorlagen_fenster.title("Mahnvorlagen verwalten")
        vorlagen_fenster.geometry("600x500")
        
        # Linke Seite: Liste der Vorlagen
        liste_frame = ttk.Frame(vorlagen_fenster, padding=10)
        liste_frame.pack(side="left", fill="both", expand=True)
        
        ttk.Label(liste_frame, text="Gespeicherte Vorlagen:").pack(anchor="w")
        
        # Vorlagenliste
        vorlagen_list_frame = ttk.Frame(liste_frame)
        vorlagen_list_frame.pack(fill="both", expand=True, pady=5)
        
        self.vorlagen_listbox = tk.Listbox(vorlagen_list_frame, height=10, width=30)
        vorlagen_scrollbar = ttk.Scrollbar(vorlagen_list_frame, orient="vertical", 
                                          command=self.vorlagen_listbox.yview)
        
        self.vorlagen_listbox.pack(side="left", fill="both", expand=True)
        vorlagen_scrollbar.pack(side="right", fill="y")
        self.vorlagen_listbox.config(yscrollcommand=vorlagen_scrollbar.set)
        
        self.vorlagen_listbox.bind('<<ListboxSelect>>', self.mahnvorlage_auswaehlen)
        
        button_frame = ttk.Frame(liste_frame)
        button_frame.pack(fill="x", pady=5)
        
        ttk.Button(button_frame, text="Neu", command=lambda: self.neue_mahnvorlage(vorlagen_fenster)).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Löschen", command=lambda: self.loesche_mahnvorlage(vorlagen_fenster)).pack(side="left", padx=5)
        
        # Rechte Seite: Vorlagendetails
        details_frame = ttk.Frame(vorlagen_fenster, padding=10)
        details_frame.pack(side="right", fill="both", expand=True)
        
        ttk.Label(details_frame, text="Name:").grid(row=0, column=0, sticky="w", pady=5)
        self.vorlage_name_var = tk.StringVar()
        ttk.Entry(details_frame, textvariable=self.vorlage_name_var, width=30).grid(
            row=0, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Mahnstufe:").grid(row=1, column=0, sticky="w", pady=5)
        self.vorlage_stufe_var = tk.StringVar(value="1")
        ttk.Combobox(details_frame, textvariable=self.vorlage_stufe_var, width=15,
                    values=["1", "2", "3", "Letzte Mahnung"]).grid(
            row=1, column=1, sticky="w", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Betreff:").grid(row=2, column=0, sticky="w", pady=5)
        self.vorlage_betreff_var = tk.StringVar(
            value="Mahnung zur Rechnung {rechnungsnr}")
        ttk.Entry(details_frame, textvariable=self.vorlage_betreff_var, width=40).grid(
            row=2, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(details_frame, text="Text:").grid(row=3, column=0, sticky="nw", pady=5)
        
        text_frame = ttk.Frame(details_frame)
        text_frame.grid(row=3, column=1, sticky="ew", pady=5, padx=5)
        
        self.vorlage_text = tk.Text(text_frame, height=10, width=40)
        self.vorlage_text.pack(side="left", fill="both", expand=True)
        text_scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=self.vorlage_text.yview)
        text_scrollbar.pack(side="right", fill="y")
        self.vorlage_text.config(yscrollcommand=text_scrollbar.set)
        
        ttk.Label(details_frame, text="Verfügbare Platzhalter:").grid(row=4, column=0, columnspan=2, sticky="w", pady=5)
        ttk.Label(details_frame, text="{rechnungsnr}, {datum}, {betrag}, {frist}").grid(
            row=5, column=0, columnspan=2, sticky="w", pady=5)
        
        # Buttons
        bottom_frame = ttk.Frame(details_frame)
        bottom_frame.grid(row=6, column=0, columnspan=2, pady=10)
        
        ttk.Button(bottom_frame, text="Speichern", 
                  command=lambda: self.speichere_mahnvorlage(vorlagen_fenster)).pack(side="left", padx=5)
        ttk.Button(bottom_frame, text="Übernehmen", 
                  command=lambda: self.uebernehme_mahnvorlage(vorlagen_fenster)).pack(side="left", padx=5)
        ttk.Button(bottom_frame, text="Schließen", 
                  command=vorlagen_fenster.destroy).pack(side="left", padx=5)
        
        # Vorlagen in Listbox füllen
        self.aktualisiere_mahnvorlagen_liste()
        
        # Fenster modal machen
        vorlagen_fenster.transient(self.root)
        vorlagen_fenster.grab_set()
        self.root.wait_window(vorlagen_fenster)
    
    def aktualisiere_mahnvorlagen_liste(self):
        """Aktualisiert die Liste der Mahnvorlagen"""
        self.vorlagen_listbox.delete(0, tk.END)
        
        # Sortierte Liste erstellen
        vorlagen_liste = []
        
        for name, vorlage in self.mahnungen.items():
            stufe = vorlage.get("stufe", "1")
            vorlagen_liste.append((name, stufe))
        
        # Nach Namen sortieren
        vorlagen_liste.sort()
        
        # Anzeigen
        for name, stufe in vorlagen_liste:
            self.vorlagen_listbox.insert(tk.END, f"{name} (Stufe {stufe})")
            # Name als Attribut speichern
            self.vorlagen_listbox.itemconfig(self.vorlagen_listbox.size()-1, {'name': name})
    
    def mahnvorlage_auswaehlen(self, event):
        """Handler für die Auswahl einer Mahnvorlage aus der Liste"""
        selection = self.vorlagen_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        # Namen aus Attribut holen
        name = self.vorlagen_listbox.itemcget(index, 'name')
        
        if name and name in self.mahnungen:
            vorlage = self.mahnungen[name]
            
            # Formular befüllen
            self.vorlage_name_var.set(name)
            self.vorlage_stufe_var.set(vorlage.get("stufe", "1"))
            self.vorlage_betreff_var.set(vorlage.get("betreff", "Mahnung zur Rechnung {rechnungsnr}"))
            
            # Text
            self.vorlage_text.delete("1.0", tk.END)
            self.vorlage_text.insert("1.0", vorlage.get("text", ""))
    
    def neue_mahnvorlage(self, parent_window):
        """Leert das Formular für eine neue Mahnvorlage"""
        self.vorlage_name_var.set("")
        self.vorlage_stufe_var.set("1")
        self.vorlage_betreff_var.set("Mahnung zur Rechnung {rechnungsnr}")
        
        # Standard-Text für die 1. Mahnung
        standard_text = (
            "Sehr geehrte Damen und Herren,\n\n"
            "leider konnten wir bislang keinen Zahlungseingang zu unserer Rechnung {rechnungsnr} vom {datum} "
            "feststellen. Wir bitten Sie, den Rechnungsbetrag in Höhe von {betrag} € "
            "innerhalb der nächsten {frist} Tage zu begleichen.\n\n"
            "Mit freundlichen Grüßen"
        )
        
        self.vorlage_text.delete("1.0", tk.END)
        self.vorlage_text.insert("1.0", standard_text)
    
    def loesche_mahnvorlage(self, parent_window):
        """Löscht die ausgewählte Mahnvorlage"""
        selection = self.vorlagen_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie eine Vorlage aus.")
            return
        
        index = selection[0]
        # Namen aus Attribut holen
        name = self.vorlagen_listbox.itemcget(index, 'name')
        
        if name and name in self.mahnungen:
            if messagebox.askyesno("Löschen bestätigen", 
                                  f"Möchten Sie die Mahnvorlage '{name}' wirklich löschen?"):
                
                # Aus Mahnungen entfernen
                del self.mahnungen[name]
                
                # Daten speichern
                self.speichere_daten()
                
                # Liste aktualisieren
                self.aktualisiere_mahnvorlagen_liste()
                
                # Formular leeren
                self.neue_mahnvorlage(parent_window)
    
    def speichere_mahnvorlage(self, parent_window):
        """Speichert die aktuelle Mahnvorlage"""
        name = self.vorlage_name_var.get().strip()
        
        if not name:
            messagebox.showerror("Fehler", "Bitte geben Sie einen Namen für die Vorlage ein!")
            return
        
        # Vorlagendaten sammeln
        vorlage = {
            "stufe": self.vorlage_stufe_var.get(),
            "betreff": self.vorlage_betreff_var.get(),
            "text": self.vorlage_text.get("1.0", "end-1c")
        }
        
        # Vorlage speichern
        self.mahnungen[name] = vorlage
        
        # Daten speichern
        self.speichere_daten()
        
        # Liste aktualisieren
        self.aktualisiere_mahnvorlagen_liste()
        
        messagebox.showinfo("Erfolg", f"Mahnvorlage '{name}' wurde gespeichert!")
    
    def uebernehme_mahnvorlage(self, parent_window):
        """Übernimmt die ausgewählte Mahnvorlage in die Mahnung"""
        selection = self.vorlagen_listbox.curselection()
        if not selection:
            messagebox.showinfo("Information", "Bitte wählen Sie eine Vorlage aus.")
            return
        
        index = selection[0]
        # Namen aus Attribut holen
        name = self.vorlagen_listbox.itemcget(index, 'name')
        
        if name and name in self.mahnungen:
            vorlage = self.mahnungen[name]
            
            # In das Hauptfenster übernehmen
            self.mahnung_neue_stufe_var.set(vorlage.get("stufe", "1"))
            self.mahnung_email_betreff_var.set(vorlage.get("betreff", "Mahnung zur Rechnung {rechnungsnr}"))
            
            # Text
            self.mahnung_email_text.delete("1.0", tk.END)
            self.mahnung_email_text.insert("1.0", vorlage.get("text", ""))
            
            # Fenster schließen
            parent_window.destroy()
    
    def automatischer_mahnlauf(self):
        """Führt einen automatischen Mahnlauf für alle überfälligen Rechnungen durch"""
        if not messagebox.askyesno("Automatischer Mahnlauf", 
                                  "Möchten Sie einen automatischen Mahnlauf für alle überfälligen Rechnungen durchführen?"):
            return
        
        # Aktuelles Datum für Fälligkeitsberechnung
        heute = datetime.datetime.now().date()
        
        # Überfällige Rechnungen sammeln
        ueberfaellige_rechnungen = []
        
        for rechnungsnr, rechnung in self.rechnungsdatenbank.items():
            # Nur offene Rechnungen berücksichtigen
            if rechnung.get("bezahlt", False):
                continue
            
            # Rechnungsdaten extrahieren
            datum = rechnung.get("datum", "")
            zahlungsziel = int(rechnung.get("zahlungsziel", "14"))
            
            # Fälligkeitsdatum berechnen
            try:
                rechnungsdatum = datetime.datetime.strptime(datum, "%d.%m.%Y").date()
                faelligkeit = rechnungsdatum + datetime.timedelta(days=zahlungsziel)
                tage_ueberfaellig = (heute - faelligkeit).days
                
                # Rechnung ist überfällig und noch nicht auf höchster Mahnstufe
                if tage_ueberfaellig > 0 and rechnung.get("mahnstufe", 0) < 4:
                    ueberfaellige_rechnungen.append({
                        "rechnungsnr": rechnungsnr,
                        "rechnung": rechnung,
                        "tage_ueberfaellig": tage_ueberfaellig
                    })
            except:
                continue
        
        if not ueberfaellige_rechnungen:
            messagebox.showinfo("Information", "Es wurden keine überfälligen Rechnungen gefunden.")
            return
        
        # Abfragen, ob alle Rechnungen gemahnt werden sollen
        mahnlauf_text = f"Es wurden {len(ueberfaellige_rechnungen)} überfällige Rechnungen gefunden. Möchten Sie für alle eine Mahnung erstellen?"
        if not messagebox.askyesno("Bestätigung", mahnlauf_text):
            return
        
        # Für jede überfällige Rechnung eine Mahnung erstellen
        erfolgreiche_mahnungen = 0
        
        for item in ueberfaellige_rechnungen:
            rechnungsnr = item["rechnungsnr"]
            rechnung = item["rechnung"]
            
            # Rechnung auswählen
            self.mahnung_rechnungsnr_var.set(rechnungsnr)
            
            # Details befüllen
            self.mahnung_rechnungsdatum_var.set(rechnung.get("datum", ""))
            
            # Zahlungsziel
            zahlungsziel = rechnung.get("zahlungsziel", "14")
            datum = rechnung.get("datum", "")
            
            try:
                rechnungsdatum = datetime.datetime.strptime(datum, "%d.%m.%Y").date()
                faelligkeit = rechnungsdatum + datetime.timedelta(days=int(zahlungsziel))
                self.mahnung_zahlungsziel_var.set(faelligkeit.strftime("%d.%m.%Y"))
            except:
                self.mahnung_zahlungsziel_var.set(f"{zahlungsziel} Tage nach Rechnungsdatum")
            
            # Betrag berechnen
            betrag = self.berechne_rechnungsbetrag(rechnung)
            self.mahnung_betrag_var.set(f"{betrag:.2f} €")
            
            # Empfänger
            self.mahnung_empfaenger_var.set(rechnung.get("empfaenger", ""))
            
            # Mahnstatus
            mahnstufe = rechnung.get("mahnstufe", 0)
            self.mahnung_stufe_var.set(str(mahnstufe))
            
            letzte_mahnung = rechnung.get("letzte_mahnung", "")
            self.mahnung_letzte_var.set(letzte_mahnung)
            
            # Nächste Mahnstufe
            naechste_stufe = mahnstufe + 1
            if naechste_stufe > 3:
                self.mahnung_neue_stufe_var.set("Letzte Mahnung")
            else:
                self.mahnung_neue_stufe_var.set(str(naechste_stufe))
            
            # Text für die Mahnung aktualisieren
            self.aktualisiere_mahnung_text(rechnung)
            
            # Mahnung erstellen
            try:
                self.erstelle_mahnung()
                erfolgreiche_mahnungen += 1
            except Exception as e:
                print(f"Fehler beim Erstellen der Mahnung für Rechnung {rechnungsnr}: {str(e)}")
        
        # Meldung über erfolgreiche Mahnungen
        messagebox.showinfo("Mahnlauf abgeschlossen", 
                           f"Es wurden {erfolgreiche_mahnungen} von {len(ueberfaellige_rechnungen)} Mahnungen erfolgreich erstellt.")
    
    def init_rechnungsbearbeitung_tab(self):
        """Initialisiert den Tab für die Rechnungsbearbeitung"""
        frame = ttk.Frame(self.tab_rechnungsbearbeitung, padding="10")
        frame.pack(fill='both', expand=True)
        
        # Oberer Bereich: Rechnungsimport
        import_frame = ttk.LabelFrame(frame, text="X-Rechnung importieren")
        import_frame.pack(fill="x", pady=10)
        
        # Datei-Auswahl
        file_frame = ttk.Frame(import_frame)
        file_frame.pack(fill="x", pady=5, padx=10)
        
        ttk.Label(file_frame, text="Datei:").grid(row=0, column=0, sticky="w", pady=5)
        self.xrechnung_dateipfad_var = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.xrechnung_dateipfad_var, width=50).grid(
            row=0, column=1, sticky="ew", pady=5, padx=5)
        ttk.Button(file_frame, text="Durchsuchen", 
                  command=self.xrechnung_datei_auswaehlen).grid(row=0, column=2, padx=5)
        
        # Import-Button
        ttk.Button(import_frame, text="X-Rechnung importieren", 
                  command=self.xrechnung_importieren).pack(pady=10)
        
        # Mittlerer Bereich: Rechnungsdetails
        details_frame = ttk.LabelFrame(frame, text="Rechnungsdetails")
        details_frame.pack(fill="both", expand=True, pady=10)
        
        # Links: Rechnungskopf
        kopf_frame = ttk.Frame(details_frame)
        kopf_frame.pack(side="left", fill="both", expand=True, padx=(10, 5), pady=10)
        
        ttk.Label(kopf_frame, text="Rechnungsnummer:").grid(row=0, column=0, sticky="w", pady=5)
        self.xrechnung_nr_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_nr_var, width=30).grid(
            row=0, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(kopf_frame, text="Rechnungsdatum:").grid(row=1, column=0, sticky="w", pady=5)
        self.xrechnung_datum_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_datum_var, width=30).grid(
            row=1, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(kopf_frame, text="Lieferant:").grid(row=2, column=0, sticky="w", pady=5)
        self.xrechnung_lieferant_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_lieferant_var, width=30).grid(
            row=2, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(kopf_frame, text="Nettobetrag:").grid(row=3, column=0, sticky="w", pady=5)
        self.xrechnung_netto_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_netto_var, width=30).grid(
            row=3, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(kopf_frame, text="USt-Betrag:").grid(row=4, column=0, sticky="w", pady=5)
        self.xrechnung_ust_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_ust_var, width=30).grid(
            row=4, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(kopf_frame, text="Bruttobetrag:").grid(row=5, column=0, sticky="w", pady=5)
        self.xrechnung_brutto_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_brutto_var, width=30).grid(
            row=5, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(kopf_frame, text="Fälligkeitsdatum:").grid(row=6, column=0, sticky="w", pady=5)
        self.xrechnung_faelligkeit_var = tk.StringVar()
        ttk.Entry(kopf_frame, textvariable=self.xrechnung_faelligkeit_var, width=30).grid(
            row=6, column=1, sticky="ew", pady=5, padx=5)
        
        # Rechts: Artikelliste
        artikel_frame = ttk.Frame(details_frame)
        artikel_frame.pack(side="right", fill="both", expand=True, padx=(5, 10), pady=10)
        
        ttk.Label(artikel_frame, text="Positionen:").pack(anchor="w")
        
        # Listbox mit Scrollbar
        list_container = ttk.Frame(artikel_frame)
        list_container.pack(fill="both", expand=True, pady=5)
        
        scrollbar = ttk.Scrollbar(list_container)
        scrollbar.pack(side="right", fill="y")
        
        self.xrechnung_artikel_listbox = tk.Listbox(list_container, height=10, width=50, 
                                                 yscrollcommand=scrollbar.set)
        self.xrechnung_artikel_listbox.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=self.xrechnung_artikel_listbox.yview)
        
        # Unterer Bereich: Aktionen
        aktionen_frame = ttk.LabelFrame(frame, text="Aktionen")
        aktionen_frame.pack(fill="x", pady=10)
        
        # Zahlungsinformationen
        zahlungs_frame = ttk.Frame(aktionen_frame)
        zahlungs_frame.pack(fill="x", pady=5, padx=10)
        
        ttk.Label(zahlungs_frame, text="IBAN:").grid(row=0, column=0, sticky="w", pady=5)
        self.xrechnung_iban_var = tk.StringVar()
        ttk.Entry(zahlungs_frame, textvariable=self.xrechnung_iban_var, width=30).grid(
            row=0, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(zahlungs_frame, text="BIC:").grid(row=1, column=0, sticky="w", pady=5)
        self.xrechnung_bic_var = tk.StringVar()
        ttk.Entry(zahlungs_frame, textvariable=self.xrechnung_bic_var, width=30).grid(
            row=1, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(zahlungs_frame, text="Verwendungszweck:").grid(row=2, column=0, sticky="w", pady=5)
        self.xrechnung_vzweck_var = tk.StringVar()
        ttk.Entry(zahlungs_frame, textvariable=self.xrechnung_vzweck_var, width=30).grid(
            row=2, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(zahlungs_frame, text="Zahlungsart:").grid(row=0, column=2, sticky="w", pady=5, padx=(20, 0))
        self.xrechnung_zahlungsart_var = tk.StringVar(value="Überweisung")
        ttk.Combobox(zahlungs_frame, textvariable=self.xrechnung_zahlungsart_var, width=15,
                     values=["Überweisung", "Lastschrift", "Kreditkarte"]).grid(
            row=0, column=3, sticky="ew", pady=5, padx=5)
        
        ttk.Label(zahlungs_frame, text="Zahlungsdatum:").grid(row=1, column=2, sticky="w", pady=5, padx=(20, 0))
        self.xrechnung_zahlungsdatum_var = tk.StringVar(value=datetime.datetime.now().strftime("%d.%m.%Y"))
        ttk.Entry(zahlungs_frame, textvariable=self.xrechnung_zahlungsdatum_var, width=15).grid(
            row=1, column=3, sticky="ew", pady=5, padx=5)
        
        # Buttons
        button_frame = ttk.Frame(aktionen_frame)
        button_frame.pack(fill="x", pady=10)
        
        ttk.Button(button_frame, text="Zahlung erstellen (SEPA XML)", 
                  command=self.erstelle_sepa_xml).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Zahlung in Bankprogramm öffnen", 
                  command=self.oeffne_bankprogramm).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Als bezahlt markieren", 
                  command=self.xrechnung_als_bezahlt_markieren).pack(side="left", padx=5)
    
    def xrechnung_datei_auswaehlen(self):
        """Öffnet einen Dateiauswahldialog für X-Rechnungen"""
        filepath = filedialog.askopenfilename(
            title="X-Rechnung auswählen",
            filetypes=[("XML-Dateien", "*.xml"), ("Alle Dateien", "*.*")]
        )
        
        if filepath:
            self.xrechnung_dateipfad_var.set(filepath)
    
    def xrechnung_importieren(self):
        """Importiert eine X-Rechnung und zeigt die Details an"""
        filepath = self.xrechnung_dateipfad_var.get()
        
        if not filepath or not os.path.exists(filepath):
            messagebox.showerror("Fehler", "Bitte wählen Sie eine gültige X-Rechnungsdatei aus!")
            return
        
        try:
            # XML-Datei parsen
            tree = ET.parse(filepath)
            root = tree.getroot()
            
            # Namespaces für X-Rechnung
            ns = {
                'xr': 'urn:ce.eu:en16931:2017:xoev-de:kosit:standard:xrechnung-1',
                'ram': 'urn:un:unece:uncefact:data:standard:ReusableAggregateBusinessInformationEntity:100',
                'udt': 'urn:un:unece:uncefact:data:standard:UnqualifiedDataType:100',
                'rsm': 'urn:un:unece:uncefact:data:standard:CrossIndustryInvoice:100'
            }
            
            # Grunddaten auslesen
            try:
                # Rechnungsnummer
                rechnungsnr = root.find('.//rsm:ExchangedDocument/ram:ID', ns).text
                self.xrechnung_nr_var.set(rechnungsnr)
                
                # Rechnungsdatum
                datum_str = root.find('.//rsm:ExchangedDocument/ram:IssueDateTime/udt:DateTimeString', ns).text
                datum = datetime.datetime.strptime(datum_str, "%Y%m%d").strftime("%d.%m.%Y")
                self.xrechnung_datum_var.set(datum)
                
                # Lieferant
                lieferant_name = root.find('.//ram:SellerTradeParty/ram:Name', ns).text
                self.xrechnung_lieferant_var.set(lieferant_name)
                
                # Beträge
                netto = root.find('.//ram:LineTotalAmount', ns).text
                self.xrechnung_netto_var.set(netto)
                
                ust = root.find('.//ram:TaxTotalAmount', ns).text
                self.xrechnung_ust_var.set(ust)
                
                brutto = root.find('.//ram:GrandTotalAmount', ns).text
                self.xrechnung_brutto_var.set(brutto)
                
                # Fälligkeitsdatum
                try:
                    faelligkeit_str = root.find('.//ram:DueDateDateTime/udt:DateTimeString', ns).text
                    faelligkeit = datetime.datetime.strptime(faelligkeit_str, "%Y%m%d").strftime("%d.%m.%Y")
                    self.xrechnung_faelligkeit_var.set(faelligkeit)
                except:
                    self.xrechnung_faelligkeit_var.set("")
                
                # Zahlungsinformationen
                try:
                    iban = root.find('.//ram:IBANID', ns).text
                    self.xrechnung_iban_var.set(iban)
                except:
                    self.xrechnung_iban_var.set("")
                
                try:
                    bic = root.find('.//ram:BICID', ns).text
                    self.xrechnung_bic_var.set(bic)
                except:
                    self.xrechnung_bic_var.set("")
                
                try:
                    vzweck = root.find('.//ram:PaymentReference', ns).text
                    self.xrechnung_vzweck_var.set(vzweck)
                except:
                    self.xrechnung_vzweck_var.set(rechnungsnr)
                
                # Artikel auslesen
                self.xrechnung_artikel_listbox.delete(0, tk.END)
                artikel_elemente = root.findall('.//ram:IncludedSupplyChainTradeLineItem', ns)
                
                for artikel in artikel_elemente:
                    bezeichnung = artikel.find('.//ram:ItemName', ns).text
                    menge = artikel.find('.//ram:BilledQuantity', ns).text
                    preis = artikel.find('.//ram:NetPrice/ram:ChargeAmount', ns).text
                    
                    try:
                        ust_satz = artikel.find('.//ram:ApplicableTradeTax/ram:RateApplicablePercent', ns).text
                    except:
                        ust_satz = ""
                    
                    display_text = f"{bezeichnung} - {menge} x {preis} € ({ust_satz}% USt)"
                    self.xrechnung_artikel_listbox.insert(tk.END, display_text)
                
                messagebox.showinfo("Erfolg", "X-Rechnung wurde erfolgreich importiert!")
                
            except Exception as e:
                messagebox.showerror("Fehler bei der Verarbeitung", 
                                   f"Fehler beim Auslesen der X-Rechnung: {str(e)}")
        
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Öffnen der XML-Datei: {str(e)}")
    
    def erstelle_sepa_xml(self):
        """Erstellt eine SEPA-XML-Datei für die Überweisung"""
        # Prüfen, ob alle benötigten Felder ausgefüllt sind
        iban = self.xrechnung_iban_var.get().strip()
        bic = self.xrechnung_bic_var.get().strip()
        betrag = self.xrechnung_brutto_var.get().strip()
        empfaenger = self.xrechnung_lieferant_var.get().strip()
        verwendungszweck = self.xrechnung_vzweck_var.get().strip()
        
        if not iban or not betrag or not empfaenger:
            messagebox.showerror("Fehler", "Bitte füllen Sie mindestens IBAN, Betrag und Empfänger aus!")
            return
        
        # IBAN und BIC validieren
        if not self.validiere_iban(iban):
            messagebox.showerror("Fehler", "Die eingegebene IBAN ist ungültig!")
            return
        
        if bic and not self.validiere_bic(bic):
            messagebox.showerror("Fehler", "Die eingegebene BIC ist ungültig!")
            return
        
        # Betrag validieren und formatieren
        try:
            betrag_dec = Decimal(betrag.replace(',', '.'))
            betrag_str = f"{betrag_dec:.2f}".replace('.', ',')  # SEPA verwendet Komma als Dezimaltrennzeichen
        except:
            messagebox.showerror("Fehler", "Der eingegebene Betrag ist ungültig!")
            return
        
        # Zahlungsdatum formatieren
        zahlungsdatum = self.xrechnung_zahlungsdatum_var.get()
        try:
            datum = datetime.datetime.strptime(zahlungsdatum, "%d.%m.%Y")
            datum_str = datum.strftime("%Y-%m-%d")
        except:
            messagebox.showerror("Fehler", "Das eingegebene Zahlungsdatum ist ungültig!")
            return
        
        # Eigene Bankdaten aus Firmendaten holen
        eigene_iban = self.firmendaten.get("IBAN", "")
        eigene_bic = self.firmendaten.get("BIC", "")
        eigener_name = self.firmendaten.get("Firmenname", "")
        
        if not eigene_iban or not eigener_name:
            messagebox.showerror("Fehler", "Bitte ergänzen Sie Ihre IBAN und Ihren Firmennamen in den Firmendaten!")
            return
        
        # SEPA-XML erstellen
        try:
            # Root-Element
            root = ET.Element("Document")
            root.set("xmlns", "urn:iso:std:iso:20022:tech:xsd:pain.001.001.03")
            root.set("xmlns:xsi", "http://www.w3.org/2001/XMLSchema-instance")
            
            # CstmrCdtTrfInitn (Customer Credit Transfer Initiation)
            cstmr_init = ET.SubElement(root, "CstmrCdtTrfInitn")
            
            # GrpHdr (Group Header)
            grp_hdr = ET.SubElement(cstmr_init, "GrpHdr")
            
            # Generiere eine Message-ID
            msg_id = f"MSG-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
            ET.SubElement(grp_hdr, "MsgId").text = msg_id
            
            # Erstellungsdatum und -zeit
            now = datetime.datetime.now()
            ET.SubElement(grp_hdr, "CreDtTm").text = now.strftime("%Y-%m-%dT%H:%M:%S")
            
            # Anzahl der Transaktionen
            ET.SubElement(grp_hdr, "NbOfTxs").text = "1"
            
            # Initiator (Auftraggeber)
            init_party = ET.SubElement(grp_hdr, "InitgPty")
            ET.SubElement(init_party, "Nm").text = eigener_name
            
            # PmtInf (Payment Information)
            pmt_inf = ET.SubElement(cstmr_init, "PmtInf")
            
            # Zahlungs-ID
            ET.SubElement(pmt_inf, "PmtInfId").text = f"PMT-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
            
            # Zahlungsart
            ET.SubElement(pmt_inf, "PmtMtd").text = "TRF"  # TRF = Transfer (Überweisung)
            
            # Anzahl der Buchungen in dieser Zahlung
            ET.SubElement(pmt_inf, "NbOfTxs").text = "1"
            
            # Ausführungsdatum
            ET.SubElement(pmt_inf, "ReqdExctnDt").text = datum_str
            
            # Debitor (Zahler)
            debtor = ET.SubElement(pmt_inf, "Dbtr")
            ET.SubElement(debtor, "Nm").text = eigener_name
            
            # Debitor-Konto
            dbtr_acct = ET.SubElement(pmt_inf, "DbtrAcct")
            dbtr_id = ET.SubElement(dbtr_acct, "Id")
            ET.SubElement(dbtr_id, "IBAN").text = eigene_iban
            
            # Debitor-Bank
            dbtr_agt = ET.SubElement(pmt_inf, "DbtrAgt")
            fin_inst_id = ET.SubElement(dbtr_agt, "FinInstnId")
            if eigene_bic:
                ET.SubElement(fin_inst_id, "BIC").text = eigene_bic
            else:
                ET.SubElement(fin_inst_id, "Othr").append(ET.Element("Id", text="NOTPROVIDED"))
            
            # Credit Transfer Transaction Information (einzelne Überweisung)
            cdtTrfTxInf = ET.SubElement(pmt_inf, "CdtTrfTxInf")
            
            # Zahlungs-ID für diese Transaktion
            pmt_id = ET.SubElement(cdtTrfTxInf, "PmtId")
            ET.SubElement(pmt_id, "EndToEndId").text = f"NOTPROVIDED"
            
            # Betrag
            amt = ET.SubElement(cdtTrfTxInf, "Amt")
            instd_amt = ET.SubElement(amt, "InstdAmt")
            instd_amt.set("Ccy", "EUR")
            instd_amt.text = betrag_str.replace(',', '.')  # XML verwendet Punkt als Dezimaltrennzeichen
            
            # Kreditor-Bank
            cdtr_agt = ET.SubElement(cdtTrfTxInf, "CdtrAgt")
            fin_inst_id = ET.SubElement(cdtr_agt, "FinInstnId")
            if bic:
                ET.SubElement(fin_inst_id, "BIC").text = bic
            else:
                ET.SubElement(fin_inst_id, "Othr").append(ET.Element("Id", text="NOTPROVIDED"))
            
            # Kreditor (Empfänger)
            cdtr = ET.SubElement(cdtTrfTxInf, "Cdtr")
            ET.SubElement(cdtr, "Nm").text = empfaenger
            
            # Kreditor-Konto
            cdtr_acct = ET.SubElement(cdtTrfTxInf, "CdtrAcct")
            cdtr_id = ET.SubElement(cdtr_acct, "Id")
            ET.SubElement(cdtr_id, "IBAN").text = iban
            
            # Verwendungszweck
            if verwendungszweck:
                rmtInf = ET.SubElement(cdtTrfTxInf, "RmtInf")
                ET.SubElement(rmtInf, "Ustrd").text = verwendungszweck
            
            # XML formatieren
            xml_str = ET.tostring(root, encoding='utf-8')
            dom = minidom.parseString(xml_str)
            pretty_xml = dom.toprettyxml(indent="  ")
            
            # Datei speichern
            filepath = filedialog.asksaveasfilename(
                defaultextension=".xml",
                filetypes=[("XML-Dateien", "*.xml")],
                initialfile=f"SEPA_Ueberweisung_{datum.strftime('%Y-%m-%d')}.xml",
                title="SEPA-XML speichern"
            )
            
            if filepath:
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(pretty_xml)
                
                messagebox.showinfo("Erfolg", f"SEPA-XML wurde unter {filepath} gespeichert!")
        
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Erstellen der SEPA-XML: {str(e)}")
    
    def validiere_iban(self, iban):
        """Einfache Validierung einer IBAN"""
        # Leerzeichen entfernen
        iban = iban.replace(" ", "")
        
        # Länge prüfen (min. 15, max. 34 Zeichen)
        if len(iban) < 15 or len(iban) > 34:
            return False
        
        # Prüfung der Länderkennung und der Prüfziffer
        if not iban[:2].isalpha() or not iban[2:4].isdigit():
            return False
        
        # Weitere Zeichen müssen alphanumerisch sein
        if not all(c.isalnum() for c in iban):
            return False
        
        # Einfache IBAN-Prüfung gilt als bestanden
        return True
    
    def validiere_bic(self, bic):
        """Einfache Validierung einer BIC"""
        # Leerzeichen entfernen
        bic = bic.replace(" ", "")
        
        # Länge prüfen (8 oder 11 Zeichen)
        if len(bic) != 8 and len(bic) != 11:
            return False
        
        # Format-Prüfung: AAAABBCC(DDD)
        # AAAA = Bank-Code (4 Buchstaben)
        # BB = Ländercode (2 Buchstaben)
        # CC = Ortscode (2 alphanumerische Zeichen)
        # DDD = optional: Filialcode (3 alphanumerische Zeichen)
        
        if not bic[:4].isalpha() or not bic[4:6].isalpha() or not all(c.isalnum() for c in bic[6:]):
            return False
        
        # Einfache BIC-Prüfung gilt als bestanden
        return True
    
    def oeffne_bankprogramm(self):
        """Öffnet das Standard-Bankprogramm oder die Online-Banking-Webseite"""
        try:
            bank = self.firmendaten.get("Bank", "").lower()
            
            # Bekannte Banken und ihre Online-Banking-URLs
            bank_urls = {
                "sparkasse": "https://www.sparkasse.de/online-banking/",
                "volksbank": "https://www.volksbank.de/banking-private/",
                "deutsche bank": "https://meine.deutsche-bank.de/",
                "commerzbank": "https://banking.commerzbank.de/",
                "postbank": "https://banking.postbank.de/",
                "ing": "https://banking.ing.de/",
                "dkb": "https://banking.dkb.de/",
                "comdirect": "https://www.comdirect.de/",
            }
            
            # URL für die Bank suchen
            url = None
            for bank_name, bank_url in bank_urls.items():
                if bank_name in bank:
                    url = bank_url
                    break
            
            # Wenn keine spezifische Bank gefunden wurde, Standard-URL verwenden
            if not url:
                url = "https://www.sparkasse.de/online-banking/"
            
            # Browser öffnen
            webbrowser.open(url)
            
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Öffnen des Bankprogramms: {str(e)}")
    
    def xrechnung_als_bezahlt_markieren(self):
        """Markiert die importierte X-Rechnung als bezahlt (für Buchhaltung)"""
        rechnungsnr = self.xrechnung_nr_var.get()
        lieferant = self.xrechnung_lieferant_var.get()
        betrag = self.xrechnung_brutto_var.get()
        
        if not rechnungsnr or not lieferant or not betrag:
            messagebox.showerror("Fehler", "Bitte importieren Sie zuerst eine X-Rechnung!")
            return
        
        # Bestätigung einholen
        if not messagebox.askyesno("Bestätigung", 
                                  f"Möchten Sie die Rechnung {rechnungsnr} von {lieferant} über {betrag} € als bezahlt markieren?"):
            return
        
        # TODO: Hier könnte eine Integration mit einer Buchhaltungssoftware erfolgen
        
        messagebox.showinfo("Erfolg", f"Rechnung {rechnungsnr} wurde als bezahlt markiert!")
    
    def erstelle_xrechnung(self, rechnung_daten):
        """Erstellt eine X-Rechnung im XML-Format"""
        if not rechnung_daten:
            return None
        
        docx_pfad = rechnung_daten.get("docx")
        rechnungsnr = rechnung_daten.get("rechnungsnr")
        empfaenger = rechnung_daten.get("empfaenger")
        
        try:
            # Rechnungsdaten aus der Datenbank holen
            rechnung = self.rechnungsdatenbank.get(rechnungsnr)
            if not rechnung:
                return None
            
            # X-Rechnung XML erstellen
            root = ET.Element("CrossIndustryInvoice")
            root.set("xmlns", "urn:un:unece:uncefact:data:standard:CrossIndustryInvoice:100")
            
            # Rechnungskopf
            header = ET.SubElement(root, "ExchangedDocumentContext")
            ET.SubElement(header, "GuidelineSpecifiedDocumentContextParameter").text = "urn:cen.eu:en16931:2017"
            
            # Rechnungsdokument
            document = ET.SubElement(root, "ExchangedDocument")
            ET.SubElement(document, "ID").text = rechnungsnr
            
            # Rechnungsdatum
            datum = rechnung.get("datum", "")
            try:
                datum_obj = datetime.datetime.strptime(datum, "%d.%m.%Y")
                datum_str = datum_obj.strftime("%Y%m%d")
            except:
                datum_str = datetime.datetime.now().strftime("%Y%m%d")
            
            date_time = ET.SubElement(document, "IssueDateTime")
            dt_string = ET.SubElement(date_time, "DateTimeString")
            dt_string.set("format", "102")  # Format: YYYYMMDD
            dt_string.text = datum_str
            
            # Rechnungstyp
            ET.SubElement(document, "TypeCode").text = "380"  # 380 = Commercial Invoice
            
            # Eigene Firmendaten
            supply_chain = ET.SubElement(root, "SupplyChainTradeTransaction")
            agreement = ET.SubElement(supply_chain, "ApplicableHeaderTradeAgreement")
            
            seller_party = ET.SubElement(agreement, "SellerTradeParty")
            ET.SubElement(seller_party, "ID").text = self.firmendaten.get("USt-ID", "")
            ET.SubElement(seller_party, "Name").text = self.firmendaten.get("Firmenname", "")
            
            seller_address = ET.SubElement(seller_party, "PostalTradeAddress")
            ET.SubElement(seller_address, "PostcodeCode").text = self.firmendaten.get("PLZ", "")
            ET.SubElement(seller_address, "LineOne").text = self.firmendaten.get("Straße", "")
            ET.SubElement(seller_address, "CityName").text = self.firmendaten.get("Ort", "")
            ET.SubElement(seller_address, "CountryID").text = "DE"  # Deutschland
            
            # Empfängerdaten
            buyer_party = ET.SubElement(agreement, "BuyerTradeParty")
            ET.SubElement(buyer_party, "ID").text = empfaenger.get("USt-ID", "")
            
            if empfaenger.get("Firma"):
                ET.SubElement(buyer_party, "Name").text = empfaenger["Firma"]
            else:
                name_str = ""
                if empfaenger.get("Vorname"):
                    name_str += empfaenger["Vorname"] + " "
                if empfaenger.get("Nachname"):
                    name_str += empfaenger["Nachname"]
                ET.SubElement(buyer_party, "Name").text = name_str.strip()
            
            buyer_address = ET.SubElement(buyer_party, "PostalTradeAddress")
            ET.SubElement(buyer_address, "PostcodeCode").text = empfaenger.get("PLZ", "")
            ET.SubElement(buyer_address, "LineOne").text = empfaenger.get("Straße", "")
            ET.SubElement(buyer_address, "CityName").text = empfaenger.get("Ort", "")
            ET.SubElement(buyer_address, "CountryID").text = empfaenger.get("Land", "DE")
            
            # Lieferdaten
            delivery = ET.SubElement(supply_chain, "ApplicableHeaderTradeDelivery")
            
            # Leistungszeitraum
            if rechnung.get("leistungszeitraum_start") and rechnung.get("leistungszeitraum_ende"):
                start_date = rechnung["leistungszeitraum_start"]
                end_date = rechnung["leistungszeitraum_ende"]
                
                try:
                    start_obj = datetime.datetime.strptime(start_date, "%d.%m.%Y")
                    end_obj = datetime.datetime.strptime(end_date, "%d.%m.%Y")
                    
                    start_str = start_obj.strftime("%Y%m%d")
                    end_str = end_obj.strftime("%Y%m%d")
                    
                    supply_period = ET.SubElement(delivery, "ActualDeliverySupplyChainEvent")
                    occurrence = ET.SubElement(supply_period, "OccurrenceDateTime")
                    
                    start_dt = ET.SubElement(occurrence, "StartDateTime")
                    start_dt_str = ET.SubElement(start_dt, "DateTimeString")
                    start_dt_str.set("format", "102")
                    start_dt_str.text = start_str
                    
                    end_dt = ET.SubElement(occurrence, "EndDateTime")
                    end_dt_str = ET.SubElement(end_dt, "DateTimeString")
                    end_dt_str.set("format", "102")
                    end_dt_str.text = end_str
                except:
                    pass
            
            # Zahlungsdaten
            settlement = ET.SubElement(supply_chain, "ApplicableHeaderTradeSettlement")
            
            # Währung
            ET.SubElement(settlement, "InvoiceCurrencyCode").text = "EUR"
            
            # Zahlungsbedingungen
            payment_terms = ET.SubElement(settlement, "SpecifiedTradePaymentTerms")
            
            # Zahlungsziel
            zahlungsziel = rechnung.get("zahlungsziel", "14")
            try:
                zahlungsziel_tage = int(zahlungsziel)
                faelligkeit = datetime.datetime.strptime(datum, "%d.%m.%Y") + datetime.timedelta(days=zahlungsziel_tage)
                faelligkeit_str = faelligkeit.strftime("%Y%m%d")
                
                due_date = ET.SubElement(payment_terms, "DueDateDateTime")
                due_dt_str = ET.SubElement(due_date, "DateTimeString")
                due_dt_str.set("format", "102")
                due_dt_str.text = faelligkeit_str
            except:
                pass
            
            # Bankdaten
            if self.firmendaten.get("IBAN") and self.firmendaten.get("BIC"):
                payment_means = ET.SubElement(settlement, "SpecifiedTradeSettlementPaymentMeans")
                ET.SubElement(payment_means, "TypeCode").text = "58"  # SEPA Credit Transfer
                
                payee_account = ET.SubElement(payment_means, "PayeePartyCreditorFinancialAccount")
                ET.SubElement(payee_account, "IBANID").text = self.firmendaten.get("IBAN")
                
                payee_agent = ET.SubElement(payment_means, "PayeeSpecifiedCreditorFinancialInstitution")
                ET.SubElement(payee_agent, "BICID").text = self.firmendaten.get("BIC")
            
            # Steuerdaten
            tax_summary = ET.SubElement(settlement, "ApplicableTradeTax")
            
            # Steuerschuldnerschaft
            if rechnung.get("steuerschuldnerschaft", False):
                ET.SubElement(tax_summary, "TypeCode").text = "VAT"
                ET.SubElement(tax_summary, "CategoryCode").text = "AE"  # Reverse Charge
                ET.SubElement(tax_summary, "BasisAmount").text = "0.00"
                ET.SubElement(tax_summary, "CalculatedAmount").text = "0.00"
                ET.SubElement(tax_summary, "RateApplicablePercent").text = "0.00"
            else:
                # Beträge berechnen
                artikel = rechnung.get("artikel", [])
                ust_betraege = {}
                nettosumme = Decimal('0.00')
                
                for artikel_item in artikel:
                    try:
                        menge = Decimal(artikel_item.get("menge", "1").replace(',', '.'))
                        preis = Decimal(artikel_item.get("preis", "0.00").replace(',', '.'))
                        ust_satz = Decimal(artikel_item.get("ust", "19").replace(',', '.'))
                        
                        position_netto = menge * preis
                        ust_betrag = position_netto * (ust_satz / 100)
                        
                        nettosumme += position_netto
                        
                        ust_key = str(ust_satz)
                        if ust_key not in ust_betraege:
                            ust_betraege[ust_key] = Decimal('0.00')
                        ust_betraege[ust_key] += ust_betrag
                        
                    except:
                        continue
                
                # Für jede Steuerrate eine Zeile
                for satz, betrag in ust_betraege.items():
                    tax_entry = ET.SubElement(settlement, "ApplicableTradeTax")
                    ET.SubElement(tax_entry, "TypeCode").text = "VAT"
                    ET.SubElement(tax_entry, "CategoryCode").text = "S"  # Standard rate
                    ET.SubElement(tax_entry, "BasisAmount").text = f"{sum([Decimal(a.get('menge', '1').replace(',', '.')) * Decimal(a.get('preis', '0.00').replace(',', '.')) for a in artikel if a.get('ust', '19').replace(',', '.') == satz]):.2f}"
                    ET.SubElement(tax_entry, "CalculatedAmount").text = f"{betrag:.2f}"
                    ET.SubElement(tax_entry, "RateApplicablePercent").text = satz
            
            # Gesamtbeträge
            monetary_summation = ET.SubElement(settlement, "SpecifiedTradeSettlementHeaderMonetarySummation")
            
            artikel = rechnung.get("artikel", [])
            nettosumme = sum([Decimal(a.get("menge", "1").replace(',', '.')) * Decimal(a.get("preis", "0.00").replace(',', '.')) for a in artikel])
            ust_betraege = {}
            
            for artikel_item in artikel:
                try:
                    menge = Decimal(artikel_item.get("menge", "1").replace(',', '.'))
                    preis = Decimal(artikel_item.get("preis", "0.00").replace(',', '.'))
                    ust_satz = Decimal(artikel_item.get("ust", "19").replace(',', '.'))
                    
                    position_netto = menge * preis
                    ust_betrag = position_netto * (ust_satz / 100)
                    
                    ust_key = str(ust_satz)
                    if ust_key not in ust_betraege:
                        ust_betraege[ust_key] = Decimal('0.00')
                    ust_betraege[ust_key] += ust_betrag
                    
                except:
                    continue
            
            gesamt_ust = sum(ust_betraege.values())
            bruttosumme = nettosumme + gesamt_ust
            
            ET.SubElement(monetary_summation, "LineTotalAmount").text = f"{nettosumme:.2f}"
            ET.SubElement(monetary_summation, "TaxBasisTotalAmount").text = f"{nettosumme:.2f}"
            ET.SubElement(monetary_summation, "TaxTotalAmount").text = f"{gesamt_ust:.2f}"
            ET.SubElement(monetary_summation, "GrandTotalAmount").text = f"{bruttosumme:.2f}"
            ET.SubElement(monetary_summation, "DuePayableAmount").text = f"{bruttosumme:.2f}"
            
            # Artikelpositionen
            i = 1
            for artikel_item in artikel:
                line_item = ET.SubElement(supply_chain, "IncludedSupplyChainTradeLineItem")
                
                # Position
                ET.SubElement(line_item, "AssociatedDocumentLineDocument").append(
                    ET.Element("LineID", text=str(i)))
                i += 1
                
                # Artikel
                artikel_bezeichnung = artikel_item.get("bezeichnung", "")
                menge = Decimal(artikel_item.get("menge", "1").replace(',', '.'))
                preis = Decimal(artikel_item.get("preis", "0.00").replace(',', '.'))
                ust_satz = Decimal(artikel_item.get("ust", "19").replace(',', '.'))
                
                position_netto = menge * preis
                ust_betrag = position_netto * (ust_satz / 100)
                position_brutto = position_netto + ust_betrag
                
                # Artikeldaten
                product = ET.SubElement(line_item, "SpecifiedTradeProduct")
                ET.SubElement(product, "Name").text = artikel_bezeichnung
                
                # Preisdaten
                price_spec = ET.SubElement(line_item, "SpecifiedLineTradeAgreement")
                net_price = ET.SubElement(price_spec, "NetPriceProductTradePrice")
                ET.SubElement(net_price, "ChargeAmount").text = f"{preis:.2f}"
                
                # Mengendaten
                delivery_spec = ET.SubElement(line_item, "SpecifiedLineTradeDelivery")
                ET.SubElement(delivery_spec, "BilledQuantity").text = f"{menge:.2f}"
                
                # Steuerdaten
                settlement_spec = ET.SubElement(line_item, "SpecifiedLineTradeSettlement")
                
                tax_spec = ET.SubElement(settlement_spec, "ApplicableTradeTax")
                ET.SubElement(tax_spec, "TypeCode").text = "VAT"
                ET.SubElement(tax_spec, "CategoryCode").text = "S"  # Standard rate
                ET.SubElement(tax_spec, "RateApplicablePercent").text = f"{ust_satz:.2f}"
                
                # Beträge
                monetary_spec = ET.SubElement(settlement_spec, "SpecifiedTradeSettlementLineMonetarySummation")
                ET.SubElement(monetary_spec, "LineTotalAmount").text = f"{position_netto:.2f}"
            
            # XML formatieren
            xml_str = ET.tostring(root, encoding='utf-8')
            dom = minidom.parseString(xml_str)
            pretty_xml = dom.toprettyxml(indent="  ")
            
            # Datei speichern
            if docx_pfad:
                # Pfad zum DOCX-Dokument extrahieren
                docx_dir = os.path.dirname(docx_pfad)
                docx_name = os.path.basename(docx_pfad)
                xml_name = f"XRechnung_{rechnungsnr}.xml"
                xml_path = os.path.join(docx_dir, xml_name)
                
                with open(xml_path, "w", encoding="utf-8") as f:
                    f.write(pretty_xml)
                
                return xml_path
            else:
                filepath = filedialog.asksaveasfilename(
                    defaultextension=".xml",
                    filetypes=[("XML-Dateien", "*.xml")],
                    initialfile=f"XRechnung_{rechnungsnr}.xml",
                    title="X-Rechnung speichern"
                )
                
                if filepath:
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(pretty_xml)
                    
                    return filepath
            
            return None
            
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Erstellen der X-Rechnung: {str(e)}")
            return None
    
    def speichere_firmendaten(self):
        # Daten aus GUI in Dictionary übertragen
        for feld, var in self.firmendaten_eintraege.items():
            self.firmendaten[feld] = var.get()
        
        self.speichere_daten()
        messagebox.showinfo("Erfolg", "Firmendaten wurden gespeichert!")
    
    def erstelle_briefkopf(self):
        # Prüfen, ob alle wichtigen Felder ausgefüllt sind
        required = ["Firmenname", "Straße", "PLZ", "Ort", "Steuernummer", "IBAN"]
        for feld in required:
            if not self.firmendaten_eintraege[feld].get():
                messagebox.showerror("Fehler", f"Bitte füllen Sie das Feld '{feld}' aus!")
                return
        
        # Briefkopf-Vorlage erstellen
        doc = docx.Document()
        
        # Logo einfügen, falls vorhanden
        if "logo_pfad" in self.firmendaten and os.path.exists(self.firmendaten["logo_pfad"]):
            try:
                doc.add_picture(self.firmendaten["logo_pfad"], width=Pt(150))
                doc.add_paragraph()  # Abstand nach Logo
            except:
                pass
        
        # Kopfzeile mit Firmendaten
        firmenname = self.firmendaten_eintraege["Firmenname"].get()
        adresse = f"{self.firmendaten_eintraege['Straße'].get()}, {self.firmendaten_eintraege['PLZ'].get()} {self.firmendaten_eintraege['Ort'].get()}"
        
        header = doc.sections[0].header
        p = header.paragraphs[0]
        p.text = f"{firmenname} | {adresse}"
        
        # Platz für Empfänger
        doc.add_paragraph("").add_run("Empfänger:").bold = True
        doc.add_paragraph("Name/Firma")
        doc.add_paragraph("Straße")
        doc.add_paragraph("PLZ Ort")
        doc.add_paragraph("\n\n")
        
        # Platz für Rechnungsinformationen
        doc.add_paragraph().add_run("RECHNUNG").bold = True
        
        table = doc.add_table(rows=4, cols=2)
        table.autofit = True
        
        zellen = table.rows[0].cells
        zellen[0].text = "Rechnungsnummer:"
        zellen[1].text = "<RECHNUNGSNR>"
        
        zellen = table.rows[1].cells
        zellen[0].text = "Datum:"
        zellen[1].text = "<DATUM>"
        
        zellen = table.rows[2].cells
        zellen[0].text = "Kundennummer:"
        zellen[1].text = "<KUNDENNR>"
        
        zellen = table.rows[3].cells
        zellen[0].text = "Leistungszeitraum:"
        zellen[1].text = "<LEISTUNGSZEITRAUM>"
        
        doc.add_paragraph("\n")
        
        # Tabelle für Artikel
        artikel_tabelle = doc.add_table(rows=1, cols=5)
        artikel_tabelle.style = 'Table Grid'
        
        kopfzeile = artikel_tabelle.rows[0].cells
        kopfzeile[0].text = "Bezeichnung"
        kopfzeile[1].text = "Anzahl"
        kopfzeile[2].text = "Netto €"
        kopfzeile[3].text = "USt %"
        kopfzeile[4].text = "Gesamt €"
        
        # Beispielzeile
        row = artikel_tabelle.add_row().cells
        row[0].text = "<ARTIKEL>"
        row[1].text = "<ANZAHL>"
        row[2].text = "<NETTO>"
        row[3].text = "<UST>"
        row[4].text = "<GESAMT>"
        
        doc.add_paragraph("\n")
        
        # Summentabelle
        sum_table = doc.add_table(rows=3, cols=2)
        sum_table.autofit = True
        
        row = sum_table.rows[0].cells
        row[0].text = "Nettobetrag:"
        row[1].text = "<NETTOSUMME> €"
        
        row = sum_table.rows[1].cells
        row[0].text = "Umsatzsteuer:"
        row[1].text = "<USTSUMME> €"
        
        row = sum_table.rows[2].cells
        row[0].text = "Gesamtbetrag:"
        row[1].text = "<BRUTTOSUMME> €"
        
        # Platz für Zahlungsbedingungen
        zahlungsbedingungen = doc.add_paragraph()
        zahlungsbedingungen.add_run("Zahlungsbedingungen: ").bold = True
        zahlungsbedingungen.add_run("Zahlbar innerhalb von <ZAHLUNGSZIEL> Tagen ohne Abzug.")
        
        # Platz für Skonto
        skonto = doc.add_paragraph()
        skonto.add_run("Skonto: ").bold = True
        skonto.add_run("<SKONTO_PROZENT>% bei Zahlung innerhalb von <SKONTO_TAGE> Tagen.")
        
        # Fußzeile mit Bankdaten
        footer = doc.sections[0].footer
        p = footer.paragraphs[0]
        
        # Erweiterte Fußzeile mit mehr Informationen
        footer_text = (
            f"{firmenname} | {adresse} | "
            f"Tel: {self.firmendaten_eintraege['Telefon'].get()} | "
            f"E-Mail: {self.firmendaten_eintraege['Email'].get()} | "
            f"Steuernummer: {self.firmendaten_eintraege['Steuernummer'].get()} | "
            f"USt-ID: {self.firmendaten_eintraege['USt-ID'].get()}"
        )
        
        p.text = footer_text
        
        # Zweite Fußzeile für Bankdaten
        footer.add_paragraph().text = (
            f"Bank: {self.firmendaten_eintraege['Bank'].get()} | "
            f"IBAN: {self.firmendaten_eintraege['IBAN'].get()} | "
            f"BIC: {self.firmendaten_eintraege['BIC'].get()}"
        )
        
        # Speichern
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Dokument", "*.docx")],
            title="Briefkopf speichern"
        )
        
        if filepath:
            doc.save(filepath)
            messagebox.showinfo("Erfolg", f"Briefkopf wurde unter {filepath} gespeichert!")
            self.firmendaten["briefkopf_pfad"] = filepath
            self.speichere_daten()
    
    def lade_briefkopf(self):
        filepath = filedialog.askopenfilename(
            filetypes=[("Word Dokument", "*.docx")],
            title="Briefkopf laden"
        )
        
        if filepath:
            self.firmendaten["briefkopf_pfad"] = filepath
            self.speichere_daten()
            messagebox.showinfo("Erfolg", "Briefkopf wurde ausgewählt!")

    def generiere_rechnungsnummer(self):
        heute = datetime.datetime.now()
        jahr = heute.strftime("%Y")
        monat = heute.strftime("%m")
        
        # Zähler aus Datei lesen oder neu anlegen
        counter_file = f"rechnung_counter_{jahr}.json"
        counter = {}
        
        if os.path.exists(counter_file):
            with open(counter_file, "r") as f:
                counter = json.load(f)
        
        if monat not in counter:
            counter[monat] = 0
        
        counter[monat] += 1
        
        # Zähler speichern
        with open(counter_file, "w") as f:
            json.dump(counter, f)
        
        # Rechnungsnummer formatieren: Jahr-Monat-Zähler
        return f"{jahr}-{monat}-{counter[monat]:03d}"
    
    def aktualisiere_empfaenger_liste(self):
        empfaenger = list(self.adressbuch.keys())
        # Verschiedene Comboboxen in allen Tabs aktualisieren
        if hasattr(self, 'empfaenger_combobox'):
            self.empfaenger_combobox['values'] = empfaenger
        if hasattr(self, 'angebot_empfaenger_combobox'):
            self.angebot_empfaenger_combobox['values'] = empfaenger
        if hasattr(self, 'turnus_empfaenger_combo'):
            self.turnus_empfaenger_combo['values'] = empfaenger
    
    def aktualisiere_adressliste(self, filter_text=None):
        self.adressen_listbox.delete(0, tk.END)
        
        # Wir speichern die Originalnamen in einer separaten Liste
        self.listbox_name_mapping = []
        
        # Sortierte und formatierte Liste erstellen
        kunden_liste = []
        for name, adresse in self.adressbuch.items():
            nachname = adresse.get("Nachname", "")
            vorname = adresse.get("Vorname", "")
            firma = adresse.get("Firma", "")
            display_text = f"{nachname}, {vorname}" if nachname or vorname else name
            if firma:
                display_text += f" - {firma}"
            
            # Bei Suchfilter überprüfen, ob der Text darin vorkommt
            if filter_text is None or filter_text.lower() in display_text.lower() or \
               filter_text.lower() in name.lower():
                kunden_liste.append((nachname, vorname, firma, name, display_text))
        
        # Nach Nachname, Vorname, Firma sortieren
        kunden_liste.sort()
        
        # Anzeigen
        for _, _, _, name, display_text in kunden_liste:
            self.adressen_listbox.insert(tk.END, display_text)
            # Originalnamen in separater Liste speichern
            self.listbox_name_mapping.append(name)
    
    def search_adressbuch(self, *args):
        search_text = self.search_var.get()
        self.aktualisiere_adressliste(search_text)
    
    def adresse_auswaehlen(self, event):
        if self.adressen_listbox.curselection():
            index = self.adressen_listbox.curselection()[0]
            # Namen aus unserer eigenen Mapping-Liste holen
            if 0 <= index < len(self.listbox_name_mapping):
                name = self.listbox_name_mapping[index]
                adresse = self.adressbuch[name]
                
                for feld, var in self.adress_eintraege.items():
                    var.set(adresse.get(feld, ""))
    
    def neue_adresse(self):
        # Felder leeren
        for var in self.adress_eintraege.values():
            var.set("")
    
    def loesche_adresse(self):
        if self.adressen_listbox.curselection():
            index = self.adressen_listbox.curselection()[0]
            # Namen aus unserer eigenen Mapping-Liste holen
            if 0 <= index < len(self.listbox_name_mapping):
                name = self.listbox_name_mapping[index]
                display_text = self.adressen_listbox.get(index)
                
                if messagebox.askyesno("Löschen bestätigen", f"Möchten Sie den Kontakt '{display_text}' wirklich löschen?"):
                    del self.adressbuch[name]
                    self.speichere_daten()
                    self.aktualisiere_adressliste()
                    self.aktualisiere_empfaenger_liste()
    
    def speichere_adresse(self):
        vorname = self.adress_eintraege["Vorname"].get()
        nachname = self.adress_eintraege["Nachname"].get()
        firma = self.adress_eintraege["Firma"].get()
        
        if not (vorname or nachname or firma):
            messagebox.showerror("Fehler", "Bitte geben Sie mindestens Vor-/Nachname oder Firma ein!")
            return
        
        # Name für Dictionary-Key erstellen
        if nachname and vorname:
            name = f"{nachname}, {vorname}"
        elif nachname:
            name = nachname
        elif vorname:
            name = vorname
        else:
            name = firma
        
        adresse = {}
        for feld, var in self.adress_eintraege.items():
            adresse[feld] = var.get()
        
        self.adressbuch[name] = adresse
        self.speichere_daten()
        self.aktualisiere_adressliste()
        self.aktualisiere_empfaenger_liste()
        
        messagebox.showinfo("Erfolg", f"Adresse '{name}' wurde gespeichert!")
    
    def update_leistungszeitraum(self, event=None):
        monat_jahr = self.monat_var.get()
        if monat_jahr and len(monat_jahr) == 7:  # Format: MM.YYYY
            monat, jahr = monat_jahr.split(".")
            monat = int(monat)
            jahr = int(jahr)
            
            # Erster Tag des Monats
            erster_tag = datetime.date(jahr, monat, 1)
            
            # Letzter Tag des Monats
            letzter_tag = erster_tag.replace(day=calendar.monthrange(jahr, monat)[1])
            
            # Datumsformat setzen
            self.leistung_start_var.set(erster_tag.strftime("%d.%m.%Y"))
            self.leistung_ende_var.set(letzter_tag.strftime("%d.%m.%Y"))
    
    def neue_rechnung(self):
        # Artikeleinträge leeren
        for eintrag in self.artikel_eintraege:
            eintrag["bezeichnung"].set("")
            eintrag["menge"].set("1")
            eintrag["preis"].set("0.00")
            eintrag["ust"].set("19")
            eintrag["gesamt_label"].text = "0.00"
        
        # Rechnungsdaten zurücksetzen
        self.rechnungsnr_var.set(self.generiere_rechnungsnummer())
        self.rechnungsdatum_var.set(datetime.datetime.now().strftime("%d.%m.%Y"))
        self.leistung_start_var.set("")
        self.leistung_ende_var.set("")
        self.steuerschuldnerschaft_var.set(False)
        self.storno_var.set(False)
        self.storno_rechnungsnr_var.set("")
        
        # Zahlungsbedingungen zurücksetzen
        self.zahlungsziel_var.set("14")
        self.skonto_prozent_var.set("0")
        self.skonto_tage_var.set("0")
        
        # Empfänger leeren
        self.empfaenger_var.set("")
        
        # Vorschau leeren
        self.vorschau_text.config(state=tk.NORMAL)
        self.vorschau_text.delete(1.0, tk.END)
        self.vorschau_text.config(state=tk.DISABLED)
    
    def berechne_summen(self):
        empfaenger = self.empfaenger_var.get()
        if not empfaenger or empfaenger not in self.adressbuch:
            messagebox.showerror("Fehler", "Bitte wählen Sie einen gültigen Empfänger!")
            return
        
        artikel = []
        nettosumme = Decimal('0.00')
        ust_betraege = {}  # Um Steuerbeträge nach Steuersätzen zu gruppieren
        
        ist_leistung = self.rechnungstyp_var.get() == "Leistung"
        menge_label = "Stunden" if ist_leistung else "Anzahl"
        preis_label = "Stundensatz" if ist_leistung else "Netto-Preis"
        
        for eintrag in self.artikel_eintraege:
            bezeichnung = eintrag["bezeichnung"].get()
            
            # Nur Artikel mit Bezeichnung berücksichtigen
            if bezeichnung:
                try:
                    menge = Decimal(eintrag["menge"].get().replace(',', '.'))
                    preis = Decimal(eintrag["preis"].get().replace(',', '.'))
                    ust_satz = Decimal(eintrag["ust"].get().replace(',', '.'))
                    
                    position_netto = menge * preis
                    
                    # Bei Steuerschuldnerschaft ist der USt-Betrag 0
                    if self.steuerschuldnerschaft_var.get() and ust_satz > 0:
                        ust_betrag = Decimal('0.00')
                    else:
                        ust_betrag = position_netto * (ust_satz / 100)
                    
                    position_brutto = position_netto + ust_betrag
                    
                    # Gesamtpreis in Label aktualisieren
                    eintrag["gesamt_label"].config(text=f"{position_brutto:.2f}")
                    
                    artikel.append({
                        "bezeichnung": bezeichnung,
                        "menge": menge,
                        "preis": preis,
                        "ust_satz": ust_satz,
                        "ust_betrag": ust_betrag,
                        "gesamt": position_brutto,
                        "menge_label": menge_label,
                        "preis_label": preis_label
                    })
                    
                    nettosumme += position_netto
                    
                    # Steuerbeträge nach Sätzen gruppieren
                    ust_key = str(ust_satz)
                    if ust_key not in ust_betraege:
                        ust_betraege[ust_key] = Decimal('0.00')
                    ust_betraege[ust_key] += ust_betrag
                    
                except (ValueError, Exception) as e:
                    messagebox.showerror("Fehler", f"Ungültige Zahl in Artikel '{bezeichnung}'! {str(e)}")
                    return
        
        if not artikel:
            messagebox.showerror("Fehler", "Bitte geben Sie mindestens einen Artikel ein!")
            return
        
        # Gesamtbeträge berechnen
        gesamt_ust = sum(ust_betraege.values())
        bruttosumme = nettosumme + gesamt_ust
        
        # Skonto berechnen
        skonto_prozent = Decimal('0.00')
        try:
            skonto_prozent = Decimal(self.skonto_prozent_var.get().replace(',', '.'))
        except:
            pass
        
        skonto_betrag = bruttosumme * (skonto_prozent / 100)
        skonto_summe = bruttosumme - skonto_betrag
        
        # Vorschau anzeigen
        self.vorschau_text.config(state=tk.NORMAL)
        self.vorschau_text.delete(1.0, tk.END)
        
        self.vorschau_text.insert(tk.END, f"Empfänger: {empfaenger}\n\n")
        self.vorschau_text.insert(tk.END, f"Artikel: {len(artikel)}\n")
        
        # Steueraufstellung
        self.vorschau_text.insert(tk.END, "\nSteueraufstellung:\n")
        for satz, betrag in ust_betraege.items():
            self.vorschau_text.insert(tk.END, f"  {satz}% USt: {betrag:.2f} €\n")
        
        # Summen
        self.vorschau_text.insert(tk.END, f"\nNettosumme: {nettosumme:.2f} €\n")
        self.vorschau_text.insert(tk.END, f"USt gesamt: {gesamt_ust:.2f} €\n")
        self.vorschau_text.insert(tk.END, f"Bruttosumme: {bruttosumme:.2f} €\n")
        
        # Skonto
        if skonto_prozent > 0:
            tage = self.skonto_tage_var.get()
            self.vorschau_text.insert(tk.END, f"\nSkonto ({skonto_prozent}% bei Zahlung innerhalb von {tage} Tagen): {skonto_betrag:.2f} €\n")
            self.vorschau_text.insert(tk.END, f"Betrag nach Skonto: {skonto_summe:.2f} €\n")
        
        # Steuerschuldnerschaft
        if self.steuerschuldnerschaft_var.get():
            self.vorschau_text.insert(tk.END, "\nHinweis: Steuerschuldnerschaft des Leistungsempfängers\n")
        
        # Storno
        if self.storno_var.get():
            original_nr = self.storno_rechnungsnr_var.get()
            if original_nr:
                self.vorschau_text.insert(tk.END, f"\nKorrektur der Rechnung {original_nr}. Diese wird storniert.\n")
            else:
                self.vorschau_text.insert(tk.END, "\nFehler: Bitte geben Sie die Original-Rechnungsnummer für den Storno an.\n")
        
        self.vorschau_text.config(state=tk.DISABLED)
        
        return {
            "empfaenger": empfaenger,
            "artikel": artikel,
            "nettosumme": nettosumme,
            "ust_betraege": ust_betraege,
            "bruttosumme": bruttosumme,
            "skonto_prozent": skonto_prozent,
            "skonto_betrag": skonto_betrag,
            "skonto_summe": skonto_summe
        }
    
    def erstelle_rechnung(self, zum_versenden=False):
        # Prüfen, ob ein Briefkopf existiert
        if not self.firmendaten.get("briefkopf_pfad") or not os.path.exists(self.firmendaten["briefkopf_pfad"]):
            messagebox.showerror("Fehler", "Bitte erstellen oder laden Sie zuerst einen Briefkopf!")
            return None
        
        # Empfänger prüfen
        empfaenger_name = self.empfaenger_var.get()
        if not empfaenger_name or empfaenger_name not in self.adressbuch:
            messagebox.showerror("Fehler", "Bitte wählen Sie einen gültigen Empfänger!")
            return None
        
        empfaenger = self.adressbuch[empfaenger_name]
        
        # Summen berechnen
        summen = self.berechne_summen()
        if not summen:
            return None
        
        nettosumme = summen["nettosumme"]
        ust_betraege = summen["ust_betraege"]
        bruttosumme = summen["bruttosumme"]
        artikel = summen["artikel"]
        skonto_prozent = summen["skonto_prozent"]
        skonto_betrag = summen["skonto_betrag"]
        skonto_summe = summen["skonto_summe"]
        
        # Leistungszeitraum
        leistungszeitraum = ""
        if self.leistung_start_var.get() and self.leistung_ende_var.get():
            leistungszeitraum = f"{self.leistung_start_var.get()} - {self.leistung_ende_var.get()}"
        
        # Rechnungsdokument erstellen
        try:
            # Vorlage kopieren
            doc = docx.Document(self.firmendaten["briefkopf_pfad"])
            
            # Storno-Hinweis
            if self.storno_var.get():
                original_nr = self.storno_rechnungsnr_var.get()
                if not original_nr:
                    messagebox.showerror("Fehler", "Bitte geben Sie die Original-Rechnungsnummer für den Storno an!")
                    return None
                
                # Storno-Hinweis als ersten Absatz hinzufügen
                p = doc.paragraphs[0]
                p.add_run(f"\nKORREKTUR/STORNO der Rechnung {original_nr}").bold = True
            
            # Empfängerdaten eintragen
            for i, para in enumerate(doc.paragraphs):
                if "Empfänger:" in para.text:
                    # Die nächsten Paragraphen mit Empfängerdaten überschreiben
                    empfaenger_zeilen = []
                    
                    # Name oder Firma
                    if empfaenger.get("Firma"):
                        empfaenger_zeilen.append(empfaenger["Firma"])
                        if empfaenger.get("Vorname") and empfaenger.get("Nachname"):
                            empfaenger_zeilen.append(f"{empfaenger['Vorname']} {empfaenger['Nachname']}")
                        elif empfaenger.get("Vorname"):
                            empfaenger_zeilen.append(empfaenger["Vorname"])
                        elif empfaenger.get("Nachname"):
                            empfaenger_zeilen.append(empfaenger["Nachname"])
                    else:
                        name_str = ""
                        if empfaenger.get("Vorname"):
                            name_str += empfaenger["Vorname"] + " "
                        if empfaenger.get("Nachname"):
                            name_str += empfaenger["Nachname"]
                        empfaenger_zeilen.append(name_str.strip())
                    
                    # Adresse
                    empfaenger_zeilen.append(empfaenger.get("Straße", ""))
                    empfaenger_zeilen.append(f"{empfaenger.get('PLZ', '')} {empfaenger.get('Ort', '')}")
                    
                    # In Dokument schreiben
                    for j, zeile in enumerate(empfaenger_zeilen):
                        if i+j+1 < len(doc.paragraphs):
                            doc.paragraphs[i+j+1].text = zeile
                    break
            
            # Rechnungsinformationen
            rechnungsnr = self.rechnungsnr_var.get()
            datum = self.rechnungsdatum_var.get()
            
            for table in doc.tables:
                if len(table.rows) >= 3 and "Rechnungsnummer:" in table.rows[0].cells[0].text:
                    table.rows[0].cells[1].text = rechnungsnr
                    table.rows[1].cells[1].text = datum
                    
                    # Debitorennummer als Kundennummer verwenden, falls vorhanden
                    if empfaenger.get("Debitorennummer"):
                        kundennr = empfaenger["Debitorennummer"]
                    else:
                        # Sonst eine einfache Kundennummer generieren
                        kundennr = f"{empfaenger_name[0].upper()}{hash(empfaenger_name) % 10000:04d}"
                    
                    table.rows[2].cells[1].text = kundennr
                    
                    # Leistungszeitraum
                    if len(table.rows) >= 4 and "Leistungszeitraum:" in table.rows[3].cells[0].text:
                        table.rows[3].cells[1].text = leistungszeitraum
                    
                    break
            
            # Artikel in Tabelle einfügen
            artikel_tabelle = None
            for table in doc.tables:
                if len(table.rows) >= 1 and len(table.rows[0].cells) >= 5:
                    if "Bezeichnung" in table.rows[0].cells[0].text and "Anzahl" in table.rows[0].cells[1].text:
                        artikel_tabelle = table
                        break
            
            if artikel_tabelle:
                # Bei Leistungen Header anpassen
                ist_leistung = self.rechnungstyp_var.get() == "Leistung"
                if ist_leistung and "Anzahl" in artikel_tabelle.rows[0].cells[1].text:
                    artikel_tabelle.rows[0].cells[1].text = "Stunden"
                    artikel_tabelle.rows[0].cells[2].text = "Stundensatz €"
                
                # Beispielzeile entfernen, falls vorhanden
                if len(artikel_tabelle.rows) > 1:
                    if "<ARTIKEL>" in artikel_tabelle.rows[1].cells[0].text:
                        artikel_tabelle._tbl.remove(artikel_tabelle.rows[1]._tr)
                
                # Artikel hinzufügen
                for artikel_item in artikel:
                    row = artikel_tabelle.add_row().cells
                    row[0].text = artikel_item["bezeichnung"]
                    row[1].text = f"{artikel_item['menge']}"
                    row[2].text = f"{artikel_item['preis']:.2f}"
                    
                    # Bei Steuerschuldnerschaft anpassen
                    if self.steuerschuldnerschaft_var.get() and artikel_item["ust_satz"] > 0:
                        row[3].text = f"{artikel_item['ust_satz']} %*"
                    else:
                        row[3].text = f"{artikel_item['ust_satz']} %"
                    
                    row[4].text = f"{artikel_item['gesamt']:.2f}"
            
            # Summentabelle finden und aktualisieren
            sum_table = None
            for table in doc.tables:
                if len(table.rows) >= 3 and "Nettobetrag:" in table.rows[0].cells[0].text:
                    sum_table = table
                    break
            
            if sum_table:
                sum_table.rows[0].cells[1].text = f"{nettosumme:.2f} €"
                
                # Bei mehreren Steuersätzen detailliert aufschlüsseln
                if len(ust_betraege) > 1:
                    # Alle vorhandenen Zeilen ab der zweiten löschen (Umsatzsteuer, Gesamtbetrag)
                    while len(sum_table.rows) > 1:
                        sum_table._tbl.remove(sum_table.rows[-1]._tr)
                    
                    # Zeilen für jeden Steuersatz einfügen
                    for satz, betrag in sorted(ust_betraege.items()):
                        row_cells = sum_table.add_row().cells
                        row_cells[0].text = f"USt {satz} %:"
                        row_cells[1].text = f"{betrag:.2f} €"
                    
                    # Zeile für Gesamtbetrag hinzufügen
                    row_cells = sum_table.add_row().cells
                    row_cells[0].text = "Gesamtbetrag:"
                    row_cells[1].text = f"{bruttosumme:.2f} €"
                    
                    # Zeile für Skonto hinzufügen, falls vorhanden
                    if skonto_prozent > 0:
                        skonto_tage = self.skonto_tage_var.get()
                        
                        row_cells = sum_table.add_row().cells
                        row_cells[0].text = f"Skonto {skonto_prozent}% bei Zahlung innerhalb von {skonto_tage} Tagen:"
                        row_cells[1].text = f"{skonto_betrag:.2f} €"
                        
                        row_cells = sum_table.add_row().cells
                        row_cells[0].text = "Betrag nach Skonto:"
                        row_cells[1].text = f"{skonto_summe:.2f} €"
                else:
                    # Standard mit einem Steuersatz
                    # Sicherstellen, dass genau 3 Zeilen existieren
                    while len(sum_table.rows) > 3:
                        sum_table._tbl.remove(sum_table.rows[-1]._tr)
                    
                    while len(sum_table.rows) < 3:
                        sum_table.add_row()
                    
                    # Werte setzen
                    sum_table.rows[1].cells[0].text = "Umsatzsteuer:"
                    sum_table.rows[1].cells[1].text = f"{gesamt_ust:.2f} €"
                    sum_table.rows[2].cells[0].text = "Gesamtbetrag:"
                    sum_table.rows[2].cells[1].text = f"{bruttosumme:.2f} €"
                    
                    # Zeilen für Skonto hinzufügen, falls vorhanden
                    if skonto_prozent > 0:
                        skonto_tage = self.skonto_tage_var.get()
                        
                        row_cells = sum_table.add_row().cells
                        row_cells[0].text = f"Skonto {skonto_prozent}% bei Zahlung innerhalb von {skonto_tage} Tagen:"
                        row_cells[1].text = f"{skonto_betrag:.2f} €"
                        
                        row_cells = sum_table.add_row().cells
                        row_cells[0].text = "Betrag nach Skonto:"
                        row_cells[1].text = f"{skonto_summe:.2f} €"
            
            # Zahlungsbedingungen aktualisieren
            zahlungsziel = self.zahlungsziel_var.get()
            zahlungsbedingungen_text = f"Zahlbar innerhalb von {zahlungsziel} Tagen ohne Abzug."
            
            for para in doc.paragraphs:
                if "Zahlungsbedingungen:" in para.text:
                    # Text nach dem Doppelpunkt durch neuen Text ersetzen
                    text_teile = para.text.split(":")
                    if len(text_teile) > 1:
                        para.text = f"{text_teile[0]}: {zahlungsbedingungen_text}"
                    break
            
            # Skonto-Text aktualisieren
            skonto_text = ""
            if skonto_prozent > 0:
                skonto_tage = self.skonto_tage_var.get()
                skonto_text = f"{skonto_prozent}% bei Zahlung innerhalb von {skonto_tage} Tagen."
            
            for para in doc.paragraphs:
                if "Skonto:" in para.text:
                    # Text nach dem Doppelpunkt durch neuen Text ersetzen
                    text_teile = para.text.split(":")
                    if len(text_teile) > 1:
                        para.text = f"{text_teile[0]}: {skonto_text}"
                    break
            
            # Steuerschuldnerschaft-Hinweis hinzufügen
            if self.steuerschuldnerschaft_var.get():
                p = doc.add_paragraph()
                p.add_run("* Steuerschuldnerschaft des Leistungsempfängers").italic = True
                
                p = doc.add_paragraph()
                p.add_run("Gemäß § 13b UStG ist der Leistungsempfänger Steuerschuldner und hat die Umsatzsteuer selbst zu berechnen und an das Finanzamt abzuführen.").italic = True
            
            # Speichern
            if zum_versenden:
                # Temporäres Speichern für E-Mail-Versand
                temp_dir = os.environ.get("TEMP", ".")
                if not os.path.exists(temp_dir):
                    temp_dir = "."
                
                temp_docx = os.path.join(temp_dir, f"Rechnung_{rechnungsnr}.docx")
                doc.save(temp_docx)
                
                # Rechnung in Datenbank speichern
                self.speichere_rechnung_in_datenbank(rechnungsnr, empfaenger_name)
                
                return {"docx": temp_docx, "rechnungsnr": rechnungsnr, "empfaenger": empfaenger}
            else:
                filepath = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Dokument", "*.docx")],
                    initialfile=f"Rechnung_{rechnungsnr}.docx",
                    title="Rechnung speichern"
                )
                
                if filepath:
                    doc.save(filepath)
                    messagebox.showinfo("Erfolg", f"Rechnung wurde unter {filepath} gespeichert!")
                    
                    # Rechnung in Datenbank speichern
                    self.speichere_rechnung_in_datenbank(rechnungsnr, empfaenger_name)
                    
                    # Rechnungsnummer automatisch erhöhen für die nächste Rechnung
                    self.rechnungsnr_var.set(self.generiere_rechnungsnummer())
                return filepath
        
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Erstellen der Rechnung: {str(e)}")
            return None
    
    def erstelle_rechnung_und_versende(self):
        # Rechnung erstellen
        ergebnis = self.erstelle_rechnung(zum_versenden=True)
        
        if ergebnis:
            docx_pfad = ergebnis["docx"]
            rechnungsnr = ergebnis["rechnungsnr"]
            empfaenger = ergebnis["empfaenger"]
            
            email = empfaenger.get("Email", "")
            
            # E-Mail-Betreff
            betreff = f"Rechnung {rechnungsnr}"
            
            # E-Mail-Text
            email_text = (
                "Sehr geehrte Damen und Herren,\n\n"
                "hiermit erlauben wir uns anbei die Rechnung für unsere Leistungen zu stellen. "
                "Bei Rückfragen stehen wir gerne zur Verfügung."
            )
            
            # Zusätzlich X-Rechnung erstellen, falls gewünscht
            xrechnung_pfad = self.erstelle_xrechnung(ergebnis)
            
            # Mail-Client öffnen
            try:
                attachments = f'"{docx_pfad}"'
                if xrechnung_pfad:
                    attachments += f' "{xrechnung_pfad}"'
                
                if sys.platform == "win32":
                    # Windows mit Outlook
                    os.system(f'start outlook.exe /c ipm.note /m "{email}" /a {attachments} /s /t "{betreff}" {email_text}')
                else:
                    # Andere Betriebssysteme - Standard-Mail-Client
                    import webbrowser
                    mailto_link = f"mailto:{email}?subject={betreff}&body={email_text}"
                    webbrowser.open(mailto_link)
                    messagebox.showinfo("Hinweis", 
                                       f"E-Mail-Client wurde geöffnet. Bitte hängen Sie die Rechnung manuell an:\n{docx_pfad}")
                    if xrechnung_pfad:
                        messagebox.showinfo("Hinweis", 
                                           f"Bitte hängen Sie auch die X-Rechnung an:\n{xrechnung_pfad}")
            
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Öffnen des E-Mail-Clients: {str(e)}")

def main():
    root = tk.Tk()
    app = RechnungsProgramm(root)
    root.mainloop()

if __name__ == "__main__":
    main()







