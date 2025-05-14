import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import json # JSON para datos de entrada EXTERNOS (causas y modelos)

# Cargar datos desde archivos JSON (dato entrada 1)
with open('causas_EM.json', 'r', encoding='utf-8') as f:
    causas = json.load(f)

with open('modelos_EM.json', 'r', encoding='utf-8') as f:
    modelos = json.load(f)

# Crear ventana principal (dato entrada)
root = tk.Tk()
root.title("Escritos Multiples - Proyecto - Franco")
root.geometry("1000x800")

# Crear frame principal con scrollbar
main_frame = tk.Frame(root)
main_frame.pack(fill=tk.BOTH, expand=True)

# Crear canvas
canvas = tk.Canvas(main_frame)
canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Añadir scrollbar vertical
scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=canvas.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# Configurar canvas
canvas.configure(yscrollcommand=scrollbar.set)
canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

# Crear otro frame dentro del canvas
second_frame = tk.Frame(canvas)
canvas.create_window((0, 0), window=second_frame, anchor="nw")

# Encabezados
tk.Label(second_frame, text="Nombre de la causa").grid(row=0, column=0, padx=10, pady=10)
tk.Label(second_frame, text="Modelo de escrito").grid(row=0, column=1, padx=10, pady=10)
tk.Label(second_frame, text="Seleccionar").grid(row=0, column=2, padx=10, pady=10)
tk.Label(second_frame, text="Sumar dato").grid(row=0, column=3, padx=10, pady=10)

# Guardar referencias
comboboxes = []
check_vars = []
sumar_dato_vars = []

# Crear filas por causa
for i, causa in enumerate(causas):
    #tk.Label(second_frame, text=causa["nombre"], width=70).grid(row=i+1, column=0, padx=1, pady=5, sticky="w")
    tk.Label(second_frame, text=causa["nombre"], width=100, wraplength=700, justify="left").grid(row=i+1, column=0, padx=1, pady=5, sticky="w")
    combo = ttk.Combobox(
        second_frame, 
        values=list(modelos.keys()), 
        state="readonly",
        width=60 # <-- ancho menú desplegable
    )
    combo.grid(row=i+1, column=1, padx=5, pady=5)
    comboboxes.append(combo)
    
    var_select = tk.BooleanVar()
    chk_select = tk.Checkbutton(second_frame, variable=var_select)
    chk_select.grid(row=i+1, column=2)
    check_vars.append(var_select)
    
    var_sumar = tk.BooleanVar()
    chk_sumar = tk.Checkbutton(second_frame, variable=var_sumar)
    chk_sumar.grid(row=i+1, column=3)
    sumar_dato_vars.append(var_sumar)
# Función para generar Word
def generar_word():
    generados = []

    for idx, (causa, combo, var, sumar_var) in enumerate(zip(causas, comboboxes, check_vars, sumar_dato_vars)):
        if var.get():
            modelo_titulo = combo.get()
            if modelo_titulo not in modelos:
                continue

            texto_modelo = modelos[modelo_titulo]
            patrocinante = causa["patrocinio"].lower()

            # Crear documento 
            doc = Document()

            # Márgenes word
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(5)
                section.left_margin = Cm(5)
                section.bottom_margin = Cm(2)
                section.right_margin = Cm(2)

            # Estilo -general-
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Título del modelo
            p = doc.add_paragraph()
            run = p.add_run(modelo_titulo)
            run.bold = True
            run.underline = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Sr. Juez
            doc.add_paragraph("Sr. Juez:")
            doc.add_paragraph("")

            # patrocinio (AGUS) (EUGE CABA) (EUGE PBA)
            if patrocinante.lower() == "maria agustina raffo, abogada, inscripta en el t°119 f°265 del c.p.a.c.f, cuit 27-32475742-8, responsable inscripta":
                    abogada = "MARIA AGUSTINA RAFFO, abogada, inscripta en el T°119 F°265 del C.P.A.C.F, CUIT 27-32475742-8, Responsable Inscripta"
            elif patrocinante.lower() == "maria eugenia lupi serres, abogada, inscripta en el t°118 f°474 del c.p.a.c.f, cuit 27-33538887-4, responsable inscripta":
                    abogada = "MARIA EUGENIA LUPI SERRES, abogada, inscripta en el T°118 F°474 del C.P.A.C.F, CUIT 27-33538887-4, Responsable Inscripta"
            elif patrocinante.lower() == "maria eugenia lupi serres, abogada, inscripta en el t°xi f°214 c.a.q., cuit 27-33538887-4, responsable inscripta":
                    abogada = "MARIA EUGENIA LUPI SERRES, abogada, inscripta en el T°XI F°214 C.A.Q., CUIT 27-33538887-4, Responsable Inscripta"
            else:
                    abogada = "(no se detectó abogado cargado en Lex)"

            parrafo_presentacion = doc.add_paragraph()
            run = parrafo_presentacion.add_run(f"{abogada}")
            run.bold = True
            parrafo_presentacion.add_run(", en autos caratulados ")
            run = parrafo_presentacion.add_run(causa["nombre"])
            run.bold = True
            parrafo_presentacion.add_run(", a VS respetuosamente digo:")
            parrafo_presentacion.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p = doc.add_paragraph(texto_modelo)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            doc.add_paragraph("")

            # Final
            p = doc.add_paragraph()
            run = p.add_run("Proveer de conformidad,")
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            p2 = doc.add_paragraph()
            run2 = p2.add_run("SERÁ JUSTICIA.")
            run2.bold = True
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Guardar documento
            nombre_archivo = f"{idx + 1} - {modelo_titulo}"
            if sumar_var.get():
                nombre_archivo += " - AGREGAR DATOS"
            nombre_archivo += ".docx"

            doc.save(nombre_archivo)
            generados.append(nombre_archivo)

    if generados:
        mensaje = "Se han generado los siguientes Word:\n\n" + "\n".join(generados)
        messagebox.showinfo("Documentos generados", mensaje)
    else:
        messagebox.showwarning("Sin selección", "No se seleccionó ninguna causa para generar escritos.")

# Botón de generar
tk.Button(second_frame, text="Generar escritos", command=generar_word, bg="lightblue").grid(row=len(causas)+2, column=1, columnspan=2, pady=20)

root.mainloop()