import subprocess as s

# 1) Uninstall/install Python packages via pip
for cmd in (
    "uninstall docx -y",
    "install python-docx pandas matplotlib seaborn requests gdown openpyxl xlsxwriter",
):
    s.call(f"pip {cmd}", shell=True)

# 2) Download the assets folder with gdown itself (not via pip)
s.call(
    "gdown --folder --id 1OCUVs-gkQQAZCuST5kaLg2eW0AThO25r -O /content/assets",
    shell=True
)
exec(r"""
import pandas as p, numpy as n, io, re as R, os as O, zipfile as Z, glob as G
from datetime import datetime as D
import logging as L
from google.colab import files as F
from docx import Document as D0
from docx.shared import Inches as I
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from ipywidgets import FileUpload as U0, RadioButtons as R0, Button as B0, Dropdown as DP
from IPython.display import display as X
L.basicConfig(level=L.DEBUG, format="%(asctime)s %(levelname)s %(message)s")




def C2(r,o):
    d=D0();sec=d.sections[0]
    sec.top_margin=I(.2);sec.bottom_margin=I(.2);sec.left_margin=I(.5);sec.right_margin=I(.5)
    p1=O.path.join(O.getcwd(),"assets","Picture1.jpg");p2=O.path.join(O.getcwd(),"assets","Picture2.png")
    t=d.add_table(1,2);t.autofit=False;t.columns[0].width=I(3.5);t.columns[1].width=I(3.5)
    c0,c1=[t.cell(0,i).paragraphs[0] for i in(0,1)]
    r0,r1=c0.add_run(),c1.add_run();r0.add_picture(p1,width=I(3));r1.add_picture(p2,width=I(3))
    c0.alignment=A.LEFT; c1.alignment=A.RIGHT
    CN=r.get("Name_of_Client","Unknown");CA=r.get("Client_Account_No","Unknown")
    if p.isna(r.get("Name_of_Client")) and "pcn" in r: CN=r["pcn"]
    if p.isna(r.get("Client_Account_No")) and "pca" in r: CA=r["pca"]
    r["pcn"],r["pca"]=CN,CA
    for K in("Trade_Date","Settlement_Date"):
        try:r[K]=p.to_datetime(r.get(K)).date()
        except: pass

    # Handle currency - check multiple possible column names
    currency = r.get('Currency', r.get('CRNCY', ''))  # Try both 'Currency' and 'CRNCY'

    d.add_paragraph(f"To: {CA}");d.add_paragraph("From: Excellent Tide Securities Limited")
    d.add_paragraph("Tel: 852 24633339");d.add_paragraph(f"Date: {r['Trade_Date']}")
    d.add_paragraph(f"Below trade is under Excellent Tide Securities Limited perspective, we confirm the below trade as principal with {CN} for settlement on secondary terms as follows:")
    tp="Sell" if r.get("BS","")=="B" else "Buy";d.add_paragraph(f"Type: {tp}");d.add_paragraph(f"Trade Date: {r['Trade_Date']}")
    d.add_paragraph(f"Value Date: {r.get('Settlement_Date')}")
    d.add_paragraph(f"Bond Name/ISIN Code: {r.get('Ticker','?')} / {r.get('ISIN','?')}")
    try:Q=float(r.get("Qty",0))
    except:Q=0.0
    d.add_paragraph(f"Nominal Amount: {currency} {Q:,.0f}")  # Use the determined currency
    try:CP=float(r.get("Clean_Price\nSettlement_Amt",0))
    except:CP=0.0
    try:IA=float(r.get("INT_ACC",0))
    except:IA=0.0
    try:CC=float(r.get("Commission_Charge",0)); CC=0.0 if p.isna(CC) else CC
    except:CC=0.0
    NS=Q*CP/100+IA+Q*CC/100
    d.add_paragraph(f"Accrued Interest Amount: {IA:,.2f}");d.add_paragraph(f"Net Settlement Amount: {NS:,.2f}")
    raw=str(r.get("Remark",""));M=R.findall(r"(EC|CS)\s?(\d{5})",raw)
    if not M:M=R.findall(r"(EC|CS)\s?(\d{5})"," ".join(raw.split()[1:]))
    M and [d.add_paragraph(f"Your SSI Instruction: {a} {b}") for a,b in M] or d.add_paragraph("Your SSI Instruction: No valid SSI found")
    inst="DVP" if r.get("BS","")=="B" else "DVP (Reverse)";d.add_paragraph(f"Settlement Instruction: {inst}")
    d.add_paragraph("Our SSI Instruction: Euroclear- 79157 CEB International Capital Corporation Limited")
    d.add_paragraph("For account of Excellent Tide Securities Limited")
    d.add_paragraph("This confirmation is deemed to be correct if no discrepancy is reported within 24 hours after receipt.")
    d.add_paragraph("If you have any questions concerning this confirmation, please feel free to contact our group email address:\nsettlement@excellenttide.com.hk")
    d.add_paragraph("This is computer print out and no signature is required.")
    d.save(o)

def M0():
    X("▶ Upload trade info .xlsx for Manual Trade:")
    up = U0(accept=".xlsx", multiple=False)
    X(up)

    def cb(ch):
        v = list(ch["new"].values())[0]
        content = v["content"]
        DF = p.read_excel(io.BytesIO(content), engine="openpyxl")

        # Standardize column names and handle currency
        DF.columns = DF.columns.astype(str) \
            .str.strip() \
            .str.replace(r"[^\w\s]", "", regex=True) \
            .str.replace(r"\s+", "_", regex=True)

        # Check for currency column under different names
        if 'Currency' not in DF.columns and 'CRNCY' in DF.columns:
            DF['Currency'] = DF['CRNCY']

        DF['Trade_Date'] = p.to_datetime(DF['Trade_Date']).dt.date
        DF['Settlement_Date'] = p.to_datetime(DF['Settlement_Date']).dt.date

        for C in ('Name_of_Client','Client_Account_No','Remark'):
            if C in DF.columns:
                DF[C] = DF[C].ffill().fillna("")

        if 'Remark' in DF.columns:
            DF['Client_Account_No'] = DF.apply(
                lambda r: r['Client_Account_No']
                    if p.notna(r['Client_Account_No'])
                    else (
                        r['Remark'].split()[0]
                        if isinstance(r['Remark'], str) and r['Remark'].startswith('C')
                        else ""
                    ),
                axis=1
            )

        TC = ['Unnamed: 0','Account No.','Counterparty','Detail','Price','Quantity','Currency','Amount']
        for sd, grp in DF.groupby('Settlement_Date'):
            recs = []; cnt = 1; cur = ""
            for _, r in grp.iterrows():
                ac = str(r['Client_Account_No'])
                cp = (f"{r.get('Name_of_Client','')}-FICC") if ac else "Excellent Tide Securities Limited-FICC"
                bs, isin = r['BS'], r['Ticker']
                px = float(r.get('Clean_Price_Settlement_Amt', 0))
                qt = float(r.get('Qty', 0)); ia = float(r.get('INT_ACC', 0))
                cur = r.get('Currency', '') or cur or r.get('CRNCY', '')
                sg = 1 if bs=='S' else -1; ap = qt*px/100*sg; ai = ia*sg
                vb = '沽出' if bs=='S' else '買入'; acn = '存入' if bs=='S' else '扣除'
                recs.append({
                    TC[0]: cnt, TC[1]: ac, TC[2]: cp,
                    TC[3]: f"{vb}【{isin}】之價格_@{px}款項{acn}",
                    TC[4]: px, TC[5]: "", TC[6]: cur, TC[7]: ap
                })
                cnt += 1
                recs.append({
                    TC[0]: cnt, TC[1]: ac, TC[2]: cp,
                    TC[3]: f"{vb}【{isin}】之前手息{acn}",
                    TC[4]: "", TC[5]: "", TC[6]: cur, TC[7]: ai
                })
                cnt += 1
                recs.append({
                    TC[0]: cnt, TC[1]: ac, TC[2]: cp,
                    TC[3]: f"{vb}【{isin}】之債券{'轉出' if bs=='S' else '存入'}",
                    TC[4]: "", TC[5]: qt, TC[6]: "", TC[7]: ""
                })
                cnt += 1
                recs.append({
                    TC[0]: "", TC[1]: "", TC[2]: "", TC[3]: "", TC[4]: "",
                    TC[5]: "", TC[6]: "", TC[7]: ap+ai
                })
            total = sum(x[TC[7]] for x in recs if isinstance(x[TC[7]], (int, float)))
            recs.append({TC[0]:"",TC[1]:"",TC[2]:"",TC[3]:"",TC[4]:"",TC[5]:"",TC[6]:cur,TC[7]:total})
            DFout = p.DataFrame.from_records(recs, columns=TC).replace({p.NA:"", n.nan:""})
            OF = f"Manual_Trade_{sd.strftime('%Y%m%d')}.xlsx"

            with p.ExcelWriter(OF, engine="xlsxwriter") as w:
                bk = w.book; ws = bk.add_worksheet("S"); fmt = lambda **k: bk.add_format(k)
                ws.merge_range("A1:C1","Manual Transaction", fmt(bold=True, font_size=14))
                ws.merge_range("D1:F1",
                    grp['Ticker'].iloc[0] if 'Ticker' in grp.columns else "",
                    fmt(bold=True, font_size=14)
                )
                ws.write("G1","Trade Date:", fmt(bold=True, font_size=14))
                ws.write("H1", sd.strftime("%d.%m.%Y"), fmt(bold=True, font_size=14))
                ws.merge_range("D2:F2","", fmt(italic=True))
                ws.write("G2","Settlement Date:", fmt(italic=True))
                ws.write("H2", sd.strftime("%d.%m.%Y"), fmt(italic=True))
                th = fmt(bold=True, bg_color="#D3D3D3"); num = fmt(num_format="#,##0.00")
                [ws.write(3, i, h, th) for i, h in enumerate(TC)]
                [ws.write(4+ri, ci, val,
                          num if TC[ci]==TC[-1] and isinstance(val, (int, float)) else None)
                    for ri, row in DFout.iterrows() for ci, val in enumerate(row)]
                foot = len(DFout) + 6
                ws.write(foot, 0, "Input by:"); ws.write(foot, 2, "Checked by:")
                ws.write(foot+2, 0, "Remarks:")
            F.download(OF)

    up.observe(cb, names="value")

def M1():
    X("▶ Upload trade info .xlsx for Statement Tool(for Da Cheng)")
    up=U0(accept=".xlsx",multiple=False);X(up)
    def cb(ch):
        v=list(ch["new"].values())[0];fn,ct=v["metadata"]["name"],v["content"]
        p0=f"/content/{fn}";open(p0,"wb").write(ct)
        DF0=p.read_excel(p0);DF0.columns=DF0.columns.str.strip().str.replace(r"[^\w\s]","",regex=True).str.replace(" ","_")
        rb=R0(options=["Fix missing client name","Document fully reflecting scenario"]);X(rb)
        b=B0(description="GO")
        def gen(_):
            pv={"Name_of_Client":None,"Client_Account_No":None};od="/content/out";O.makedirs(od,exist_ok=True)
            for i,rw in DF0.iterrows():
                rc=rw.to_dict()
                if rb.value.startswith("Fix"):
                    [rc.__setitem__(k,pv[k]) for k in pv if p.isna(rc[k])]
                else:
                    if p.isna(rc["Name_of_Client"]) or p.isna(rc["Client_Account_No"]):continue
                pv={k:rc[k] for k in pv}
                nm=str(rc["Name_of_Client"]).replace(" ","_")
                iso=rc.get("ISIN","Unknown_ISIN")
                C2(rc,f"{od}/TC_{nm}_{iso}_{i+1}.docx")
            zf=f"{od}/st.zip";zz=Z.ZipFile(zf,"w");[zz.write(fn,O.path.basename(fn)) for fn in G.glob(f"{od}/*.docx")];zz.close()
            F.download(zf)
        b.on_click(gen);X(b)
    up.observe(cb,names="value")
from ipywidgets import Output, Tab, HTML, VBox, Layout
from IPython.display import display

# 1) Header (financial style, polished and animated)
header = HTML(
    "<div style='"
    "background: linear-gradient(100deg, #0B3954, #1E6091); "
    "padding: 22px 30px; "
    "border-radius: 14px; "
    "box-shadow: 0 8px 16px rgba(0,0,0,0.1); "
    "margin-bottom: 20px; "
    "transition: all 0.3s ease;'>"
    "<h1 style='"
    "color: #FFFFFF; "
    "font-family: \"Helvetica Neue\", Arial, sans-serif; "
    "font-size: 32px; "
    "font-weight: 700; "
    "margin: 0;'>"
    "Excellent Tide Securities</h1>"
    "<p style='"
    "color: #D0EDFF; "
    "font-size: 16px; "
    "margin: 8px 0 0; "
    "font-style: italic;'>"
    "Interactive Trade & Settlement Portal</p>"
    "</div>"
)

# 2) Outputs for each tool
out_m0 = Output()
with out_m0:
    M0()      # Settlement tool

out_m1 = Output()
with out_m1:
    M1()      # Statement tool

# 3) Tab container (modern tab styling)
tabs = Tab(children=[out_m0, out_m1])
tabs.set_title(0, 'Settlement Tools')
tabs.set_title(1, 'Statement Tools')
tabs.layout = Layout(
    width='98%',
    border='1.5px solid #0B3954',
    border_radius='12px',
    padding='12px',
    background_color='#F7F9FB',  # lighter gray-blue
    box_shadow='0 4px 12px rgba(0, 0, 0, 0.05)'
)

# 4) Wrap header+tabs in a box that can have background_color
app = VBox(
    [header, tabs],
    layout=Layout(
        width='100%',
        padding='20px',
        background_color='#EDF3F8'  # consistent, soft background
    )
)

# 5) Render it
display(app)
""")