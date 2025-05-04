# Программа поиска дубликатов файлов с графическим интерфейсом

Вот полный код программы на Python с графическим интерфейсом для поиска дубликатов файлов:

```python
import os
import hashlib
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from docx import Document
from threading import Thread
import time


class FileDuplicateFinder:
    def __init__(self, root):
        self.root = root
        self.root.title("Поиск дубликатов файлов")
        self.root.geometry("800x600")
        
        # Переменные
        self.selected_drives = []
        self.duplicates = {}
        self.total_files = 0
        self.processed_files = 0
        self.is_running = False
        self.is_paused = False
        self.cancel_requested = False
        
        # Создание интерфейса
        self.create_widgets()
        
    def create_widgets(self):
        # Фрейм для выбора дисков
        drive_frame = LabelFrame(self.root, text="Выберите диски для поиска", padx=10, pady=10)
        drive_frame.pack(fill=X, padx=10, pady=5)
        
        # Получаем список доступных дисков
        drives = [f"{d}:\\" for d in 'ABCDEFGHIJKLMNOPQRSTUVWXYZ' if os.path.exists(f"{d}:\\")]
        
        # Создаем чекбоксы для каждого диска
        self.drive_vars = []
        for i, drive in enumerate(drives):
            var = BooleanVar()
            cb = Checkbutton(drive_frame, text=drive, variable=var)
            cb.grid(row=i//4, column=i%4, sticky=W)
            self.drive_vars.append((drive, var))
        
        # Фрейм управления
        control_frame = Frame(self.root, padx=10, pady=10)
        control_frame.pack(fill=X, padx=10, pady=5)
        
        # Кнопки управления
        self.start_btn = Button(control_frame, text="Начать поиск", command=self.start_search)
        self.start_btn.pack(side=LEFT, padx=5)
        
        self.pause_btn = Button(control_frame, text="Пауза", command=self.toggle_pause, state=DISABLED)
        self.pause_btn.pack(side=LEFT, padx=5)
        
        self.cancel_btn = Button(control_frame, text="Отмена", command=self.cancel_search, state=DISABLED)
        self.cancel_btn.pack(side=LEFT, padx=5)
        
        self.save_btn = Button(control_frame, text="Сохранить результаты", command=self.save_results, state=DISABLED)
        self.save_btn.pack(side=RIGHT, padx=5)
        
        # Прогресс бар
        self.progress = ttk.Progressbar(self.root, orient=HORIZONTAL, length=100, mode='determinate')
        self.progress.pack(fill=X, padx=10, pady=5)
        
        # Статус
        self.status_var = StringVar()
        self.status_var.set("Готов к работе")
        status_label = Label(self.root, textvariable=self.status_var, anchor=W)
        status_label.pack(fill=X, padx=10, pady=5)
        
        # Результаты
        result_frame = LabelFrame(self.root, text="Результаты поиска", padx=10, pady=10)
        result_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        
        self.result_text = Text(result_frame, wrap=NONE)
        scroll_y = Scrollbar(result_frame, orient=VERTICAL, command=self.result_text.yview)
        scroll_x = Scrollbar(result_frame, orient=HORIZONTAL, command=self.result_text.xview)
        self.result_text.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        
        scroll_y.pack(side=RIGHT, fill=Y)
        scroll_x.pack(side=BOTTOM, fill=X)
        self.result_text.pack(fill=BOTH, expand=True)
        
    def start_search(self):
        # Получаем выбранные диски
        self.selected_drives = [drive for drive, var in self.drive_vars if var.get()]
        
        if not self.selected_drives:
            messagebox.showwarning("Ошибка", "Выберите хотя бы один диск для поиска")
            return
        
        # Сброс предыдущих результатов
        self.duplicates = {}
        self.total_files = 0
        self.processed_files = 0
        self.cancel_requested = False
        self.is_paused = False
        
        # Обновление интерфейса
        self.start_btn.config(state=DISABLED)
        self.pause_btn.config(state=NORMAL)
        self.cancel_btn.config(state=NORMAL)
        self.save_btn.config(state=DISABLED)
        self.result_text.delete(1.0, END)
        self.status_var.set("Подсчет файлов...")
        self.root.update()
        
        # Подсчет общего количества файлов
        self.count_total_files()
        
        if self.cancel_requested:
            return
        
        # Начало поиска дубликатов
        self.status_var.set("Поиск дубликатов...")
        self.is_running = True
        
        # Запуск в отдельном потоке
        search_thread = Thread(target=self.find_duplicates)
        search_thread.start()
        
    def count_total_files(self):
        for drive in self.selected_drives:
            for root, _, files in os.walk(drive):
                if self.cancel_requested:
                    return
                
                # Пропускаем системные папки
                if any(ignore in root for ignore in ['$RECYCLE.BIN', 'System Volume Information']):
                    continue
                
                self.total_files += len(files)
                
                # Обновляем статус каждые 1000 файлов
                if self.total_files % 1000 == 0:
                    self.status_var.set(f"Подсчет файлов... Найдено {self.total_files}")
                    self.root.update()
    
    def find_duplicates(self):
        file_hashes = {}
        
        for drive in self.selected_drives:
            for root, _, files in os.walk(drive):
                if self.cancel_requested:
                    self.search_completed(False)
                    return
                
                # Ожидание, если стоит пауза
                while self.is_paused and not self.cancel_requested:
                    time.sleep(0.1)
                
                if self.cancel_requested:
                    self.search_completed(False)
                    return
                
                # Пропускаем системные папки
                if any(ignore in root for ignore in ['$RECYCLE.BIN', 'System Volume Information']):
                    continue
                
                for file in files:
                    if self.cancel_requested:
                        self.search_completed(False)
                        return
                    
                    # Ожидание, если стоит пауза
                    while self.is_paused and not self.cancel_requested:
                        time.sleep(0.1)
                    
                    if self.cancel_requested:
                        self.search_completed(False)
                        return
                    
                    try:
                        file_path = os.path.join(root, file)
                        
                        # Пропускаем символические ссылки и другие специальные файлы
                        if not os.path.isfile(file_path):
                            continue
                        
                        # Вычисляем хеш файла
                        file_hash = self.calculate_hash(file_path)
                        
                        # Добавляем в словарь
                        if file_hash in file_hashes:
                            file_hashes[file_hash].append(file_path)
                        else:
                            file_hashes[file_hash] = [file_path]
                        
                        # Увеличиваем счетчик обработанных файлов
                        self.processed_files += 1
                        
                        # Обновляем прогресс
                        progress = (self.processed_files / self.total_files) * 100
                        self.progress['value'] = progress
                        self.status_var.set(
                            f"Обработано {self.processed_files} из {self.total_files} файлов ({progress:.1f}%)"
                        )
                        self.root.update()
                        
                    except (PermissionError, OSError) as e:
                        continue
        
        # Фильтруем только дубликаты (файлы с одинаковым хешем)
        self.duplicates = {k: v for k, v in file_hashes.items() if len(v) > 1}
        
        # Завершение поиска
        self.search_completed(True)
    
    def calculate_hash(self, file_path, chunk_size=8192):
        hasher = hashlib.md5()
        
        with open(file_path, 'rb') as f:
            while chunk := f.read(chunk_size):
                if self.cancel_requested or self.is_paused:
                    break
                hasher.update(chunk)
        
        return hasher.hexdigest()
    
    def toggle_pause(self):
        self.is_paused = not self.is_paused
        if self.is_paused:
            self.pause_btn.config(text="Продолжить")
            self.status_var.set(f"Пауза... Обработано {self.processed_files} из {self.total_files} файлов")
        else:
            self.pause_btn.config(text="Пауза")
            self.status_var.set(f"Продолжение... Обработано {self.processed_files} из {self.total_files} файлов")
    
    def cancel_search(self):
        self.cancel_requested = True
        self.status_var.set("Отмена операции...")
        self.pause_btn.config(state=DISABLED)
        self.cancel_btn.config(state=DISABLED)
    
    def search_completed(self, success):
        self.is_running = False
        
        # Обновление интерфейса
        self.start_btn.config(state=NORMAL)
        self.pause_btn.config(state=DISABLED)
        self.cancel_btn.config(state=DISABLED)
        
        if success and not self.cancel_requested:
            self.save_btn.config(state=NORMAL)
            self.status_var.set(f"Поиск завершен. Найдено {len(self.duplicates)} групп дубликатов.")
            
            # Вывод результатов
            self.result_text.delete(1.0, END)
            for group in self.duplicates.values():
                self.result_text.insert(END, "Дубликаты:\n")
                for file in group:
                    self.result_text.insert(END, f"  - {file}\n")
                self.result_text.insert(END, "\n")
        else:
            self.status_var.set("Поиск отменен.")
    
    def save_results(self):
        if not self.duplicates:
            messagebox.showinfo("Информация", "Нет результатов для сохранения")
            return
        
        # Запрашиваем место сохранения
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            title="Сохранить результаты"
        )
        
        if not file_path:
            return
        
        try:
            # Создаем документ Word
            doc = Document()
            
            # Добавляем заголовок
            doc.add_heading('Результаты поиска дубликатов файлов', 0)
            
            # Добавляем таблицу
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Файл 1 (Имя и путь)'
            hdr_cells[1].text = 'Файл 2 (Имя и путь)'
            
            # Добавляем данные
            for group in self.duplicates.values():
                # Для каждой пары файлов в группе дубликатов
                for i in range(len(group)):
                    for j in range(i+1, len(group)):
                        row_cells = table.add_row().cells
                        row_cells[0].text = group[i]
                        row_cells[1].text = group[j]
            
            # Сохраняем документ
            doc.save(file_path)
            messagebox.showinfo("Успех", f"Результаты успешно сохранены в {file_path}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")


if __name__ == "__main__":
    root = Tk()
    app = FileDuplicateFinder(root)
    root.mainloop()
```

## Описание программы

### Особенности программы:
1. **Графический интерфейс** с использованием Tkinter
2. **Выбор дисков** для поиска через чекбоксы
3. **Отображение прогресса** поиска (количество обработанных файлов и процент выполнения)
4. **Управление процессом**:
   - Кнопка "Пауза" для временной остановки поиска
   - Кнопка "Отмена" для прерывания процесса
5. **Сохранение результатов** в формате Word (.docx) с таблицей дубликатов
6. **Многопоточность** - поиск выполняется в отдельном потоке, чтобы не блокировать интерфейс

### Как использовать:
1. Запустите программу
2. Выберите диски для поиска
3. Нажмите "Начать поиск"
4. Дождитесь завершения (или используйте кнопки управления)
5. Просмотрите результаты в интерфейсе
6. Сохраните результаты в Word-документ при необходимости

### Требования:
Для работы программы необходимо установить следующие библиотеки:
```
pip install tkinter python-docx
```
