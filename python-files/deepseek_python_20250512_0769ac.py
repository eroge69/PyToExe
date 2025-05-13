import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from datetime import datetime, timedelta
from calendar import monthrange
import docx
from docx.shared import Pt
import json
import os
import random
from random import choice, randint

class Cadet:
    def __init__(self, name, category, platoon="", custom_duty_count=None, available_duties=None, tags=None, rank="Рядовой вн. сл."):
        self.name = name
        self.category = category  # "Дежурные юноши", "Дневальные юноши", "Дежурные девушки", "Дневальные девушки"
        self.platoon = platoon
        self.custom_duty_count = custom_duty_count
        self.available_duties = available_duties if available_duties else []
        self.tags = tags if tags else []
        self.assigned_duties = []
        self.last_duty_date = None
        self.rank = rank
    
    def can_serve(self, duty_date, duty_name):
        # Проверка интервала между нарядами (минимум 2 дня)
        if self.last_duty_date and (duty_date - self.last_duty_date).days < 2:
            return False
        
        # Проверка меток
        for tag, value in self.tags:
            if tag == "Сборник":
                return False
            elif tag == "Больной" or tag == "Рапорт":
                start_date, end_date = [datetime.strptime(d, "%Y-%m-%d").date() for d in value.split(" - ")]
                if start_date <= duty_date <= end_date:
                    return False
            elif tag == "Каптер":
                # Каптер заступает только на определенный пост со вторника на среду
                if duty_date.weekday() == 1:  # Вторник
                    return duty_name == "Каптер"
                else:
                    if duty_name == "Каптер":
                        return False
        
        # Проверка доступности наряда
        if self.available_duties and duty_name not in self.available_duties:
            return False
        
        return True
    
    def add_duty(self, duty_date, duty_name):
        self.assigned_duties.append((duty_date, duty_name))
        self.last_duty_date = duty_date

class Duty:
    def __init__(self, name, requirements, all_cadets=False, duty_days=None, every_day=False):
        self.name = name
        self.requirements = requirements  # Словарь: {"Дежурные юноши": 1, "Дневальные девушки": 2, ...}
        self.all_cadets = all_cadets  # Могут заступать все курсанты категории
        self.duty_days = duty_days if duty_days else []
        self.every_day = every_day

class DutyDistributionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Распределение нарядов. Курс набора 2024 года")
        self.root.geometry("1000x700")
        
        # Русские названия дней недели
        self.weekdays_ru = {
            0: "Понедельник",
            1: "Вторник",
            2: "Среда",
            3: "Четверг",
            4: "Пятница",
            5: "Суббота",
            6: "Воскресенье"
        }
        
        # Данные
        self.cadets = []
        self.duties = []
        self.platoons = []
        self.duty_distribution = {}  # {date: {duty_name: [cadets]}
        self.settings = {
            "average_duties": {
                "Дежурные юноши": 4,
                "Дневальные юноши": 4,
                "Дежурные девушки": 4,
                "Дневальные девушки": 4
            },
            "duty_interval": 3,
            "course": "2024"
        }
        
        # Создание вкладок
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Вкладка "Курсанты"
        self.create_cadets_tab()
        
        # Вкладка "Наряды"
        self.create_duties_tab()
        
        # Вкладка "Распределение"
        self.create_distribution_tab()
        
        # Вкладка "Просмотр и редактирование"
        self.create_view_tab()
        
        # Вкладка "Статистика"
        self.create_stats_tab()
        
        # Загрузка данных при запуске
        self.load_data()
    
    def create_cadets_tab(self):
        self.cadets_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.cadets_tab, text="Курсанты")
        
        # Категории курсантов
        self.categories = ["Дежурные юноши", "Дневальные юноши", "Дежурные девушки", "Дневальные девушки"]
        
        # Фреймы для каждой категории
        self.category_frames = {}
        for i, category in enumerate(self.categories):
            frame = ttk.LabelFrame(self.cadets_tab, text=category)
            frame.grid(row=0, column=i, padx=5, pady=5, sticky="nsew")
            self.category_frames[category] = frame
            
            # Список курсантов
            listbox = tk.Listbox(frame, height=20, width=25)
            listbox.pack(fill=tk.BOTH, expand=True)
            listbox.bind("<Double-Button-1>", lambda e, c=category: self.edit_cadet(c))
            setattr(self, f"{category.lower().replace(' ', '_')}_listbox", listbox)
            
            # Кнопка добавления
            btn_add = ttk.Button(frame, text="Добавить", command=lambda c=category: self.add_cadet(c))
            btn_add.pack(fill=tk.X)
            
            # Кнопка удаления
            btn_remove = ttk.Button(frame, text="Удалить", command=lambda c=category: self.remove_cadet(c))
            btn_remove.pack(fill=tk.X)
            
            # Кнопка загрузки из файла
            btn_load = ttk.Button(frame, text="Загрузить из файла", command=lambda c=category: self.load_cadets_from_file(c))
            btn_load.pack(fill=tk.X)
        
        # Настройки
        settings_frame = ttk.LabelFrame(self.cadets_tab, text="Настройки распределения")
        settings_frame.grid(row=1, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
        
        # Среднее количество нарядов
        ttk.Label(settings_frame, text="Среднее количество нарядов:").grid(row=0, column=0, padx=5, pady=5)
        for i, category in enumerate(self.categories):
            ttk.Label(settings_frame, text=category).grid(row=0, column=i*2+1, padx=5)
            spin = ttk.Spinbox(settings_frame, from_=1, to=10, width=3)
            spin.grid(row=0, column=i*2+2, padx=5)
            spin.set(self.settings["average_duties"][category])
            spin.bind("<FocusOut>", lambda e, c=category: self.update_average_duties(c))
            setattr(self, f"avg_duties_{category.lower().replace(' ', '_')}", spin)
        
        # Интервал между нарядами
        ttk.Label(settings_frame, text="Интервал между нарядами (дней):").grid(row=1, column=0, padx=5, pady=5)
        self.interval_spin = ttk.Spinbox(settings_frame, from_=1, to=7, width=3)
        self.interval_spin.grid(row=1, column=1, padx=5)
        self.interval_spin.set(self.settings["duty_interval"])
        self.interval_spin.bind("<FocusOut>", lambda e: self.update_interval())
        
        # Выбор курса
        ttk.Label(settings_frame, text="Курс:").grid(row=1, column=2, padx=5, pady=5)
        self.course_combo = ttk.Combobox(settings_frame, values=["1 курс", "2 курс", "3 курс", "4 курс", "5 курс"], width=5)
        self.course_combo.grid(row=1, column=3, padx=5)
        self.course_combo.set(self.settings["course"])
        self.course_combo.bind("<<ComboboxSelected>>", lambda e: self.update_course())
        
        # Управление взводами
        ttk.Label(settings_frame, text="Добавить взвод:").grid(row=2, column=0, padx=5, pady=5)
        self.platoon_entry = ttk.Entry(settings_frame, width=15)
        self.platoon_entry.grid(row=2, column=1, padx=5)
        self.add_platoon_btn = ttk.Button(settings_frame, text="Добавить", command=self.add_platoon)
        self.add_platoon_btn.grid(row=2, column=2, padx=5)
        
        # Список взводов
        self.platoon_listbox = tk.Listbox(settings_frame, height=5, width=20)
        self.platoon_listbox.grid(row=2, column=3, padx=5, pady=5, rowspan=2)
        self.platoon_listbox.bind("<Double-Button-1>", lambda e: self.edit_platoon())
        
        # Кнопка сохранения
        save_btn = ttk.Button(self.cadets_tab, text="Сохранить данные", command=self.save_data)
        save_btn.grid(row=2, column=0, columnspan=4, pady=10)
    
    def create_duties_tab(self):
        self.duties_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.duties_tab, text="Наряды")
        
        # Список нарядов
        self.duties_listbox = tk.Listbox(self.duties_tab, height=15)
        self.duties_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.duties_listbox.bind("<Double-Button-1>", lambda e: self.edit_duty())
        
        # Кнопки управления
        btn_frame = ttk.Frame(self.duties_tab)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        add_btn = ttk.Button(btn_frame, text="Добавить", command=self.add_duty)
        add_btn.pack(side=tk.LEFT, padx=5)
        
        edit_btn = ttk.Button(btn_frame, text="Редактировать", command=self.edit_duty)
        edit_btn.pack(side=tk.LEFT, padx=5)
        
        remove_btn = ttk.Button(btn_frame, text="Удалить", command=self.remove_duty)
        remove_btn.pack(side=tk.LEFT, padx=5)
    
    def create_distribution_tab(self):
        self.distribution_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.distribution_tab, text="Распределение")
        
        # Параметры распределения
        params_frame = ttk.LabelFrame(self.distribution_tab, text="Параметры распределения")
        params_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # Месяц и год
        ttk.Label(params_frame, text="Месяц и год:").grid(row=0, column=0, padx=5, pady=5)
        self.month_combo = ttk.Combobox(params_frame, values=list(range(1, 13)), width=3)
        self.month_combo.grid(row=0, column=1, padx=5)
        self.month_combo.set(datetime.now().month)
        
        self.year_combo = ttk.Combobox(params_frame, values=list(range(2023, 2026)), width=5)
        self.year_combo.grid(row=0, column=2, padx=5)
        self.year_combo.set(datetime.now().year)
        
        # Первый день распределения
        ttk.Label(params_frame, text="Первый день:").grid(row=0, column=3, padx=5)
        self.first_day_spin = ttk.Spinbox(params_frame, from_=1, to=31, width=3)
        self.first_day_spin.grid(row=0, column=4, padx=5)
        self.first_day_spin.set(1)
        
        # Количество дней
        ttk.Label(params_frame, text="Количество дней:").grid(row=0, column=5, padx=5)
        self.days_count_spin = ttk.Spinbox(params_frame, from_=1, to=31, width=3)
        self.days_count_spin.grid(row=0, column=6, padx=5)
        self.days_count_spin.bind("<FocusOut>", self.update_days_count)
        
        # Дежурное подразделение
        self.duty_platoon_var = tk.IntVar()
        self.duty_platoon_check = ttk.Checkbutton(params_frame, text="Дежурное подразделение", 
                                                 variable=self.duty_platoon_var, command=self.toggle_duty_platoon)
        self.duty_platoon_check.grid(row=1, column=0, columnspan=3, padx=5, pady=5, sticky="w")
        
        self.duty_platoon_frame = ttk.Frame(params_frame)
        
        ttk.Label(self.duty_platoon_frame, text="Взвод:").grid(row=0, column=0, padx=5)
        self.duty_platoon_combo = ttk.Combobox(self.duty_platoon_frame, values=[], width=15)
        self.duty_platoon_combo.grid(row=0, column=1, padx=5)
        
        ttk.Label(self.duty_platoon_frame, text="Неделя:").grid(row=0, column=2, padx=5)
        self.duty_platoon_week_combo = ttk.Combobox(self.duty_platoon_frame, values=list(range(1, 5)), width=3)
        self.duty_platoon_week_combo.grid(row=0, column=3, padx=5)
        
        ttk.Label(self.duty_platoon_frame, text="Дни недели:").grid(row=0, column=4, padx=5)
        self.duty_platoon_days_entry = ttk.Entry(self.duty_platoon_frame, width=15)
        self.duty_platoon_days_entry.grid(row=0, column=5, padx=5)
        
        self.add_duty_platoon_btn = ttk.Button(self.duty_platoon_frame, text="Добавить", command=self.add_duty_platoon)
        self.add_duty_platoon_btn.grid(row=0, column=6, padx=5)
        
        self.duty_platoon_listbox = tk.Listbox(self.duty_platoon_frame, height=4, width=60)
        self.duty_platoon_listbox.grid(row=1, column=0, columnspan=7, padx=5, pady=5)
        
        # Фрейм для выбора нарядов
        duties_frame = ttk.LabelFrame(self.distribution_tab, text="Выбор нарядов для распределения")
        duties_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Прокрутка для нарядов
        self.duties_canvas = tk.Canvas(duties_frame)
        self.duties_scrollbar = ttk.Scrollbar(duties_frame, orient="vertical", command=self.duties_canvas.yview)
        self.duties_scrollable_frame = ttk.Frame(self.duties_canvas)
        
        self.duties_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.duties_canvas.configure(
                scrollregion=self.duties_canvas.bbox("all")
            )
        )
        
        self.duties_canvas.create_window((0, 0), window=self.duties_scrollable_frame, anchor="nw")
        self.duties_canvas.configure(yscrollcommand=self.duties_scrollbar.set)
        
        self.duties_canvas.pack(side="left", fill="both", expand=True)
        self.duties_scrollbar.pack(side="right", fill="y")
        
        # Кнопка генерации
        generate_btn = ttk.Button(self.distribution_tab, text="Сгенерировать распределение", command=self.generate_distribution)
        generate_btn.pack(pady=10)
        
        # Обновляем список нарядов
        self.update_duty_selection()
    
    def create_view_tab(self):
        self.view_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.view_tab, text="Просмотр и редактирование")
        
        # Просмотр распределения
        self.view_frame = ttk.Frame(self.view_tab)
        self.view_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.view_canvas = tk.Canvas(self.view_frame)
        self.view_scrollbar = ttk.Scrollbar(self.view_frame, orient="vertical", command=self.view_canvas.yview)
        self.view_scrollable_frame = ttk.Frame(self.view_canvas)
        
        self.view_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.view_canvas.configure(
                scrollregion=self.view_canvas.bbox("all")
            )
        )
        
        self.view_canvas.create_window((0, 0), window=self.view_scrollable_frame, anchor="nw")
        self.view_canvas.configure(yscrollcommand=self.view_scrollbar.set)
        
        self.view_canvas.pack(side="left", fill="both", expand=True)
        self.view_scrollbar.pack(side="right", fill="y")
        
        # Кнопки управления
        btn_frame = ttk.Frame(self.view_tab)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        edit_btn = ttk.Button(btn_frame, text="Редактировать", command=self.edit_distribution)
        edit_btn.pack(side=tk.LEFT, padx=5)
        
        save_btn = ttk.Button(btn_frame, text="Сохранить в Word", command=self.export_to_word)
        save_btn.pack(side=tk.RIGHT, padx=5)
    
    def create_stats_tab(self):
        self.stats_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.stats_tab, text="Статистика")
        
        # Статистика по курсантам
        self.stats_tree = ttk.Treeview(self.stats_tab, columns=("name", "count", "details"), show="headings")
        self.stats_tree.heading("name", text="Курсант")
        self.stats_tree.heading("count", text="Количество нарядов")
        self.stats_tree.heading("details", text="Детали")
        
        # Добавляем прокрутку
        scrollbar = ttk.Scrollbar(self.stats_tab, orient="vertical", command=self.stats_tree.yview)
        self.stats_tree.configure(yscrollcommand=scrollbar.set)
        
        self.stats_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Кнопки управления
        btn_frame = ttk.Frame(self.stats_tab)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        refresh_btn = ttk.Button(btn_frame, text="Обновить", command=self.update_stats)
        refresh_btn.pack(side=tk.LEFT, padx=5)
        
        delete_btn = ttk.Button(btn_frame, text="Удалить месяц", command=self.delete_month_stats)
        delete_btn.pack(side=tk.RIGHT, padx=5)
    
    def load_data(self):
        # Загрузка данных из файлов (если есть)
        try:
            if os.path.exists("cadets.json"):
                with open("cadets.json", "r") as f:
                    data = json.load(f)
                    self.cadets = [Cadet(**item) for item in data["cadets"]]
                    self.platoons = data.get("platoons", [])
                    self.settings = data.get("settings", self.settings)
            
            if os.path.exists("duties.json"):
                with open("duties.json", "r") as f:
                    duties_data = json.load(f)
                    self.duties = []
                    for item in duties_data:
                        self.duties.append(Duty(
                            item["name"], 
                            item["requirements"], 
                            item.get("all_cadets", False),
                            item.get("duty_days", []),
                            item.get("every_day", False)
                        ))
            
            if os.path.exists("distribution.json"):
                with open("distribution.json", "r") as f:
                    self.duty_distribution = json.load(f)
            
            self.update_cadets_lists()
            self.update_duties_list()
            self.update_platoon_list()
            self.update_stats()
            self.update_view_tab()
            self.update_duty_selection()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить данные: {str(e)}")
    
    def save_data(self):
        try:
            # Сохранение курсантов
            cadets_data = {
                "cadets": [{
                    "name": c.name,
                    "category": c.category,
                    "platoon": c.platoon,
                    "custom_duty_count": c.custom_duty_count,
                    "available_duties": c.available_duties,
                    "tags": c.tags,
                    "rank": c.rank
                } for c in self.cadets],
                "platoons": self.platoons,
                "settings": self.settings
            }
            
            with open("cadets.json", "w") as f:
                json.dump(cadets_data, f, indent=2)
            
            # Сохранение нарядов
            duties_data = []
            for duty in self.duties:
                duties_data.append({
                    "name": duty.name,
                    "requirements": duty.requirements,
                    "all_cadets": duty.all_cadets,
                    "duty_days": duty.duty_days,
                    "every_day": duty.every_day
                })
            
            with open("duties.json", "w") as f:
                json.dump(duties_data, f, indent=2)
            
            # Сохранение распределения
            with open("distribution.json", "w") as f:
                json.dump(self.duty_distribution, f, indent=2)
            
            messagebox.showinfo("Сохранено", "Данные успешно сохранены")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить данные: {str(e)}")
    
    def update_cadets_lists(self):
        for category in self.categories:
            listbox = getattr(self, f"{category.lower().replace(' ', '_')}_listbox")
            listbox.delete(0, tk.END)
            
            for cadet in [c for c in self.cadets if c.category == category]:
                listbox.insert(tk.END, cadet.name)
    
    def update_duties_list(self):
        self.duties_listbox.delete(0, tk.END)
        for duty in self.duties:
            reqs = ", ".join([f"{k}: {v}" for k, v in duty.requirements.items()])
            self.duties_listbox.insert(tk.END, f"{duty.name} ({reqs})")
    
    def update_duty_selection(self):
        # Очищаем предыдущие элементы
        for widget in self.duties_scrollable_frame.winfo_children():
            widget.destroy()
        
        # Создаем новые элементы для каждого наряда
        for i, duty in enumerate(self.duties):
            duty_frame = ttk.LabelFrame(self.duties_scrollable_frame, text=duty.name)
            duty_frame.pack(fill=tk.X, padx=5, pady=5)
            
            # Выбор дней для наряда
            ttk.Label(duty_frame, text="Дни заступления (через пробел):").grid(row=0, column=0, padx=5, pady=2)
            days_entry = ttk.Entry(duty_frame, width=20)
            days_entry.grid(row=0, column=1, padx=5, pady=2)
            days_entry.insert(0, " ".join(map(str, duty.duty_days)))
            
            # Галочка "Каждый день"
            every_day_var = tk.IntVar(value=int(duty.every_day))
            every_day_cb = ttk.Checkbutton(duty_frame, text="Каждый день", variable=every_day_var)
            every_day_cb.grid(row=0, column=2, padx=5, pady=2)
            
            # Сохраняем ссылки на элементы для каждого наряда
            duty.days_entry = days_entry
            duty.every_day_var = every_day_var
    
    def update_platoon_list(self):
        self.platoon_listbox.delete(0, tk.END)
        for platoon in self.platoons:
            self.platoon_listbox.insert(tk.END, platoon)
        
        # Обновление комбобоксов
        self.duty_platoon_combo["values"] = self.platoons
        self.course_combo.set(self.settings["course"])
    
    def update_stats(self):
        self.stats_tree.delete(*self.stats_tree.get_children())
        
        cadet_stats = {}
        for date_str, duties in self.duty_distribution.items():
            date = datetime.strptime(date_str, "%Y-%m-%d").date()
            for duty_name, cadets in duties.items():
                for cadet_name in cadets:
                    if cadet_name not in cadet_stats:
                        cadet_stats[cadet_name] = {"count": 0, "details": []}
                    cadet_stats[cadet_name]["count"] += 1
                    cadet_stats[cadet_name]["details"].append(f"{duty_name} ({date.strftime('%d.%m')})")
        
        for cadet_name, stats in sorted(cadet_stats.items(), key=lambda x: x[1]["count"], reverse=True):
            self.stats_tree.insert("", tk.END, values=(
                cadet_name,
                stats["count"],
                "; ".join(stats["details"])
            ))
    
    def update_view_tab(self):
        # Очищаем предыдущее содержимое
        for widget in self.view_scrollable_frame.winfo_children():
            widget.destroy()
        
        if not self.duty_distribution:
            return
        
        # Группируем наряды по дням
        days = sorted(self.duty_distribution.items(), key=lambda x: x[0])
        
        for date_str, duties in days:
            date = datetime.strptime(date_str, "%Y-%m-%d").date()
            day_name = self.weekdays_ru[date.weekday()]
            
            # Для каптеров определяем период с вторника по среду
            if any(duty_name == "Каптер" for duty_name in duties.keys()):
                next_day = date + timedelta(days=1)
                next_day_name = self.weekdays_ru[next_day.weekday()]
                date_range = f"{date.strftime('%d.%m.%Y')}-{next_day.strftime('%d.%m.%Y')}"
                
                # Создаем фрейм для периода
                period_frame = ttk.LabelFrame(self.view_scrollable_frame, 
                                            text=f"{date_range} ({day_name}-{next_day_name})")
                period_frame.pack(fill=tk.X, padx=5, pady=5, expand=True)
                
                # Добавляем наряды
                for duty_name, cadets in duties.items():
                    if not cadets:
                        continue
                    
                    duty_frame = ttk.Frame(period_frame)
                    duty_frame.pack(fill=tk.X, padx=5, pady=2)
                    
                    ttk.Label(duty_frame, text=duty_name, font=('Helvetica', 10, 'bold')).pack(anchor="w")
                    
                    for cadet in cadets:
                        cadet_obj = next((c for c in self.cadets if c.name == cadet), None)
                        if not cadet_obj:
                            continue
                        
                        role = "Дежурный" if "Дежурный" in cadet_obj.category else "Дневальный"
                        
                        cadet_frame = ttk.Frame(duty_frame)
                        cadet_frame.pack(fill=tk.X, padx=20, pady=1)
                        
                        ttk.Label(cadet_frame, text=f"{role}", width=15, anchor="w").pack(side=tk.LEFT)
                        ttk.Label(cadet_frame, text=cadet_obj.rank, width=20, anchor="w").pack(side=tk.LEFT)
                        ttk.Label(cadet_frame, text=cadet, width=25, anchor="w").pack(side=tk.LEFT)
            else:
                # Обычный день
                day_frame = ttk.LabelFrame(self.view_scrollable_frame, text=f"{date.strftime('%d.%m.%Y')} ({day_name})")
                day_frame.pack(fill=tk.X, padx=5, pady=5, expand=True)
                
                # Добавляем наряды
                for duty_name, cadets in duties.items():
                    if not cadets:
                        continue
                    
                    duty_frame = ttk.Frame(day_frame)
                    duty_frame.pack(fill=tk.X, padx=5, pady=2)
                    
                    ttk.Label(duty_frame, text=duty_name, font=('Helvetica', 10, 'bold')).pack(anchor="w")
                    
                    for cadet in cadets:
                        cadet_obj = next((c for c in self.cadets if c.name == cadet), None)
                        if not cadet_obj:
                            continue
                        
                        role = "Дежурный" if "Дежурный" in cadet_obj.category else "Дневальный"
                        
                        cadet_frame = ttk.Frame(duty_frame)
                        cadet_frame.pack(fill=tk.X, padx=20, pady=1)
                        
                        ttk.Label(cadet_frame, text=f"{role}", width=15, anchor="w").pack(side=tk.LEFT)
                        ttk.Label(cadet_frame, text=cadet_obj.rank, width=20, anchor="w").pack(side=tk.LEFT)
                        ttk.Label(cadet_frame, text=cadet, width=25, anchor="w").pack(side=tk.LEFT)
    
    def add_cadet(self, category):
        name = simpledialog.askstring("Добавить курсанта", "Введите ФИО курсанта:")
        if name:
            self.cadets.append(Cadet(name, category))
            self.update_cadets_lists()
    
    def remove_cadet(self, category):
        listbox = getattr(self, f"{category.lower().replace(' ', '_')}_listbox")
        selection = listbox.curselection()
        if not selection:
            return
        
        cadet_name = listbox.get(selection[0])
        cadet = next((c for c in self.cadets if c.name == cadet_name and c.category == category), None)
        if not cadet:
            return
        
        if messagebox.askyesno("Подтверждение", f"Удалить курсанта {cadet_name}?"):
            self.cadets.remove(cadet)
            self.update_cadets_lists()
            self.save_data()  # Сохраняем изменения
    
    def load_cadets_from_file(self, category):
        file_path = filedialog.askopenfilename(filetypes=[("Текстовые файлы", "*.txt")])
        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    for line in f:
                        name = line.strip()
                        if name:
                            self.cadets.append(Cadet(name, category))
                
                self.update_cadets_lists()
                messagebox.showinfo("Успешно", "Курсанты успешно добавлены")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить файл: {str(e)}")
    
    def edit_cadet(self, category):
        listbox = getattr(self, f"{category.lower().replace(' ', '_')}_listbox")
        selection = listbox.curselection()
        if not selection:
            return
        
        cadet_name = listbox.get(selection[0])
        cadet = next((c for c in self.cadets if c.name == cadet_name and c.category == category), None)
        if not cadet:
            return
        
        # Создаем окно редактирования
        edit_window = tk.Toplevel(self.root)
        edit_window.title(f"Редактирование: {cadet_name}")
        edit_window.geometry("600x500")
        
        # ФИО
        ttk.Label(edit_window, text="ФИО:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        name_entry = ttk.Entry(edit_window, width=30)
        name_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        name_entry.insert(0, cadet.name)
        
        # Взвод
        ttk.Label(edit_window, text="Взвод:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        platoon_combo = ttk.Combobox(edit_window, values=self.platoons, width=15)
        platoon_combo.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        platoon_combo.set(cadet.platoon)
        
        # Звание
        ttk.Label(edit_window, text="Звание:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        rank_combo = ttk.Combobox(edit_window, 
                                 values=["Рядовой вн. сл.", "Прапорщик вн. сл.", "Младший Сержант вн. сл.", "Сержант вн. сл."], 
                                 width=15)
        rank_combo.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        rank_combo.set(cadet.rank)
        
        # Индивидуальное количество нарядов
        ttk.Label(edit_window, text="Индив. кол-во нарядов:").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        custom_duty_spin = ttk.Spinbox(edit_window, from_=0, to=20, width=3)
        custom_duty_spin.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        custom_duty_spin.set(cadet.custom_duty_count if cadet.custom_duty_count is not None else "")
        
        # Доступные наряды
        ttk.Label(edit_window, text="Доступные наряды:").grid(row=4, column=0, padx=5, pady=5, sticky="ne")
        duties_frame = ttk.Frame(edit_window)
        duties_frame.grid(row=4, column=1, padx=5, pady=5, sticky="w")
        
        duty_vars = {}
        for i, duty in enumerate(self.duties):
            var = tk.IntVar(value=1 if duty.name in cadet.available_duties else 0)
            duty_vars[duty.name] = var
            cb = ttk.Checkbutton(duties_frame, text=duty.name, variable=var)
            cb.grid(row=i//2, column=i%2, sticky="w")
        
        # Метки
        ttk.Label(edit_window, text="Метки:").grid(row=5, column=0, padx=5, pady=5, sticky="ne")
        tags_frame = ttk.Frame(edit_window)
        tags_frame.grid(row=5, column=1, padx=5, pady=5, sticky="w")
        
        # Список меток
        tags_listbox = tk.Listbox(tags_frame, height=5, width=40)
        tags_listbox.pack(side=tk.LEFT, fill=tk.BOTH)
        
        for tag, value in cadet.tags:
            tags_listbox.insert(tk.END, f"{tag}: {value}")
        
        # Кнопки управления метками
        tags_btn_frame = ttk.Frame(tags_frame)
        tags_btn_frame.pack(side=tk.LEFT, padx=5)
        
        add_tag_btn = ttk.Button(tags_btn_frame, text="Добавить", command=lambda: self.add_tag(tags_listbox))
        add_tag_btn.pack(fill=tk.X)
        
        edit_tag_btn = ttk.Button(tags_btn_frame, text="Изменить", command=lambda: self.edit_tag(tags_listbox))
        edit_tag_btn.pack(fill=tk.X, pady=5)
        
        remove_tag_btn = ttk.Button(tags_btn_frame, text="Удалить", command=lambda: self.remove_tag(tags_listbox))
        remove_tag_btn.pack(fill=tk.X)
        
        # Кнопки сохранения/отмены
        btn_frame = ttk.Frame(edit_window)
        btn_frame.grid(row=6, column=0, columnspan=2, pady=10)
        
        save_btn = ttk.Button(btn_frame, text="Сохранить", command=lambda: self.save_cadet_changes(
            cadet, name_entry.get(), platoon_combo.get(), rank_combo.get(), 
            custom_duty_spin.get(), duty_vars, tags_listbox, edit_window))
        save_btn.pack(side=tk.LEFT, padx=10)
        
        cancel_btn = ttk.Button(btn_frame, text="Отмена", command=edit_window.destroy)
        cancel_btn.pack(side=tk.RIGHT, padx=10)
    
    def add_tag(self, listbox):
        tag_types = ["Сборник", "РПК", "Больной", "Рапорт", "Каптер"]
        tag_window = tk.Toplevel(self.root)
        tag_window.title("Добавить метку")
        tag_window.geometry("400x300")
        
        ttk.Label(tag_window, text="Тип метки:").pack(pady=5)
        tag_combo = ttk.Combobox(tag_window, values=tag_types)
        tag_combo.pack(pady=5)
        
        # Для меток "Больной" и "Рапорт" добавляем выбор дат
        date_frame = ttk.Frame(tag_window)
        date_frame.pack(pady=5)
        
        ttk.Label(date_frame, text="С:").grid(row=0, column=0, padx=5)
        start_day = ttk.Spinbox(date_frame, from_=1, to=31, width=3)
        start_day.grid(row=0, column=1, padx=5)
        start_day.set(datetime.now().day)
        
        start_month = ttk.Spinbox(date_frame, from_=1, to=12, width=3)
        start_month.grid(row=0, column=2, padx=5)
        start_month.set(datetime.now().month)
        
        start_year = ttk.Spinbox(date_frame, from_=2023, to=2025, width=5)
        start_year.grid(row=0, column=3, padx=5)
        start_year.set(datetime.now().year)
        
        ttk.Label(date_frame, text="По:").grid(row=1, column=0, padx=5)
        end_day = ttk.Spinbox(date_frame, from_=1, to=31, width=3)
        end_day.grid(row=1, column=1, padx=5)
        end_day.set(datetime.now().day + 3)
        
        end_month = ttk.Spinbox(date_frame, from_=1, to=12, width=3)
        end_month.grid(row=1, column=2, padx=5)
        end_month.set(datetime.now().month)
        
        end_year = ttk.Spinbox(date_frame, from_=2023, to=2025, width=5)
        end_year.grid(row=1, column=3, padx=5)
        end_year.set(datetime.now().year)
        
        # Примечание
        ttk.Label(tag_window, text="Примечание:").pack(pady=5)
        note_entry = ttk.Entry(tag_window, width=30)
        note_entry.pack(pady=5)
        
        def save_tag():
            tag = tag_combo.get()
            note = note_entry.get()
            
            if tag in ["Больной", "Рапорт"]:
                try:
                    start_date = f"{start_year.get()}-{start_month.get():02d}-{start_day.get():02d}"
                    end_date = f"{end_year.get()}-{end_month.get():02d}-{end_day.get():02d}"
                    value = f"{start_date} - {end_date}"
                    if note:
                        value += f" ({note})"
                except ValueError:
                    messagebox.showerror("Ошибка", "Некорректная дата")
                    return
            elif tag == "Каптер":
                value = "Каптер"
                if note:
                    value += f" ({note})"
            else:
                value = note if note else ""
            
            listbox.insert(tk.END, f"{tag}: {value}")
            tag_window.destroy()
        
        ttk.Button(tag_window, text="Добавить", command=save_tag).pack(pady=10)
    
    def edit_tag(self, listbox):
        selection = listbox.curselection()
        if not selection:
            return
        
        tag_str = listbox.get(selection[0])
        tag, value = tag_str.split(": ", 1)
        
        tag_window = tk.Toplevel(self.root)
        tag_window.title("Изменить метку")
        tag_window.geometry("400x300")
        
        ttk.Label(tag_window, text="Тип метки:").pack(pady=5)
        tag_combo = ttk.Combobox(tag_window, values=["Сборник", "РПК", "Больной", "Рапорт", "Каптер"])
        tag_combo.pack(pady=5)
        tag_combo.set(tag)
        
        # Для меток "Больной" и "Рапорт" добавляем выбор дат
        date_frame = ttk.Frame(tag_window)
        date_frame.pack(pady=5)
        
        note = ""
        if tag in ["Больной", "Рапорт"]:
            parts = value.split(" (")
            date_part = parts[0]
            if len(parts) > 1:
                note = parts[1].rstrip(")")
            
            start_date, end_date = date_part.split(" - ")
            start_year_val, start_month_val, start_day_val = start_date.split("-")
            end_year_val, end_month_val, end_day_val = end_date.split("-")
            
            ttk.Label(date_frame, text="С:").grid(row=0, column=0, padx=5)
            start_day = ttk.Spinbox(date_frame, from_=1, to=31, width=3)
            start_day.grid(row=0, column=1, padx=5)
            start_day.set(int(start_day_val))
            
            start_month = ttk.Spinbox(date_frame, from_=1, to=12, width=3)
            start_month.grid(row=0, column=2, padx=5)
            start_month.set(int(start_month_val))
            
            start_year = ttk.Spinbox(date_frame, from_=2023, to=2025, width=5)
            start_year.grid(row=0, column=3, padx=5)
            start_year.set(int(start_year_val))
            
            ttk.Label(date_frame, text="По:").grid(row=1, column=0, padx=5)
            end_day = ttk.Spinbox(date_frame, from_=1, to=31, width=3)
            end_day.grid(row=1, column=1, padx=5)
            end_day.set(int(end_day_val))
            
            end_month = ttk.Spinbox(date_frame, from_=1, to=12, width=3)
            end_month.grid(row=1, column=2, padx=5)
            end_month.set(int(end_month_val))
            
            end_year = ttk.Spinbox(date_frame, from_=2023, to=2025, width=5)
            end_year.grid(row=1, column=3, padx=5)
            end_year.set(int(end_year_val))
        elif tag == "Каптер":
            parts = value.split(" (")
            if len(parts) > 1:
                note = parts[1].rstrip(")")
            ttk.Label(date_frame, text="Для выбранной метки даты не требуются").grid(row=0, column=0, columnspan=4)
        else:
            note = value
            ttk.Label(date_frame, text="Для выбранной метки даты не требуются").grid(row=0, column=0, columnspan=4)
        
        # Примечание
        ttk.Label(tag_window, text="Примечание:").pack(pady=5)
        note_entry = ttk.Entry(tag_window, width=30)
        note_entry.pack(pady=5)
        note_entry.insert(0, note)
        
        def save_changes():
            new_tag = tag_combo.get()
            
            if new_tag in ["Больной", "Рапорт"]:
                try:
                    new_start_date = f"{start_year.get()}-{start_month.get():02d}-{start_day.get():02d}"
                    new_end_date = f"{end_year.get()}-{end_month.get():02d}-{end_day.get():02d}"
                    new_value = f"{new_start_date} - {new_end_date}"
                    new_note = note_entry.get()
                    if new_note:
                        new_value += f" ({new_note})"
                except ValueError:
                    messagebox.showerror("Ошибка", "Некорректная дата")
                    return
            elif new_tag == "Каптер":
                new_value = "Каптер"
                new_note = note_entry.get()
                if new_note:
                    new_value += f" ({new_note})"
            else:
                new_value = note_entry.get()
            
            listbox.delete(selection[0])
            listbox.insert(selection[0], f"{new_tag}: {new_value}")
            tag_window.destroy()
        
        ttk.Button(tag_window, text="Сохранить", command=save_changes).pack(pady=10)
    
    def remove_tag(self, listbox):
        selection = listbox.curselection()
        if selection:
            listbox.delete(selection[0])
    
    def save_cadet_changes(self, cadet, new_name, platoon, rank, custom_duty_count, duty_vars, tags_listbox, window):
        # Обновляем данные курсанта
        cadet.name = new_name
        cadet.platoon = platoon
        cadet.rank = rank
        cadet.custom_duty_count = int(custom_duty_count) if custom_duty_count else None
        
        # Обновляем доступные наряды
        cadet.available_duties = [duty_name for duty_name, var in duty_vars.items() if var.get() == 1]
        
        # Обновляем метки
        cadet.tags = []
        for i in range(tags_listbox.size()):
            tag_str = tags_listbox.get(i)
            tag, value = tag_str.split(": ", 1)
            cadet.tags.append((tag, value))
        
        self.update_cadets_lists()
        self.save_data()  # Сохраняем изменения
        window.destroy()
    
    def add_platoon(self):
        platoon_name = self.platoon_entry.get()
        if platoon_name and platoon_name not in self.platoons:
            self.platoons.append(platoon_name)
            self.update_platoon_list()
            self.platoon_entry.delete(0, tk.END)
    
    def edit_platoon(self):
        selection = self.platoon_listbox.curselection()
        if not selection:
            return
        
        old_name = self.platoon_listbox.get(selection[0])
        new_name = simpledialog.askstring("Изменить взвод", "Введите новое название взвода:", initialvalue=old_name)
        
        if new_name and new_name != old_name:
            if new_name in self.platoons:
                messagebox.showerror("Ошибка", "Взвод с таким названием уже существует")
                return
            
            self.platoons[self.platoons.index(old_name)] = new_name
            
            # Обновляем взвод у курсантов
            for cadet in self.cadets:
                if cadet.platoon == old_name:
                    cadet.platoon = new_name
            
            self.update_platoon_list()
            self.update_cadets_lists()
            self.save_data()  # Сохраняем изменения
    
    def update_average_duties(self, category):
        spin = getattr(self, f"avg_duties_{category.lower().replace(' ', '_')}")
        try:
            self.settings["average_duties"][category] = int(spin.get())
            self.save_data()  # Сохраняем изменения
        except ValueError:
            pass
    
    def update_interval(self):
        try:
            self.settings["duty_interval"] = int(self.interval_spin.get())
            self.save_data()  # Сохраняем изменения
        except ValueError:
            pass
    
    def update_course(self):
        self.settings["course"] = self.course_combo.get()
        self.save_data()  # Сохраняем изменения
    
    def add_duty(self):
        add_window = tk.Toplevel(self.root)
        add_window.title("Добавить наряд")
        add_window.geometry("500x400")
        
        ttk.Label(add_window, text="Название наряда:").pack(pady=5)
        name_entry = ttk.Entry(add_window, width=30)
        name_entry.pack(pady=5)
        
        ttk.Label(add_window, text="Требуемое количество:").pack(pady=5)
        reqs_frame = ttk.Frame(add_window)
        reqs_frame.pack(pady=5)
        
        reqs_entries = {}
        for i, category in enumerate(self.categories):
            ttk.Label(reqs_frame, text=category).grid(row=i, column=0, padx=5, pady=2, sticky="e")
            spin = ttk.Spinbox(reqs_frame, from_=0, to=10, width=3)
            spin.grid(row=i, column=1, padx=5, pady=2, sticky="w")
            spin.set(0)
            reqs_entries[category] = spin
        
        # Галочка "Все курсанты могут заступать"
        self.all_cadets_var = tk.IntVar()
        all_cadets_cb = ttk.Checkbutton(add_window, 
                                       text="Все курсанты категории могут заступать", 
                                       variable=self.all_cadets_var,
                                       command=lambda: self.toggle_all_cadets_for_duty(self.all_cadets_var.get()))
        all_cadets_cb.pack(pady=5)
        
        # Дни заступления
        ttk.Label(add_window, text="Дни заступления (через пробел):").pack(pady=5)
        days_entry = ttk.Entry(add_window, width=30)
        days_entry.pack(pady=5)
        
        # Галочка "Каждый день"
        self.every_day_var = tk.IntVar()
        every_day_cb = ttk.Checkbutton(add_window, text="Каждый день", variable=self.every_day_var)
        every_day_cb.pack(pady=5)
        
        def save_duty():
            name = name_entry.get()
            if not name:
                messagebox.showerror("Ошибка", "Введите название наряда")
                return
            
            requirements = {}
            has_requirements = False
            for category, spin in reqs_entries.items():
                count = int(spin.get())
                if count > 0:
                    requirements[category] = count
                    has_requirements = True
            
            if not has_requirements:
                messagebox.showerror("Ошибка", "Укажите хотя бы одну категорию с количеством > 0")
                return
            
            # Получаем дни заступления
            duty_days = []
            if self.every_day_var.get():
                duty_days = list(range(1, 32))  # Максимальное количество дней
            else:
                days_str = days_entry.get()
                if days_str:
                    try:
                        duty_days = [int(d) for d in days_str.split()]
                    except ValueError:
                        messagebox.showerror("Ошибка", "Некорректный формат дней (должны быть числа через пробел)")
                        return
            
            self.duties.append(Duty(
                name, 
                requirements, 
                bool(self.all_cadets_var.get()),
                duty_days,
                bool(self.every_day_var.get())
            ))
            
            # Если выбрано "Все курсанты могут заступать", добавляем этот наряд всем курсантам
            if self.all_cadets_var.get():
                for cadet in self.cadets:
                    for category in requirements.keys():
                        if cadet.category == category and name not in cadet.available_duties:
                            cadet.available_duties.append(name)
            
            self.update_duties_list()
            self.update_duty_selection()
            self.save_data()  # Сохраняем изменения
            add_window.destroy()
        
        btn_frame = ttk.Frame(add_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Сохранить", command=save_duty).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Отмена", command=add_window.destroy).pack(side=tk.RIGHT, padx=10)
    
    def toggle_all_cadets_for_duty(self, state):
        if state:
            if messagebox.askyesno("Подтверждение", 
                                 "Вы уверены, что хотите разрешить всем курсантам заступать в этот наряд?\n"
                                 "Это действие можно будет изменить вручную для каждого курсанта."):
                return True
            else:
                self.all_cadets_var.set(0)
                return False
        return True
    
    def edit_duty(self):
        selection = self.duties_listbox.curselection()
        if not selection:
            return
        
        duty_index = selection[0]
        duty = self.duties[duty_index]
        
        edit_window = tk.Toplevel(self.root)
        edit_window.title(f"Редактирование: {duty.name}")
        edit_window.geometry("500x400")
        
        ttk.Label(edit_window, text="Название наряда:").pack(pady=5)
        name_entry = ttk.Entry(edit_window, width=30)
        name_entry.pack(pady=5)
        name_entry.insert(0, duty.name)
        
        ttk.Label(edit_window, text="Требуемое количество:").pack(pady=5)
        reqs_frame = ttk.Frame(edit_window)
        reqs_frame.pack(pady=5)
        
        reqs_entries = {}
        for i, category in enumerate(self.categories):
            ttk.Label(reqs_frame, text=category).grid(row=i, column=0, padx=5, pady=2, sticky="e")
            spin = ttk.Spinbox(reqs_frame, from_=0, to=10, width=3)
            spin.grid(row=i, column=1, padx=5, pady=2, sticky="w")
            spin.set(duty.requirements.get(category, 0))
            reqs_entries[category] = spin
        
        # Галочка "Все курсанты могут заступать"
        self.all_cadets_var = tk.IntVar(value=int(duty.all_cadets))
        all_cadets_cb = ttk.Checkbutton(edit_window, 
                                       text="Все курсанты категории могут заступать", 
                                       variable=self.all_cadets_var,
                                       command=lambda: self.toggle_all_cadets_for_duty(self.all_cadets_var.get()))
        all_cadets_cb.pack(pady=5)
        
        # Дни заступления
        ttk.Label(edit_window, text="Дни заступления (через пробел):").pack(pady=5)
        days_entry = ttk.Entry(edit_window, width=30)
        days_entry.pack(pady=5)
        days_entry.insert(0, " ".join(map(str, duty.duty_days)))
        
        # Галочка "Каждый день"
        self.every_day_var = tk.IntVar(value=int(duty.every_day))
        every_day_cb = ttk.Checkbutton(edit_window, text="Каждый день", variable=self.every_day_var)
        every_day_cb.pack(pady=5)
        
        def save_changes():
            name = name_entry.get()
            if not name:
                messagebox.showerror("Ошибка", "Введите название наряда")
                return
            
            requirements = {}
            has_requirements = False
            for category, spin in reqs_entries.items():
                count = int(spin.get())
                if count > 0:
                    requirements[category] = count
                    has_requirements = True
            
            if not has_requirements:
                messagebox.showerror("Ошибка", "Укажите хотя бы одну категорию с количеством > 0")
                return
            
            # Получаем дни заступления
            duty_days = []
            if self.every_day_var.get():
                duty_days = list(range(1, 32))  # Максимальное количество дней
            else:
                days_str = days_entry.get()
                if days_str:
                    try:
                        duty_days = [int(d) for d in days_str.split()]
                    except ValueError:
                        messagebox.showerror("Ошибка", "Некорректный формат дней (должны быть числа через пробел)")
                        return
            
            # Обновляем данные наряда
            old_name = duty.name
            duty.name = name
            duty.requirements = requirements
            duty.all_cadets = bool(self.all_cadets_var.get())
            duty.duty_days = duty_days
            duty.every_day = bool(self.every_day_var.get())
            
            # Если изменилось название наряда, обновляем у курсантов
            if old_name != name:
                for cadet in self.cadets:
                    if old_name in cadet.available_duties:
                        cadet.available_duties.remove(old_name)
                        cadet.available_duties.append(name)
            
            # Если выбрано "Все курсанты могут заступать", добавляем этот наряд всем курсантам
            if self.all_cadets_var.get():
                for cadet in self.cadets:
                    for category in requirements.keys():
                        if cadet.category == category and name not in cadet.available_duties:
                            cadet.available_duties.append(name)
            
            self.update_duties_list()
            self.update_duty_selection()
            self.save_data()  # Сохраняем изменения
            edit_window.destroy()
        
        btn_frame = ttk.Frame(edit_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Сохранить", command=save_changes).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Отмена", command=edit_window.destroy).pack(side=tk.RIGHT, padx=10)
    
    def remove_duty(self):
        selection = self.duties_listbox.curselection()
        if not selection:
            return
        
        duty_index = selection[0]
        duty = self.duties[duty_index]
        
        if messagebox.askyesno("Подтверждение", f"Удалить наряд '{duty.name}'?"):
            # Удаляем этот наряд из списка доступных у всех курсантов
            for cadet in self.cadets:
                if duty.name in cadet.available_duties:
                    cadet.available_duties.remove(duty.name)
            
            self.duties.pop(duty_index)
            self.update_duties_list()
            self.update_duty_selection()
            self.save_data()  # Сохраняем изменения
    
    def update_days_count(self, event=None):
        try:
            month = int(self.month_combo.get())
            year = int(self.year_combo.get())
            _, days = monthrange(year, month)
            self.days_count_spin.config(to=days)
            self.days_count_spin.set(days)
        except ValueError:
            pass
    
    def toggle_duty_platoon(self):
        if self.duty_platoon_var.get():
            self.duty_platoon_frame.grid(row=2, column=0, columnspan=7, padx=5, pady=5, sticky="ew")
        else:
            self.duty_platoon_frame.grid_forget()
    
    def add_duty_platoon(self):
        platoon = self.duty_platoon_combo.get()
        week = self.duty_platoon_week_combo.get()
        days = self.duty_platoon_days_entry.get()
        
        if not platoon or not week or not days:
            messagebox.showerror("Ошибка", "Заполните все поля: взвод, неделя и дни")
            return
        
        self.duty_platoon_listbox.insert(tk.END, f"Взвод {platoon} - Неделя {week} - Дни: {days}")
        self.duty_platoon_combo.set("")
        self.duty_platoon_week_combo.set("")
        self.duty_platoon_days_entry.delete(0, tk.END)
    
    def generate_distribution(self):
        try:
            month = int(self.month_combo.get())
            year = int(self.year_combo.get())
            first_day = int(self.first_day_spin.get())
            days_count = int(self.days_count_spin.get())
        except ValueError:
            messagebox.showerror("Ошибка", "Проверьте правильность введенных параметров")
            return
        
        # Проверка наличия курсантов и нарядов
        if not self.cadets:
            messagebox.showerror("Ошибка", "Нет данных о курсантах")
            return
        
        if not self.duties:
            messagebox.showerror("Ошибка", "Нет данных о нарядах")
            return
        
        # Проверяем наличие каптеров
        kapter_cadets = [c for c in self.cadets if any(tag[0] == "Каптер" for tag in c.tags)]
        if not kapter_cadets and any(d.name == "Каптер" for d in self.duties):
            messagebox.showerror("Ошибка", "Нет курсантов с меткой 'Каптер' для наряда 'Каптер'")
            return
        
        # Собираем информацию о днях для каждого наряда
        duty_days = {}
        for duty in self.duties:
            if duty.every_day:
                duty_days[duty.name] = list(range(1, days_count + 1))
            else:
                duty_days[duty.name] = duty.duty_days
        
        # Создаем распределение
        distribution = {}
        current_date = datetime(year, month, first_day).date()
        problems = []  # Для хранения проблемных мест
        
        # Сначала распределяем каптеров (если есть)
        kapter_duty = next((d for d in self.duties if d.name == "Каптер"), None)
        if kapter_duty:
            tuesdays = []
            for day in range(days_count):
                date = current_date + timedelta(days=day)
                if date.weekday() == 1:  # Вторник
                    tuesdays.append(day + 1)
            
            # Проверяем, есть ли дни для каптеров
            if not kapter_duty.every_day and not any(day in kapter_duty.duty_days for day in tuesdays):
                messagebox.showerror("Ошибка", "Нет вторников в выбранных днях для наряда 'Каптер'")
                return
            
            # Распределяем каптеров по вторникам
            kapter_index = 0
            for day in range(days_count):
                date = current_date + timedelta(days=day)
                if (day + 1) in duty_days.get("Каптер", []) and date.weekday() == 1:  # Вторник
                    date_str = date.strftime("%Y-%m-%d")
                    distribution[date_str] = {"Каптер": []}
                    
                    # Выбираем каптера (по очереди)
                    if kapter_index >= len(kapter_cadets):
                        kapter_index = 0
                    
                    cadet = kapter_cadets[kapter_index]
                    distribution[date_str]["Каптер"].append(cadet.name)
                    cadet.add_duty(date, "Каптер")
                    kapter_index += 1
        
        # Затем распределяем остальные наряды
        for day in range(days_count):
            date = current_date + timedelta(days=day)
            date_str = date.strftime("%Y-%m-%d")
            
            # Пропускаем день, если он уже распределен (например, для каптеров)
            if date_str in distribution:
                continue
            
            # Проверяем, есть ли дежурное подразделение на эту неделю
            duty_platoon = None
            if self.duty_platoon_var.get():
                week = (date.day - 1) // 7 + 1
                for i in range(self.duty_platoon_listbox.size()):
                    item = self.duty_platoon_listbox.get(i)
                    if f"Неделя {week}" in item:
                        duty_platoon = item.split("Взвод ")[1].split(" - ")[0]
                        days_str = item.split("Дни: ")[1]
                        duty_days_platoon = [int(d) for d in days_str.split()]
                        if (day + 1) not in duty_days_platoon:
                            duty_platoon = None
                        break
            
            distribution[date_str] = {}
            
            # Распределяем наряды на этот день
            for duty in self.duties:
                # Пропускаем каптеров (они уже распределены)
                if duty.name == "Каптер":
                    continue
                
                # Проверяем, нужно ли распределять этот наряд в этот день
                if (day + 1) not in duty_days.get(duty.name, []):
                    continue
                
                distribution[date_str][duty.name] = []
                
                for category, required in duty.requirements.items():
                    # Если есть дежурное подразделение и это юноши, пропускаем
                    if duty_platoon and category in ["Дежурные юноши", "Дневальные юноши"]:
                        continue
                    
                    # Получаем доступных курсантов
                    if duty.all_cadets:
                        available_cadets = [c for c in self.cadets if c.category == category and c.can_serve(date, duty.name)]
                    else:
                        available_cadets = [
                            c for c in self.cadets 
                            if c.category == category and c.can_serve(date, duty.name) and 
                            (not c.available_duties or duty.name in c.available_duties)
                        ]
                    
                    # Сортируем по количеству нарядов (меньше сначала)
                    available_cadets.sort(key=lambda x: len(x.assigned_duties))
                    
                    # Проверяем, хватает ли курсантов
                    if len(available_cadets) < required:
                        problems.append(f"{date_str}: {duty.name} ({category}) - не хватает {required - len(available_cadets)} курсантов")
                    
                    # Выбираем нужное количество (сколько есть)
                    selected = available_cadets[:required]
                    
                    for cadet in selected:
                        distribution[date_str][duty.name].append(cadet.name)
                        cadet.add_duty(date, duty.name)
        
        # Показываем предупреждения, если есть проблемы
        if problems:
            warning_window = tk.Toplevel(self.root)
            warning_window.title("Предупреждения при генерации")
            warning_window.geometry("600x400")
            
            # Прокрутка для окна с ошибками
            canvas = tk.Canvas(warning_window)
            scrollbar = ttk.Scrollbar(warning_window, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(
                    scrollregion=canvas.bbox("all")
                )
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            ttk.Label(scrollable_frame, text="Обнаружены проблемы при генерации:", font=('Helvetica', 10, 'bold')).pack(pady=5)
            
            for problem in problems:
                ttk.Label(scrollable_frame, text=problem).pack(anchor="w", padx=5, pady=2)
            
            btn_frame = ttk.Frame(scrollable_frame)
            btn_frame.pack(pady=10)
            
            ttk.Button(btn_frame, text="Продолжить", command=lambda: [self.finalize_distribution(distribution), warning_window.destroy()]).pack(side=tk.LEFT, padx=10)
            ttk.Button(btn_frame, text="Отмена", command=warning_window.destroy).pack(side=tk.RIGHT, padx=10)
        else:
            self.finalize_distribution(distribution)
    
    def finalize_distribution(self, distribution):
        self.duty_distribution = distribution
        self.update_view_tab()
        self.save_data()  # Сохраняем изменения
        messagebox.showinfo("Готово", "Распределение успешно сгенерировано")
    
    def edit_distribution(self):
        # Сначала получаем дни, которые уже прошли и не должны изменяться
        today = datetime.now().date()
        past_days = {}
        
        for date_str, duties in self.duty_distribution.items():
            date = datetime.strptime(date_str, "%Y-%m-%d").date()
            if date < today:
                past_days[date_str] = duties
        
        # Создаем окно для выбора дней, которые не нужно изменять
        edit_window = tk.Toplevel(self.root)
        edit_window.title("Редактирование распределения")
        edit_window.geometry("600x400")
        
        ttk.Label(edit_window, text="Выберите дни, которые не нужно изменять:").pack(pady=5)
        
        # Список дней для выбора
        days_listbox = tk.Listbox(edit_window, selectmode=tk.MULTIPLE, height=10)
        scrollbar = ttk.Scrollbar(edit_window, orient="vertical", command=days_listbox.yview)
        days_listbox.configure(yscrollcommand=scrollbar.set)
        
        days_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Заполняем список дней
        days = sorted(self.duty_distribution.items(), key=lambda x: x[0])
        for date_str, duties in days:
            date = datetime.strptime(date_str, "%Y-%m-%d").date()
            day_name = self.weekdays_ru[date.weekday()]
            days_listbox.insert(tk.END, f"{date.strftime('%d.%m.%Y')} ({day_name})")
            
            # Выделяем уже прошедшие дни
            if date_str in past_days:
                days_listbox.selection_set(tk.END)
        
        def proceed_with_editing():
            # Получаем выбранные дни, которые не нужно изменять
            selected_indices = days_listbox.curselection()
            protected_days = {}
            
            for i in selected_indices:
                date_str = days[i][0]
                protected_days[date_str] = self.duty_distribution[date_str]
            
            edit_window.destroy()
            self.perform_distribution_editing(protected_days)
        
        btn_frame = ttk.Frame(edit_window)
        btn_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(btn_frame, text="Продолжить", command=proceed_with_editing).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Отмена", command=edit_window.destroy).pack(side=tk.RIGHT, padx=10)
    
    def perform_distribution_editing(self, protected_days):
        # Создаем новое распределение, сохраняя защищенные дни
        new_distribution = protected_days.copy()
        
        try:
            month = int(self.month_combo.get())
            year = int(self.year_combo.get())
            first_day = int(self.first_day_spin.get())
            days_count = int(self.days_count_spin.get())
        except ValueError:
            messagebox.showerror("Ошибка", "Проверьте правильность введенных параметров")
            return
        
        # Собираем информацию о днях для каждого наряда
        duty_days = {}
        for duty in self.duties:
            if duty.every_day:
                duty_days[duty.name] = list(range(1, days_count + 1))
            else:
                duty_days[duty.name] = duty.duty_days
        
        # Создаем распределение для оставшихся дней
        current_date = datetime(year, month, first_day).date()
        problems = []  # Для хранения проблемных мест
        
        # Сначала распределяем каптеров (если есть)
        kapter_cadets = [c for c in self.cadets if any(tag[0] == "Каптер" for tag in c.tags)]
        kapter_duty = next((d for d in self.duties if d.name == "Каптер"), None)
        
        if kapter_duty and kapter_cadets:
            tuesdays = []
            for day in range(days_count):
                date = current_date + timedelta(days=day)
                date_str = date.strftime("%Y-%m-%d")
                if date.weekday() == 1 and date_str not in protected_days:  # Вторник
                    tuesdays.append((day + 1, date_str))
            
            # Распределяем каптеров по вторникам
            kapter_index = 0
            for day, date_str in tuesdays:
                if (day) in duty_days.get("Каптер", []):
                    new_distribution[date_str] = {"Каптер": []}
                    
                    # Выбираем каптера (по очереди)
                    if kapter_index >= len(kapter_cadets):
                        kapter_index = 0
                    
                    cadet = kapter_cadets[kapter_index]
                    new_distribution[date_str]["Каптер"].append(cadet.name)
                    
                    # Обновляем данные курсанта
                    date = current_date + timedelta(days=day-1)
                    cadet.add_duty(date, "Каптер")
                    kapter_index += 1
        
        # Затем распределяем остальные наряды
        for day in range(days_count):
            date = current_date + timedelta(days=day)
            date_str = date.strftime("%Y-%m-%d")
            
            # Пропускаем день, если он уже распределен или защищен
            if date_str in new_distribution:
                continue
            
            new_distribution[date_str] = {}
            
            # Проверяем, есть ли дежурное подразделение на эту неделю
            duty_platoon = None
            if self.duty_platoon_var.get():
                week = (date.day - 1) // 7 + 1
                for i in range(self.duty_platoon_listbox.size()):
                    item = self.duty_platoon_listbox.get(i)
                    if f"Неделя {week}" in item:
                        duty_platoon = item.split("Взвод ")[1].split(" - ")[0]
                        days_str = item.split("Дни: ")[1]
                        duty_days_platoon = [int(d) for d in days_str.split()]
                        if (day + 1) not in duty_days_platoon:
                            duty_platoon = None
                        break
            
            # Распределяем наряды на этот день
            for duty in self.duties:
                # Пропускаем каптеров (они уже распределены)
                if duty.name == "Каптер":
                    continue
                
                # Проверяем, нужно ли распределять этот наряд в этот день
                if (day + 1) not in duty_days.get(duty.name, []):
                    continue
                
                new_distribution[date_str][duty.name] = []
                
                for category, required in duty.requirements.items():
                    # Если есть дежурное подразделение и это юноши, пропускаем
                    if duty_platoon and category in ["Дежурные юноши", "Дневальные юноши"]:
                        continue
                    
                    # Получаем доступных курсантов
                    if duty.all_cadets:
                        available_cadets = [c for c in self.cadets if c.category == category and c.can_serve(date, duty.name)]
                    else:
                        available_cadets = [
                            c for c in self.cadets 
                            if c.category == category and c.can_serve(date, duty.name) and 
                            (not c.available_duties or duty.name in c.available_duties)
                        ]
                    
                    # Сортируем по количеству нарядов (меньше сначала)
                    available_cadets.sort(key=lambda x: len(x.assigned_duties))
                    
                    # Проверяем, хватает ли курсантов
                    if len(available_cadets) < required:
                        problems.append(f"{date_str}: {duty.name} ({category}) - не хватает {required - len(available_cadets)} курсантов")
                    
                    # Выбираем нужное количество (сколько есть)
                    selected = available_cadets[:required]
                    
                    for cadet in selected:
                        new_distribution[date_str][duty.name].append(cadet.name)
                        cadet.add_duty(date, duty.name)
        
        # Обновляем распределение
        self.duty_distribution = new_distribution
        self.update_view_tab()
        self.save_data()  # Сохраняем изменения
        
        if problems:
            messagebox.showwarning("Предупреждение", 
                                "Распределение обновлено, но обнаружены проблемы:\n\n" + 
                                "\n".join(problems))
        else:
            messagebox.showinfo("Готово", "Распределение успешно обновлено")
    
    def delete_month_stats(self):
        month = simpledialog.askinteger("Удалить месяц", "Введите номер месяца (1-12):", minvalue=1, maxvalue=12)
        year = simpledialog.askinteger("Удалить месяц", "Введите год:", minvalue=2000, maxvalue=2100)
        
        if month and year:
            # Удаляем все записи за указанный месяц
            keys_to_delete = []
            for date_str in self.duty_distribution:
                date = datetime.strptime(date_str, "%Y-%m-%d").date()
                if date.year == year and date.month == month:
                    keys_to_delete.append(date_str)
            
            for key in keys_to_delete:
                del self.duty_distribution[key]
            
            # Обновляем данные курсантов
            for cadet in self.cadets:
                cadet.assigned_duties = [(d, n) for d, n in cadet.assigned_duties if not (d.year == year and d.month == month)]
                # Обновляем last_duty_date
                if cadet.assigned_duties:
                    cadet.last_duty_date = max(d for d, n in cadet.assigned_duties)
                else:
                    cadet.last_duty_date = None
            
            self.update_view_tab()
            self.update_stats()
            self.save_data()  # Сохраняем изменения
            messagebox.showinfo("Готово", f"Данные за {month}.{year} удалены")
    
    def export_to_word(self):
        if not self.duty_distribution:
            messagebox.showerror("Ошибка", "Нет данных для экспорта")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if not file_path:
            return
        
        try:
            doc = docx.Document()
            
            # Заголовок
            doc.add_heading(f"Распределение нарядов. Курс {self.settings['course']}", level=1)
            
            # Группируем по дням
            days = sorted(self.duty_distribution.items(), key=lambda x: x[0])
            
            for date_str, duties in days:
                date = datetime.strptime(date_str, "%Y-%m-%d").date()
                day_name = self.weekdays_ru[date.weekday()]
                
                # Для каптеров определяем период с вторника по среду
                if any(duty_name == "Каптер" for duty_name in duties.keys()):
                    next_day = date + timedelta(days=1)
                    next_day_name = self.weekdays_ru[next_day.weekday()]
                    date_range = f"{date.strftime('%d.%m.%Y')}-{next_day.strftime('%d.%m.%Y')}"
                    
                    # Добавляем заголовок периода
                    doc.add_heading(f"{date_range} ({day_name}-{next_day_name})", level=2)
                    
                    # Добавляем наряды
                    for duty_name, cadets in duties.items():
                        if not cadets:
                            continue
                        
                        # Название наряда
                        doc.add_heading(duty_name, level=3)
                        
                        # Таблица с курсантами
                        table = doc.add_table(rows=1, cols=3)
                        table.style = "Table Grid"
                        
                        # Заголовки таблицы
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = "Должность"
                        hdr_cells[1].text = "Звание"
                        hdr_cells[2].text = "ФИО"
                        
                        for cadet in cadets:
                            cadet_obj = next((c for c in self.cadets if c.name == cadet), None)
                            if not cadet_obj:
                                continue
                            
                            role = "Дежурный" if "Дежурный" in cadet_obj.category else "Дневальный"
                            
                            row_cells = table.add_row().cells
                            row_cells[0].text = role
                            row_cells[1].text = cadet_obj.rank
                            row_cells[2].text = cadet
                else:
                    # Обычный день
                    doc.add_heading(f"{date.strftime('%d.%m.%Y')} ({day_name})", level=2)
                    
                    # Добавляем наряды
                    for duty_name, cadets in duties.items():
                        if not cadets:
                            continue
                        
                        # Название наряда
                        doc.add_heading(duty_name, level=3)
                        
                        # Таблица с курсантами
                        table = doc.add_table(rows=1, cols=3)
                        table.style = "Table Grid"
                        
                        # Заголовки таблицы
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = "Должность"
                        hdr_cells[1].text = "Звание"
                        hdr_cells[2].text = "ФИО"
                        
                        for cadet in cadets:
                            cadet_obj = next((c for c in self.cadets if c.name == cadet), None)
                            if not cadet_obj:
                                continue
                            
                            role = "Дежурный" if "Дежурный" in cadet_obj.category else "Дневальный"
                            
                            row_cells = table.add_row().cells
                            row_cells[0].text = role
                            row_cells[1].text = cadet_obj.rank
                            row_cells[2].text = cadet
            
            doc.save(file_path)
            messagebox.showinfo("Готово", f"Документ сохранен как {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить документ: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = DutyDistributionApp(root)
    root.mainloop()